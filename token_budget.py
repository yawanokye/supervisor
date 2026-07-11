from __future__ import annotations

import io
import math
import os
import re
from dataclasses import dataclass
from typing import Any, Iterable, Mapping, Optional

import fitz
from docx import Document
from sqlalchemy.orm import Session

from .database import ReviewRecord, TokenLedger, User


WORDS_PER_PAGE = max(250, int(os.getenv("TOKEN_ESTIMATE_WORDS_PER_PAGE", "450")))
TOKEN_ACCOUNTING_ENABLED = os.getenv("TOKEN_ACCOUNTING_ENABLED", "true").strip().lower() in {
    "1", "true", "yes", "on"
}
TOKEN_QUOTA_ENFORCEMENT = os.getenv("TOKEN_QUOTA_ENFORCEMENT", "false").strip().lower() in {
    "1", "true", "yes", "on"
}
TOKEN_RESERVE_MULTIPLIER = max(1.0, float(os.getenv("TOKEN_RESERVE_MULTIPLIER", "1.15")))
DEFAULT_SUPERVISOR_TOKENS = max(0, int(os.getenv("SUPERVISOR_DEFAULT_TOKEN_ALLOCATION", "0")))
MAX_ALLOCATION_TOKENS = max(1_000_000, int(os.getenv("MAX_SUPERVISOR_TOKEN_ALLOCATION", "100000000000")))

TOKENS_PER_PAGE = {
    "supervisory_light": max(500, int(os.getenv("TOKENS_PER_PAGE_SUPERVISORY_LIGHT", "2500"))),
    "supervisory_standard": max(500, int(os.getenv("TOKENS_PER_PAGE_SUPERVISORY_STANDARD", "3500"))),
    "supervisory_advanced": max(500, int(os.getenv("TOKENS_PER_PAGE_SUPERVISORY_ADVANCED", "5000"))),
    "external_assessment": max(500, int(os.getenv("TOKENS_PER_PAGE_EXTERNAL_ASSESSMENT", "6500"))),
}
RESERVED_TOKENS_PER_PAGE = {
    key: int(math.ceil(rate * TOKEN_RESERVE_MULTIPLIER))
    for key, rate in TOKENS_PER_PAGE.items()
}


@dataclass(frozen=True)
class PageTokenEstimate:
    pages: int
    tokens_per_page: int
    base_tokens: int
    reserved_tokens: int
    workflow_key: str


def _docx_word_count(data: bytes) -> int:
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return len(re.findall(r"\b\w+[\w'’-]*\b", " ".join(parts), flags=re.UNICODE))


def estimate_document_pages(data: bytes, filename: str) -> int:
    """Estimate source pages without altering or fully reviewing the document.

    PDF pages are exact. DOCX pages use a conservative word-count estimate because
    pagination depends on fonts, spacing, margins and the version of Microsoft Word.
    """
    suffix = (filename or "").lower()
    if suffix.endswith(".pdf"):
        with fitz.open(stream=data, filetype="pdf") as pdf:
            return max(1, int(pdf.page_count))
    if suffix.endswith(".docx"):
        words = _docx_word_count(data)
        return max(1, int(math.ceil(words / WORDS_PER_PAGE)))
    return max(1, int(math.ceil(len(data) / 3000)))


def workflow_key(workflow_type: str, review_depth: str) -> str:
    if workflow_type == "external_assessment":
        return "external_assessment"
    depth = (review_depth or "standard").strip().lower()
    if depth not in {"light", "standard", "advanced"}:
        depth = "standard"
    return f"supervisory_{depth}"


def estimate_review_tokens(
    pages: int,
    workflow_type: str,
    review_depth: str,
) -> PageTokenEstimate:
    key = workflow_key(workflow_type, review_depth)
    rate = TOKENS_PER_PAGE[key]
    page_count = max(1, int(pages))
    base = page_count * rate
    reserved = int(math.ceil(base * TOKEN_RESERVE_MULTIPLIER / 1000.0) * 1000)
    return PageTokenEstimate(
        pages=page_count,
        tokens_per_page=rate,
        base_tokens=base,
        reserved_tokens=reserved,
        workflow_key=key,
    )


def page_capacity(tokens: int, key: str) -> int:
    rate = RESERVED_TOKENS_PER_PAGE[key]
    return max(0, int(tokens or 0) // rate)


def supervisor_capacity(user: User) -> dict[str, int]:
    available = int(user.token_balance or 0)
    return {
        "available_tokens": available,
        "reserved_tokens": int(user.token_reserved_total or 0),
        "used_tokens": int(user.token_used_total or 0),
        "standard_pages": page_capacity(available, "supervisory_standard"),
        "advanced_pages": page_capacity(available, "supervisory_advanced"),
        "external_pages": page_capacity(available, "external_assessment"),
    }


def allocation_to_tokens(amount: int, unit: str) -> int:
    value = max(0, int(amount))
    if unit == "standard_pages":
        return value * RESERVED_TOKENS_PER_PAGE["supervisory_standard"]
    if unit == "external_pages":
        return value * RESERVED_TOKENS_PER_PAGE["external_assessment"]
    if unit != "tokens":
        raise ValueError("Choose tokens, standard supervisory pages or external-examination pages.")
    return value


def _ledger(
    db: Session,
    *,
    user: User,
    amount: int,
    transaction_type: str,
    note: str,
    review_record: Optional[ReviewRecord] = None,
    created_by_user_id: Optional[int] = None,
) -> None:
    db.add(TokenLedger(
        lecturer_id=user.id,
        review_record_id=review_record.id if review_record else None,
        transaction_type=transaction_type,
        amount=int(amount),
        balance_after=int(user.token_balance or 0),
        note=(note or "")[:500] or None,
        created_by_user_id=created_by_user_id,
    ))


def adjust_allocation(
    db: Session,
    *,
    user: User,
    token_amount: int,
    mode: str,
    note: str,
    created_by_user_id: Optional[int],
) -> int:
    token_amount = max(0, int(token_amount))
    if token_amount > MAX_ALLOCATION_TOKENS:
        raise ValueError("The allocation exceeds the configured institutional limit.")
    if mode == "set":
        delta = token_amount - int(user.token_balance or 0)
    elif mode == "add":
        delta = token_amount
    else:
        raise ValueError("Choose Add to balance or Set available balance.")
    user.token_balance = int(user.token_balance or 0) + delta
    user.token_allocated_total = max(
        0,
        int(user.token_allocated_total or 0) + max(0, delta),
    )
    _ledger(
        db,
        user=user,
        amount=delta,
        transaction_type="allocation_set" if mode == "set" else "allocation_add",
        note=note,
        created_by_user_id=created_by_user_id,
    )
    return delta


def should_meter(user: User) -> bool:
    if not TOKEN_ACCOUNTING_ENABLED or user.role != "lecturer":
        return False
    return TOKEN_QUOTA_ENFORCEMENT or any([
        int(user.token_balance or 0),
        int(user.token_reserved_total or 0),
        int(user.token_allocated_total or 0),
        int(user.token_used_total or 0),
    ])


def reserve_review_tokens(
    db: Session,
    *,
    user: User,
    token_amount: int,
    pages: int,
    workflow_label: str,
    job_id: str,
) -> bool:
    if not should_meter(user):
        return False
    required = max(0, int(token_amount))
    available = int(user.token_balance or 0)
    if available < required:
        standard_pages = page_capacity(available, "supervisory_standard")
        external_pages = page_capacity(available, "external_assessment")
        raise ValueError(
            f"Insufficient review-token allocation. This {pages}-page submission is estimated to require "
            f"{required:,} tokens, but {available:,} are available. The remaining allocation supports about "
            f"{standard_pages} standard supervisory pages or {external_pages} external-examination pages."
        )
    user.token_balance = available - required
    user.token_reserved_total = int(user.token_reserved_total or 0) + required
    _ledger(
        db,
        user=user,
        amount=-required,
        transaction_type="review_reservation",
        note=f"Reserved for {workflow_label}: {job_id}",
    )
    return True


def reserve_existing_review_tokens(db: Session, record: ReviewRecord, user: User) -> bool:
    if record.token_accounting_status == "reserved":
        return True
    required = int(record.token_estimate or record.token_reserved or 0)
    if required <= 0 or not should_meter(user):
        return False
    reserve_review_tokens(
        db,
        user=user,
        token_amount=required,
        pages=int(record.estimated_pages or 1),
        workflow_label=(
            "external assessment rebuild"
            if record.workflow_type == "external_assessment"
            else "supervisory review rebuild"
        ),
        job_id=record.job_id,
    )
    record.token_reserved = required
    record.token_accounting_status = "reserved"
    return True


def usage_token_total(ai_review: Optional[Mapping[str, Any]]) -> int:
    if not ai_review:
        return 0
    usages: Iterable[Mapping[str, Any]] = ai_review.get("usage") or []
    total = 0
    for row in usages:
        total += max(0, int(row.get("input_tokens") or 0))
        total += max(0, int(row.get("output_tokens") or 0))
    return total


def settle_review_tokens(
    db: Session,
    *,
    record: ReviewRecord,
    user: User,
    actual_tokens: int,
) -> None:
    if record.token_accounting_status != "reserved":
        return
    reserved = max(0, int(record.token_reserved or 0))
    actual = max(0, int(actual_tokens or 0))
    # A provider may omit usage on a successful response. In that rare case,
    # retain the conservative reservation rather than recording zero usage.
    charged = actual if actual > 0 else reserved
    adjustment = reserved - charged
    user.token_reserved_total = max(0, int(user.token_reserved_total or 0) - reserved)
    user.token_balance = int(user.token_balance or 0) + adjustment
    user.token_used_total = int(user.token_used_total or 0) + charged
    record.tokens_used = charged
    record.token_accounting_status = "settled"
    _ledger(
        db,
        user=user,
        amount=adjustment,
        transaction_type="review_settlement",
        note=(
            f"Settled {record.workflow_type.replace('_', ' ')} for {record.filename}. "
            f"Reserved {reserved:,}; charged {charged:,}."
        ),
        review_record=record,
    )


def release_review_reservation(
    db: Session,
    *,
    record: ReviewRecord,
    user: User,
    note: str,
) -> None:
    if record.token_accounting_status != "reserved":
        return
    reserved = max(0, int(record.token_reserved or 0))
    user.token_reserved_total = max(0, int(user.token_reserved_total or 0) - reserved)
    user.token_balance = int(user.token_balance or 0) + reserved
    record.token_accounting_status = "released"
    _ledger(
        db,
        user=user,
        amount=reserved,
        transaction_type="reservation_release",
        note=note,
        review_record=record,
    )
