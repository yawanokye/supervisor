from __future__ import annotations

import io
import json

import pytest
from docx import Document

from app.academic_ai_engine import _batch_prompt, _valid_issue
from app.annotated_exporter import build_annotated_docx
from app.context_guard import build_context_lock
from app.document_parser import extract_docx


def _docx_bytes(doc: Document) -> bytes:
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def test_table_rows_keep_caption_number_title_and_section_path():
    doc = Document()
    doc.add_heading("CHAPTER FOUR", level=1)
    doc.add_heading("4.2 Regression Results", level=2)
    doc.add_paragraph("Table 4.1 Regression estimates")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Variable"
    table.cell(0, 1).text = "Coefficient"
    table.cell(1, 0).text = "Exchange rate"
    table.cell(1, 1).text = "0.42"

    rows = extract_docx(_docx_bytes(doc))
    table_rows = [row for row in rows if row.get("source_kind") == "table_row"]

    assert table_rows
    assert all(row.get("table_number") == "4.1" for row in table_rows)
    assert all(row.get("table_title") == "Regression estimates" for row in table_rows)
    assert any("4.2 Regression Results" in row.get("section_path", []) for row in table_rows)


def test_annotated_table_comment_names_section_and_table_without_arbitrary_red_text():
    doc = Document()
    doc.add_heading("4.2 Regression Results", level=2)
    doc.add_paragraph("Table 4.1 Regression estimates")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Variable"
    table.cell(0, 1).text = "Coefficient"
    table.cell(1, 0).text = "Exchange rate"
    table.cell(1, 1).text = "0.42"
    source = _docx_bytes(doc)

    rows = extract_docx(source)
    evidence_row = next(
        row for row in rows
        if row.get("source_kind") == "table_row" and row.get("table_row") == 2
    )
    review = {
        "academic_findings": [
            {
                "status": "partly_meets_requirement",
                "section": "4.2 Regression Results",
                "section_reference": "4.2 Regression Results",
                "required_action": "Explain the coefficient and relate the interpretation to the stated objective.",
                "comment": "The coefficient is presented without adequate interpretation.",
                "problematic_quote": "",
                "evidence": [
                    {
                        **evidence_row,
                        "document_role": "current",
                        "is_heading": False,
                    }
                ],
                "headings": ["4.2 Regression Results"],
                "annotation_eligible": True,
            }
        ]
    }

    annotated = build_annotated_docx(source, review)
    out = Document(io.BytesIO(annotated))
    comments = [comment.text for comment in out.comments]

    assert any("4.2 Regression Results, Table 4.1: Regression estimates" in text for text in comments)
    assert any("Explain the coefficient" in text for text in comments)
    assert '[1]' not in out.paragraphs[1].text
    assert 'w:commentRangeStart' in out.element.body.xml


def test_degree_standard_is_independent_of_review_depth():
    section = {
        "section_key": "S001P01",
        "heading": "Introduction",
        "part": 1,
        "paragraphs": [
            {
                "paragraph": 1,
                "text": "This study examines a clearly stated research problem.",
                "heading": "Introduction",
                "is_heading": False,
                "document_role": "current",
                "section_path": ["Introduction"],
            }
        ],
    }

    phd_review = {"summary": {"academic_level": "PhD", "research_approach": "quantitative"}}
    phd_packet = json.loads(_batch_prompt(phd_review, [section], [], {}, "light"))
    assert phd_packet["review_context"]["declared_degree_label"] == "PhD thesis"
    assert "original contribution to knowledge" in phd_packet["review_context"]["review_benchmark"]
    assert phd_packet["review_context"]["review_depth"] == "light"

    bachelors_review = {"summary": {"academic_level": "Bachelors", "research_approach": "quantitative"}}
    bachelors_packet = json.loads(_batch_prompt(bachelors_review, [section], [], {}, "advanced"))
    assert bachelors_packet["review_context"]["declared_degree_label"] == "Bachelor’s dissertation"
    assert "modest but explicit contribution" in bachelors_packet["review_context"]["review_benchmark"]
    assert bachelors_packet["review_context"]["review_depth"] == "advanced"


def test_issue_evidence_cannot_drift_to_another_section():
    paragraphs = [
        {
            "paragraph": 1,
            "text": "The research problem is clearly stated.",
            "heading": "1.2 Statement of the Problem",
            "section_path": ["CHAPTER ONE", "1.2 Statement of the Problem"],
            "document_role": "current",
        },
        {
            "paragraph": 2,
            "text": "The sample comprised 315 respondents.",
            "heading": "3.4 Population and Sample",
            "section_path": ["CHAPTER THREE", "3.4 Population and Sample"],
            "document_role": "current",
        },
    ]
    index = {f"P{row['paragraph']}": row for row in paragraphs}
    lock = build_context_lock(paragraphs, {"academic_level": "Research Masters / MPhil"})
    issue = {
        "finding_id": "F1",
        "category": "methodological_rigour",
        "section": "1.2 Statement of the Problem",
        "issue_title": "Sample justification is unclear",
        "severity": "major",
        "confidence": 0.9,
        "evidence_paragraph_ids": ["P2"],
        "problematic_quote": "The sample comprised 315 respondents.",
        "assessment": "The sample size is stated without a justification.",
        "academic_consequence": "The adequacy of the sample cannot be judged.",
        "required_action": "Provide the sample-size determination and assumptions.",
    }

    assert _valid_issue(
        issue,
        index,
        lock,
        allowed_ids={"P1"},
        canonical_section="1.2 Statement of the Problem",
    ) is None


def test_table_caption_evidence_is_linked_to_the_table_for_annotation():
    doc = Document()
    doc.add_heading("5.3 Measurement Model", level=2)
    doc.add_paragraph("Table 5.2 Construct reliability and validity")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Construct"
    table.cell(0, 1).text = "AVE"
    table.cell(1, 0).text = "Evidence quality"
    table.cell(1, 1).text = "0.61"
    source = _docx_bytes(doc)

    rows = extract_docx(source)
    caption = next(row for row in rows if row.get("source_kind") == "table_caption")
    assert caption.get("table_index") == 1

    review = {
        "academic_findings": [
            {
                "status": "partly_meets_requirement",
                "section": "5.3 Measurement Model",
                "section_reference": "5.3 Measurement Model",
                "required_action": "Explain the basis for accepting the reported validity result.",
                "comment": "The table requires a clearer interpretation.",
                "problematic_quote": "",
                "evidence": [{**caption, "document_role": "current", "is_heading": False}],
                "headings": ["5.3 Measurement Model"],
                "annotation_eligible": True,
            }
        ]
    }

    annotated = build_annotated_docx(source, review)
    out = Document(io.BytesIO(annotated))
    comments = [comment.text for comment in out.comments]
    assert any(
        "5.3 Measurement Model, Table 5.2: Construct reliability and validity" in text
        for text in comments
    )
    assert len(comments) == 1
    assert "w:commentRangeStart" in out.element.body.xml


@pytest.mark.parametrize("academic_level, expected_label", [
    ("Bachelors", "Bachelor’s dissertation"),
    ("Non-Research Masters", "Non-Research Master’s project"),
    ("Research Masters (MPhil)", "Research Master’s or MPhil dissertation"),
    ("Professional Doctorate", "Professional Doctorate thesis"),
    ("PhD", "PhD thesis"),
])
@pytest.mark.parametrize("depth", ["light", "standard", "advanced"])
def test_every_degree_level_keeps_its_own_benchmark_at_every_depth(
    academic_level: str, expected_label: str, depth: str
):
    section = {
        "section_key": "S001P01",
        "heading": "Introduction",
        "part": 1,
        "paragraphs": [{
            "paragraph": 1,
            "text": "The study states its research problem.",
            "heading": "Introduction",
            "is_heading": False,
            "document_role": "current",
            "section_path": ["Introduction"],
        }],
    }
    review = {"summary": {"academic_level": academic_level, "research_approach": "quantitative"}}
    packet = json.loads(_batch_prompt(review, [section], [], {}, depth))

    assert packet["review_context"]["declared_degree_label"] == expected_label
    assert packet["review_context"]["review_depth"] == depth
    assert packet["coverage_contract"]["degree_standard_must_not_change_with_depth"] is True
