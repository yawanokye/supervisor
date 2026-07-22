from __future__ import annotations

import re

from docx import Document

from app.annotated_exporter import _label_existing_source_comments, _preferred_evidence
from app.final_review_quality import build_canonical_finding_rows
from app.human_supervisory_editor import add_human_judgement_findings
from app.natural_supervisor_comment import natural_supervisor_comment
from app.review_release_guard import filter_and_rewrite_release_findings


def _row(paragraph: int, heading: str, text: str, *, is_heading: bool = False):
    return {
        "paragraph": paragraph,
        "paragraph_id": f"P{paragraph}",
        "chapter_number": 1,
        "heading": heading,
        "section_reference": heading,
        "section_path": ["Chapter One", heading],
        "text": text,
        "is_heading": is_heading,
        "document_role": "current",
    }


def _finding(fid: str, paragraph: int, section: str, item: str, comment: str, action: str, category: str):
    text = {
        2: "The main digital capability and operational performance constructs are introduced in broad terms.",
        3: "The practical problem is discussed generally without local evidence or a precise unresolved issue.",
        4: "The scope refers to organisations generally and to one selected institution without defining the population boundary.",
    }[paragraph]
    return {
        "finding_id": fid,
        "status": "partly_meets_requirement",
        "severity": "major",
        "confidence": 0.94,
        "section": section,
        "section_reference": section,
        "category": category,
        "item": item,
        "comment": comment,
        "assessment": comment,
        "required_action": action,
        "problematic_quote": text,
        "annotation_eligible": True,
        "evidence": [{**_row(paragraph, section, text), "document_role": "current"}],
    }


def test_natural_comment_is_three_sentences_without_visible_database_labels():
    row = {
        "item": "The analytical structure is inconsistent",
        "assessment": "The purpose promises one task. The objectives introduce another task. The questions use a third label.",
        "required_action": "Choose one set of tasks and align the purpose, objectives and questions. Use the same terminology throughout.",
        "academic_consequence": "The methodology cannot answer an unstable set of questions.",
    }
    text = natural_supervisor_comment(row)
    sentences = re.findall(r"[^.!?]+[.!?]", text)
    assert 1 <= len(sentences) <= 3
    assert not any(label in text for label in ("Issue:", "Problem identified:", "Action required:", "Verification:"))
    assert "Choose one set of tasks" in text


def test_duplicate_construct_problem_and_scope_families_are_consolidated(monkeypatch):
    monkeypatch.setenv("VPROF_HUMAN_JUDGEMENT_PASS", "false")
    runtime = [
        _row(1, "Background to the Study", "Background to the Study", is_heading=True),
        _row(2, "Background to the Study", "The main digital capability and operational performance constructs are introduced in broad terms."),
        _row(3, "Statement of the Problem", "The practical problem is discussed generally without local evidence or a precise unresolved issue."),
        _row(4, "Scope of the Study", "The scope refers to organisations generally and to one selected institution without defining the population boundary."),
    ]
    findings = [
        _finding("F-CON-1", 2, "Background to the Study", "Undefined central constructs", "The constructs are not defined.", "Define the central constructs.", "construct_definition"),
        _finding("F-CON-2", 3, "Statement of the Problem", "Use of the main construct is undefined", "The same constructs shift across sections.", "Define and use the labels consistently.", "construct_definition"),
        _finding("F-PRO-1", 3, "Statement of the Problem", "Problem not specified", "The problem lacks local evidence.", "State the evidenced practical problem and gap.", "problem_statement"),
        _finding("F-PRO-2", 3, "Statement of the Problem", "Problem statement needs a precise research gap", "The unresolved issue is vague.", "State what previous research has not established.", "problem_statement"),
        _finding("F-SCO-1", 4, "Scope of the Study", "The scope section is incomplete", "The unit and period are absent.", "State the setting, unit, period and exclusions.", "scope_completeness"),
        _finding("F-SCO-2", 4, "Scope of the Study", "The population and institutional scope are inconsistent", "The work moves between a broad population and one case.", "Use one population and setting consistently.", "scope_alignment"),
    ]
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Digital Capability and Operational Performance"},
        "academic_findings": findings,
        "_runtime_context": {"current_paragraphs": runtime},
    }
    rows = build_canonical_finding_rows(review, force=True)
    titles = [row["item"].lower() for row in rows]
    assert sum("construct" in title for title in titles) == 1
    assert sum("problem statement" in title or "problem" in title and "evidence" in title for title in titles) == 1
    assert sum("scope" in title or "population" in title for title in titles) == 1
    assert [row["finding_number"] for row in rows] == list(range(1, len(rows) + 1))


def test_exact_anchor_prefers_substantive_paragraph_over_section_heading():
    heading = _row(10, "Background to the Study", "Background to the Study", is_heading=True)
    paragraph = _row(11, "Background to the Study", "A factual claim is made here without a supporting citation.")
    ordered = _preferred_evidence(
        {"section": "Background to the Study", "section_reference": "Background to the Study"},
        [heading, paragraph],
    )
    assert ordered[0]["paragraph"] == 11


def test_previous_empty_comment_is_removed_and_resolved_missing_section_is_labelled():
    document = Document()
    heading = document.add_paragraph("Purpose of the Study")
    body = document.add_paragraph("The purpose of the study is to examine the stated relationship.")
    document.add_comment(runs=[heading.runs[0]], text="Purpose of the study is missing.", author="Earlier Reviewer", initials="ER")
    document.add_comment(runs=[body.runs[0]], text="", author="Earlier Reviewer", initials="ER")
    _label_existing_source_comments(document)
    comments = list(document.comments)
    assert len(comments) == 1
    assert "appears addressed in the current version" in comments[0].text
    assert "Purpose of the study is missing" in comments[0].text


def test_sparse_organisation_comment_is_not_released_when_outline_is_present():
    runtime = [
        _row(20, "Organisation of the Study", "Organisation of the Study", is_heading=True),
        _row(21, "Organisation of the Study", "Chapter One introduces the study. Chapter Two reviews the literature. Chapter Three presents the methods."),
    ]
    review = {"summary": {"academic_level": "MPhil"}, "_runtime_context": {"current_paragraphs": runtime}}
    row = {
        "finding_id": "AI-ORG-SPARSE",
        "status": "partly_meets_requirement",
        "severity": "minor",
        "section": "Organisation of the Study",
        "section_reference": "Organisation of the Study",
        "item": "Sparse chapter descriptions",
        "comment": "The chapter descriptions are very brief.",
        "required_action": "Expand each chapter description.",
        "evidence": [{**runtime[1], "document_role": "current"}],
    }
    assert filter_and_rewrite_release_findings([row], review) == []


def test_limitations_and_absolute_claim_checks_are_current_submission_specific():
    runtime = [
        _row(1, "Background to the Study", "The proposed platform completely prevents all reporting errors."),
        _row(2, "Limitations of the Study", "The study faced limited time, financial constraints and reluctance of respondents. Measures were taken to ensure validity and reliability."),
    ]
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Digital Reporting and Service Quality"},
        "_runtime_context": {"current_paragraphs": runtime},
    }
    rows = add_human_judgement_findings([], review, ["digital reporting", "service quality"])
    ids = {row.get("finding_id") for row in rows}
    assert "HUMAN-LIMITATIONS-CONSEQUENCES" in ids
    assert "HUMAN-ABSOLUTE-CLAIM" in ids


def test_production_rules_do_not_contain_benchmark_names_or_locations():
    from pathlib import Path

    source = "\n".join(path.read_text(encoding="utf-8", errors="ignore") for path in Path("app").rglob("*.py"))
    for value in ("Aboabo", "Tamale", "Zurkarnain", "Assinman", "Priscilla Boafowaa"):
        assert value.lower() not in source.lower()
