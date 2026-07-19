from __future__ import annotations

import io
import re
import zipfile

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.document_parser import extract_docx
from app.final_review_quality import build_canonical_finding_rows


def _finding(fid: str, paragraph: int, section: str, item: str, comment: str, action: str, *, severity: str = "major"):
    return {
        "finding_id": fid,
        "status": "does_not_meet_requirement",
        "severity": severity,
        "confidence": 0.96,
        "chapter_number": 1,
        "section": section,
        "section_reference": section,
        "category": "other",
        "item": item,
        "comment": comment,
        "required_action": action,
        "evidence": [{
            "document_role": "current",
            "chapter_number": 1,
            "paragraph": paragraph,
            "heading": section,
            "section_reference": section,
            "section_path": ["CHAPTER ONE", section],
            "text": comment,
        }],
        "annotation_eligible": True,
    }


def _runtime():
    return [
        {"paragraph": 1, "chapter_number": 1, "heading": "CHAPTER ONE", "is_heading": True, "text": "CHAPTER ONE", "section_path": ["CHAPTER ONE"]},
        {"paragraph": 2, "chapter_number": 1, "heading": "Introduction", "is_heading": True, "text": "Introduction", "section_path": ["CHAPTER ONE", "Introduction"]},
        {"paragraph": 3, "chapter_number": 1, "heading": "Introduction", "text": "The banking sector is important(Alnaa & Matey, 2024), (Alnaa & Matey, 2024).", "section_path": ["CHAPTER ONE", "Introduction"]},
        {"paragraph": 4, "chapter_number": 1, "heading": "Background of the Study", "is_heading": True, "text": "Background of the Study", "section_path": ["CHAPTER ONE", "Background of the Study"]},
        {"paragraph": 5, "chapter_number": 1, "heading": "Background of the Study", "text": "The study discusses commercial banks and rural banks in Ghana.", "section_path": ["CHAPTER ONE", "Background of the Study"]},
        {"paragraph": 6, "chapter_number": 1, "heading": "Problem Statement", "is_heading": True, "text": "Problem Statement", "section_path": ["CHAPTER ONE", "Problem Statement"]},
        {"paragraph": 7, "chapter_number": 1, "heading": "Problem Statement", "text": "Fraud continues, but the scale of the problem is not stated.", "section_path": ["CHAPTER ONE", "Problem Statement"]},
        {"paragraph": 8, "chapter_number": 1, "heading": "General Objective", "is_heading": True, "text": "General Objective", "section_path": ["CHAPTER ONE", "Research Objectives", "General Objective"]},
        {"paragraph": 9, "chapter_number": 1, "heading": "General Objective", "text": "The primary objective of this study is to examine the effect of internal controls on fraud detection and prevention among rural banks in Ghana, specifically Assinman Rural Bank PLC.", "section_path": ["CHAPTER ONE", "Research Objectives", "General Objective"]},
        {"paragraph": 10, "chapter_number": 1, "heading": "Specific Objectives", "is_heading": True, "text": "Specific Objectives", "section_path": ["CHAPTER ONE", "Research Objectives", "Specific Objectives"]},
        {"paragraph": 11, "chapter_number": 1, "heading": "Specific Objectives", "text": "To assess the internal control mechanisms used for fraud detection. To determine the effectiveness of the controls. To examine the impact of behavioural and institutional factors (Pressure, Opportunity and Rationalisation) on fraud incidence at Assinman Rural Bank PLC.", "section_path": ["CHAPTER ONE", "Research Objectives", "Specific Objectives"]},
        {"paragraph": 12, "chapter_number": 1, "heading": "Research Questions", "is_heading": True, "text": "Research Questions", "section_path": ["CHAPTER ONE", "Research Questions"]},
        {"paragraph": 13, "chapter_number": 1, "heading": "Research Questions", "text": "What is the rate of successful fraud detection and prevention? Which Fraud Triangle factors significantly contribute to fraud occurrence?", "section_path": ["CHAPTER ONE", "Research Questions"]},
        {"paragraph": 14, "chapter_number": 1, "heading": "Scope of the Study", "is_heading": True, "text": "Scope of the Study", "section_path": ["CHAPTER ONE", "Scope of the Study"]},
        {"paragraph": 15, "chapter_number": 1, "heading": "Scope of the Study", "text": "The study uses Assinman Rural Bank PLC as the case study.", "section_path": ["CHAPTER ONE", "Scope of the Study"]},
        {"paragraph": 16, "chapter_number": 1, "heading": "Limitations of the Study", "is_heading": True, "text": "Limitations of the Study", "section_path": ["CHAPTER ONE", "Limitations of the Study"]},
        {"paragraph": 17, "chapter_number": 1, "heading": "Limitations of the Study", "text": "Access may be difficult, but the findings will be generalised to all rural banks.", "section_path": ["CHAPTER ONE", "Limitations of the Study"]},
    ]


def _missing(label: str, fid: str, paragraph: int, section: str, action: str):
    row = _finding(fid, paragraph, section, f"{label} is missing from Chapter One", f"{label} is missing from Chapter One.", action)
    row.update({
        "missing_section_label": label,
        "section_contract_label": label,
        "section_contract_verified": True,
        "section_status": "missing",
    })
    return row


def test_examples_are_clean_and_specific_to_the_bank_study():
    findings = [
        _missing("Purpose of the Study", "M1", 7, "Problem Statement", "Add the purpose after the problem statement."),
        _missing("Research Hypotheses", "M2", 13, "Research Questions", "Add hypotheses after the questions."),
        _missing("Definition of Terms", "M3", 15, "Scope of the Study", "Add definitions before the organisation of the study."),
    ]
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Effect of Internal Controls on Fraud Detection and Prevention at Assinman Rural Bank PLC"},
        "academic_findings": findings,
        "_runtime_context": {"current_paragraphs": _runtime()},
    }
    rows = build_canonical_finding_rows(review, force=True)
    comments = {row.get("missing_section_label"): row["student_comment"] for row in rows if row.get("missing_section_label")}
    purpose = comments["Purpose of the Study"].lower()
    hypothesis = comments["Research Hypotheses"].lower()
    definitions = comments["Definition of Terms"].lower()
    assert "internal controls" in purpose and "assinman rural bank plc" in purpose
    assert "pressure" in purpose and "fraud incidence" in purpose
    assert "h₀" in hypothesis and "do not significantly predict" in hypothesis
    assert "pressure" in hypothesis and "opportunity" in hypothesis and "rationalisation" in hypothesis
    assert "fraud detection and prevention" in definitions
    assert not any(bad in " ".join(comments.values()).lower() for bad in (
        "most significantly contribute, significantly contribute",
        "contribute the occurrence",
        "the study context",
    ))


def test_scope_conflict_and_scope_completeness_are_both_retained():
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Internal Controls and Fraud Prevention at Assinman Rural Bank PLC"},
        "academic_findings": [],
        "_runtime_context": {"current_paragraphs": _runtime()},
    }
    rows = build_canonical_finding_rows(review, force=True)
    ids = {row.get("finding_id") for row in rows}
    assert "HUMAN-SCOPE-CONSISTENCY" in ids
    assert "HUMAN-SCOPE-COMPLETENESS" in ids
    conflict = next(row for row in rows if row.get("finding_id") == "HUMAN-SCOPE-CONSISTENCY")
    assert "commercial banks" in conflict["student_comment"].lower()
    assert "rural banks" in conflict["student_comment"].lower()
    assert "assinman rural bank plc" in conflict["student_comment"].lower()


def test_two_limitations_findings_are_consolidated_into_one_human_comment():
    first = _finding(
        "L1", 17, "Limitations of the Study",
        "The limitations section does not explain its constraints",
        "The section mentions access but does not explain the effect on interpretation.",
        "Identify the design, sampling, measurement and access constraints.",
    )
    second = _finding(
        "L2", 17, "Limitations of the Study",
        "Generalisation is overstated",
        "A single case cannot automatically represent all rural banks.",
        "Use analytical or contextual transferability unless wider inference is supported.",
    )
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Internal Controls and Fraud Prevention at Assinman Rural Bank PLC"},
        "academic_findings": [first, second],
        "_runtime_context": {"current_paragraphs": _runtime()},
    }
    rows = build_canonical_finding_rows(review, force=True)
    limitation_rows = [row for row in rows if "limitation" in (row.get("section_reference") or "").lower()]
    assert len(limitation_rows) == 1
    text = limitation_rows[0]["student_comment"].lower()
    assert "access" in text and "transferability" in text


def test_citation_issue_is_visible_comment_one_and_appendix_is_professional(monkeypatch):
    monkeypatch.setenv("VPROF_COMMENT_MERGE_BY_SECTION", "false")
    monkeypatch.setenv("VPROF_EXPORT_ONE_COMMENT_PER_FINDING", "true")
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("Internal controls support fraud prevention(Amanamah, 2024), (Amanamah, 2024).")
    doc.add_heading("Problem Statement", level=2)
    doc.add_paragraph("The problem is stated generally without local evidence.")
    source = io.BytesIO(); doc.save(source)
    parsed = extract_docx(source.getvalue())
    problem = _finding("P1", 5, "Problem Statement", "The problem needs evidence", "The problem is stated generally.", "Add recent local evidence.")
    problem_target = next(row for row in parsed if "without local evidence" in row.get("text", ""))
    problem["evidence"] = [{**problem_target, "document_role": "current"}]
    review = {
        "summary": {"academic_level": "MPhil", "reviewer_name": "Supervisor", "study_title": "Internal Controls and Fraud Prevention"},
        "academic_findings": [problem],
        "_runtime_context": {"current_paragraphs": parsed},
    }
    output = build_annotated_docx(source.getvalue(), review)
    reviewed = Document(io.BytesIO(output))
    body = "\n".join(paragraph.text for paragraph in reviewed.paragraphs)
    assert "[1]" not in body
    assert "SUPERVISORY REVIEW SUMMARY" in body
    assert "Overall decision:" in body
    assert "Priority corrections" in body
    assert "Specific corrections required" in body
    assert "1. [" in body
    with zipfile.ZipFile(io.BytesIO(output)) as archive:
        comments_xml = archive.read("word/comments.xml").decode("utf-8")
    plain = re.sub(r"<[^>]+>", " ", comments_xml)
    assert re.search(r"\b1\.\s+Issue:\s+The citation presentation", plain)
    assert "missing spaces" in plain
