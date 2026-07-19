from io import BytesIO

from docx import Document

from app.report_exporter import build_docx_report


def test_summary_report_summarises_material_corrections_without_repeating_every_comment():
    findings = []
    for index in range(1, 6):
        findings.append({
            "finding_id": f"F{index}",
            "section": "Statement of the Problem",
            "item": f"Issue {index}",
            "severity": "major" if index < 3 else "moderate",
            "status": "partly_meets_requirement",
            "required_action": f"Required correction number {index}.",
            "comment": f"Comment {index}.",
            "evidence": [{"chapter_number": 1, "paragraph": index, "paragraph_id": f"P{index}", "text": f"Problem sentence {index}.", "document_role": "current"}],
        })

    review = {
        "summary": {
            "filename": "Chapter One.docx",
            "document_label": "Chapter 1",
            "academic_level": "Research Masters / MPhil",
            "research_approach": "quantitative",
            "review_depth": "standard",
            "academic_sections_reviewed": 1,
            "readiness_label": "Major revision required",
            "readiness_meaning": "Revise all material issues.",
            "critical_issues": 0,
            "major_issues": 2,
            "moderate_issues": 3,
            "selected_chapter": 1,
            "current_chapters_detected": [1],
        },
        "overall_academic_assessment": "The chapter requires revision.",
        "academic_strengths": [],
        "academic_section_reviews": [{
            "heading": "Statement of the Problem",
            "section_assessment": "The section requires revision.",
            "section_score": 45,
            "coverage_warning": "",
        }],
        "academic_findings": findings,
        "priority_actions": [],
        "alignment_results": [],
        "revision_results": [],
        "_runtime_context": {"current_paragraphs": [
            {"chapter_number": 1, "paragraph": i, "paragraph_id": f"P{i}", "text": f"Problem sentence {i}.", "document_role": "current"}
            for i in range(1, 6)
        ]},
    }

    document = Document(BytesIO(build_docx_report(review)))
    report_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Numbered comments and detailed corrections" in report_text
    assert "5 sequentially numbered comments" in report_text
    assert "without repeating every comment word for word" in report_text
    assert "Required correction number 1" in report_text
    assert "Required correction number 5" in report_text
    # The report is decision-led. It does not reproduce the full five-row
    # correction register because the exact numbered guidance is in the thesis.
    assert not any(
        len(table.columns) == 4 and table.rows and table.rows[0].cells[0].text == "No."
        for table in document.tables
    )

