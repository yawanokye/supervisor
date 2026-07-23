from __future__ import annotations

import io

from docx import Document

from app.annotated_exporter import (
    build_annotated_docx,
    native_annotation_audit,
)
from app.document_parser import parse_document
from app.inline_annotated_exporter import (
    build_inline_annotated_docx,
    inline_annotation_audit,
)


def _source_with_previous_comment() -> bytes:
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    intro = doc.add_paragraph("This chapter introduces the study and lists the sections that follow.")
    doc.add_comment(intro.runs[0], text="Earlier supervisor comment", author="Previous Supervisor", initials="PS")
    doc.add_heading("Background to the Study", level=2)
    doc.add_paragraph("Digital systems improve inventory control in many organisations.")
    doc.add_heading("Statement of the Problem", level=2)
    doc.add_paragraph("The problem is described generally without evidence from the declared study setting.")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _review(source: bytes) -> dict:
    rows = parse_document(source, "chapter.docx")
    background = next(row for row in rows if "improve inventory" in row.get("text", ""))
    problem = next(row for row in rows if "described generally" in row.get("text", ""))
    return {
        "summary": {"academic_level": "Bachelors", "reviewer_name": "V-Professor"},
        "academic_findings": [
            {
                "finding_id": "A1",
                "status": "does_not_meet_requirement",
                "severity": "major",
                "section": "Background to the Study",
                "section_reference": "Background to the Study",
                "item": "The empirical claim is unsupported",
                "issue_title": "The empirical claim is unsupported",
                "assessment": "The paragraph presents an empirical benefit without a citation.",
                "required_action": "Add a relevant citation or qualify the statement.",
                "evidence": [{**background, "document_role": "current"}],
                "problematic_quote": background["text"],
                "annotation_eligible": True,
            },
            {
                "finding_id": "A2",
                "status": "does_not_meet_requirement",
                "severity": "major",
                "section": "Statement of the Problem",
                "section_reference": "Statement of the Problem",
                "item": "The problem is not evidenced in the declared setting",
                "issue_title": "The problem is not evidenced in the declared setting",
                "assessment": "The problem remains general and is not supported by local evidence.",
                "required_action": "State the practical problem and support it with evidence from the closest applicable context.",
                "evidence": [{**problem, "document_role": "current"}],
                "problematic_quote": problem["text"],
                "annotation_eligible": True,
            },
        ],
        "_runtime_context": {"current_paragraphs": rows},
    }


def test_previous_comments_cannot_make_empty_current_export_pass():
    source = _source_with_previous_comment()
    review = _review(source)
    source_audit = native_annotation_audit(source, review, comment_author="V-Professor")
    assert source_audit["previous_comment_count"] == 0
    assert source_audit["current_comment_count"] == 0
    assert source_audit["passed"] is False


def test_native_and_inline_exports_represent_every_final_finding_number():
    source = _source_with_previous_comment()
    review = _review(source)
    native = build_annotated_docx(source, review, comment_author="V-Professor")
    native_audit = native_annotation_audit(native, review, comment_author="V-Professor")
    assert native_audit["previous_comment_count"] == 1
    assert native_audit["current_comment_count"] >= 1
    assert native_audit["missing_finding_numbers"] == []
    assert native_audit["passed"] is True

    inline = build_inline_annotated_docx(source, review, comment_author="V-Professor")
    inline_audit = inline_annotation_audit(inline, review)
    assert inline_audit["note_count"] >= 1
    assert inline_audit["missing_finding_numbers"] == []
    assert inline_audit["passed"] is True


def test_absolute_claim_fallback_is_not_silently_dropped_from_either_annotation():
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("Background to the Study", level=2)
    doc.add_paragraph(
        "Blockchain eliminates fraud in procurement by recording every transaction in an immutable ledger (Author et al., 2025)."
    )
    stream = io.BytesIO()
    doc.save(stream)
    source = stream.getvalue()
    rows = parse_document(source, "chapter.docx")
    paragraph = next(row for row in rows if "eliminates fraud" in row.get("text", ""))
    review = {
        "summary": {"academic_level": "Bachelors", "reviewer_name": "V-Professor"},
        "academic_findings": [
            {
                "finding_id": "A1",
                "status": "does_not_meet_requirement",
                "severity": "major",
                "section": "Background to the Study",
                "section_reference": "Background to the Study",
                "item": "The absolute claim is too strong",
                "issue_title": "The absolute claim is too strong",
                "assessment": "The claim is stated as certain rather than proportionate to the evidence.",
                "required_action": "Replace the absolute wording with a proportionate expression unless the cited source supports it directly.",
                "student_comment": "The statement ‘eliminates fraud in procurement by recording every transaction in an immutable ledger (Author et al’ is too absolute. Replace it with proportionate wording unless the cited source supports the stronger claim directly.",
                "evidence": [{**paragraph, "document_role": "current"}],
                "problematic_quote": paragraph["text"],
                "annotation_eligible": True,
            }
        ],
        "_runtime_context": {"current_paragraphs": rows},
    }

    native = build_annotated_docx(source, review, comment_author="V-Professor")
    native_audit = native_annotation_audit(native, review, comment_author="V-Professor")
    assert native_audit["missing_finding_numbers"] == []
    assert native_audit["passed"] is True

    inline = build_inline_annotated_docx(source, review, comment_author="V-Professor")
    inline_audit = inline_annotation_audit(inline, review)
    assert inline_audit["missing_finding_numbers"] == []
    assert inline_audit["passed"] is True
