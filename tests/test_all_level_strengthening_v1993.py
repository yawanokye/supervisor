from __future__ import annotations

import io

from docx import Document

from app.comment_quality import prepare_public_issues
from app.deterministic_supervisory_checklist import deterministic_supervisory_checklist_issues
from app.document_parser import parse_document


def _fixture_paragraphs():
    document = Document()
    document.add_heading("CHAPTER ONE", level=0)
    sections = [
        ("1.1 Background of the Study", "Climate finance is important. Many countries face challenges. The study examines transparency."),
        ("1.2 Statement of the Problem", "There is a problem. Prior studies are limited. This study fills the gap."),
        ("1.3 Purpose of the Study", "The purpose is to examine climate action."),
        ("1.4 Research Objectives", "To assess awareness and operational performance among firms."),
        ("1.5 Research Questions", "What is the effect of transparency?"),
        ("1.6 Significance of the Study", "The results reveal that policymakers benefit."),
        ("1.7 Scope of the Study", "The study covers Ghana."),
        ("1.8 Definition of Terms", "Transparency means transparency."),
        ("1.9 Organisation of the Study", "The study contains five chapters."),
    ]
    for heading, text in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(text)
    stream = io.BytesIO()
    document.save(stream)
    return parse_document(stream.getvalue(), "synthetic-chapter-one.docx")


def test_mphil_chapter_one_deterministic_review_has_broad_coverage():
    issues = deterministic_supervisory_checklist_issues(
        _fixture_paragraphs(),
        academic_level="Research Masters / MPhil",
        research_approach="quantitative",
        max_issues=80,
    )
    public, _stats = prepare_public_issues(issues)
    categories = {row.get("category") for row in public}
    assert len(public) >= 6
    assert "theoretical_grounding" in categories
    assert "research_gap_and_problem" in categories
    assert "objectives_questions_hypotheses" in categories
    assert "chapter_structure" in categories
    assert not any("no definitions follow" in (row.get("assessment") or "").lower() for row in public)


def test_all_declared_levels_receive_level_specific_contribution_comment():
    levels = [
        "Bachelors",
        "Non-Research Masters",
        "Research Masters / MPhil",
        "Professional Doctorate",
        "PhD",
    ]
    paragraphs = _fixture_paragraphs()
    for level in levels:
        issues = deterministic_supervisory_checklist_issues(
            paragraphs,
            academic_level=level,
            research_approach="quantitative",
            max_issues=80,
        )
        public, _stats = prepare_public_issues(issues)
        titles = " ".join(row.get("issue_title", "") for row in public).lower()
        assert "contribution" in titles


def test_comment_polishing_removes_malformed_imperatives():
    sample = [{
        "category": "research_gap_and_problem",
        "section": "Statement of the Problem",
        "issue_title": "Local evidence is weak",
        "assessment": "The problem statement is generic.",
        "academic_consequence": "The justification is weak.",
        "required_action": "Revise the marked passage so that it incorporate local evidence from the Central Region.",
        "confidence": 0.9,
        "severity": "major",
    }]
    public, _stats = prepare_public_issues(sample)
    assert public
    assert "so that it incorporate" not in public[0]["required_action"].lower()
    assert "by incorporating" in public[0]["required_action"].lower()
