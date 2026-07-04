from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.document_parser import parse_document
from app.external_assessment import _validate_stage_output
from app.external_assessment_guard import (
    build_document_manifest,
    filter_contradicted_rows,
    find_presence_contradictions,
    find_unsupported_numeric_claims,
    find_unsupported_reference_risk_claims,
    select_balanced_evidence,
)


def _paragraph(
    number: int,
    text: str,
    *,
    chapter: int | None = None,
    heading: str = "",
    is_heading: bool = False,
    marker: int | None = None,
) -> dict:
    return {
        "paragraph": number,
        "chapter_number": chapter,
        "chapter_marker_number": marker,
        "heading": heading,
        "text": text,
        "is_heading": is_heading,
        "source_kind": "paragraph",
    }


def seven_chapter_rows() -> list[dict]:
    detail = (
        " The section provides sustained scholarly explanation, source-based reasoning, "
        "methodological detail and interpretation appropriate to doctoral examination."
    ) * 24
    return [
        _paragraph(1, "FORENSIC INVESTIGATION TECHNIQUES AND PROSECUTION OUTCOMES IN GHANA"),
        _paragraph(2, "BY"),
        _paragraph(3, "EDMUND ANIMLEY"),
        _paragraph(4, "Index No. 10292830"),
        _paragraph(5, "A thesis submitted to the University of Cape Coast for the award of Doctor of Philosophy"),
        _paragraph(6, "CHAPTER ONE", chapter=1, is_heading=True, marker=1),
        _paragraph(7, "INTRODUCTION", chapter=1, heading="Introduction", is_heading=True),
        _paragraph(8, "The statement of the problem, purpose, research objectives, research questions, significance and chapter summary are developed." + detail, chapter=1, heading="Statement of the Problem"),
        _paragraph(9, "CHAPTER TWO", chapter=2, is_heading=True, marker=2),
        _paragraph(10, "CONTEXT OF FINANCIAL CRIME", chapter=2, heading="Context of Financial Crime", is_heading=True),
        _paragraph(11, "This chapter establishes the institutional and legal context for forensic investigation." + detail, chapter=2, heading="Institutional Context"),
        _paragraph(12, "CHAPTER THREE", chapter=3, is_heading=True, marker=3),
        _paragraph(13, "LITERATURE REVIEW", chapter=3, heading="Literature Review", is_heading=True),
        _paragraph(14, "The theoretical framework, empirical review, conceptual model and hypothesis development lead to hypotheses H1 to H9." + detail, chapter=3, heading="Conceptual Model and Hypotheses"),
        _paragraph(15, "CHAPTER FOUR", chapter=4, is_heading=True, marker=4),
        _paragraph(16, "RESEARCH METHODOLOGY", chapter=4, heading="Research Methodology", is_heading=True),
        _paragraph(17, "The positivist research philosophy, explanatory research design, population, sample size, sampling procedure, questionnaire, data collection procedure and ethical considerations are justified." + detail, chapter=4, heading="Research Design and Procedures"),
        _paragraph(18, "CHAPTER FIVE", chapter=5, is_heading=True, marker=5),
        _paragraph(19, "RESULTS", chapter=5, heading="Results", is_heading=True),
        _paragraph(20, "The PLS-SEM results report the measurement model, indicator loadings, reliability, AVE, HTMT, Fornell-Larcker, VIF, structural model, bootstrapping, path coefficients, R squared, effect sizes and PLSpredict. The usable sample was 315 and one reported path coefficient was 0.709." + detail, chapter=5, heading="Measurement and Structural Model Results"),
        _paragraph(21, "CHAPTER SIX", chapter=6, is_heading=True, marker=6),
        _paragraph(22, "DISCUSSION", chapter=6, heading="Discussion", is_heading=True),
        _paragraph(23, "The discussion of findings interprets each hypothesis against theory, prior studies, alternative explanations and study limitations." + detail, chapter=6, heading="Discussion of Findings"),
        _paragraph(24, "CHAPTER SEVEN", chapter=7, is_heading=True, marker=7),
        _paragraph(25, "CONCLUSIONS AND RECOMMENDATIONS", chapter=7, heading="Conclusions and Recommendations", is_heading=True),
        _paragraph(26, "The chapter presents conclusions, recommendations, theoretical contribution, policy implications, limitations and future research." + detail, chapter=7, heading="Contribution and Recommendations"),
        _paragraph(27, "REFERENCES: Mensah, K. (2024). Forensic investigation evidence. Journal of Evidence, 8(2), 1-20. https://doi.org/10.1000/example", is_heading=True),
        _paragraph(28, "APPENDIX A: SURVEY INSTRUMENT", is_heading=True),
        _paragraph(29, "The questionnaire and survey instrument contain the study constructs and response scales." + detail, heading="Questionnaire"),
        _paragraph(30, "APPENDIX B: ETHICAL CLEARANCE", is_heading=True),
        _paragraph(31, "Ethical clearance and informed consent documentation are attached." + detail, heading="Ethical Clearance"),
    ]


def test_edmund_animley_regression_manifest_detects_complete_thesis_structure() -> None:
    rows = seven_chapter_rows()
    manifest = build_document_manifest(
        rows,
        summary={"academic_level": "PhD", "research_approach": "quantitative"},
    )

    assert manifest["detected_chapters"] == [1, 2, 3, 4, 5, 6, 7]
    assert manifest["role_chapters"]["literature_theory"][0] == 3
    assert manifest["role_chapters"]["methodology"][0] == 4
    assert manifest["role_chapters"]["results"][0] == 5
    assert manifest["role_chapters"]["discussion"][0] == 6
    assert manifest["role_chapters"]["conclusions"][0] == 7
    assert manifest["inferred_metadata"]["candidate_name"] == "EDMUND ANIMLEY"
    assert manifest["inferred_metadata"]["candidate_number"] == "10292830"

    for key in (
        "conceptual_model",
        "hypotheses",
        "research_paradigm_and_design",
        "sampling_and_data_collection",
        "measurement_model",
        "structural_model",
        "results",
        "discussion",
        "conclusions",
        "research_instrument",
        "ethics",
        "appendices",
        "references",
    ):
        assert manifest["presence_signals"][key]["status"] == "present", key

    method_keys = {
        item["key"] for item in manifest["method_rubric"]["detected_methods"]
    }
    assert "pls_sem" in method_keys
    assert manifest["source_content_hash"]
    assert manifest["manifest_hash"]


def test_balanced_evidence_reaches_later_results_discussion_and_conclusion_chapters() -> None:
    rows = seven_chapter_rows()
    manifest = build_document_manifest(rows, summary={"academic_level": "PhD"})
    evidence = select_balanced_evidence(
        rows,
        manifest,
        target_roles=("results", "discussion", "conclusions", "ethics"),
        max_chars=26000,
    )
    chapters = {item["chapter_number"] for item in evidence}
    assert {5, 6, 7}.issubset(chapters)
    assert any("measurement model" in item["text"].lower() for item in evidence)
    assert any("discussion of findings" in item["text"].lower() for item in evidence)
    assert any("conclusions" in item["text"].lower() for item in evidence)


def test_absence_claims_are_blocked_when_content_is_present_or_not_confidently_located() -> None:
    manifest = build_document_manifest(
        seven_chapter_rows(),
        summary={"academic_level": "PhD"},
    )
    claims = {
        "finding_1": "Chapter Three is entirely missing.",
        "finding_2": "No ethical clearance is provided.",
        "finding_3": "The candidate number was not supplied.",
    }
    contradictions = find_presence_contradictions(claims, manifest)
    components = {item["component"] for item in contradictions}
    assert {"chapter_3", "ethics", "candidate_number"}.issubset(components)

    safe = find_presence_contradictions(
        {
            "finding_1": "No major concerns were identified in the results.",
            "finding_2": "No results were statistically significant.",
            "finding_3": "No findings supported the null hypothesis.",
        },
        manifest,
    )
    assert safe == []

    sparse_manifest = build_document_manifest(
        [_paragraph(1, "A short introduction to a study", chapter=1)],
        summary={"academic_level": "PhD"},
    )
    uncertain = find_presence_contradictions(
        {"finding": "The measurement model is missing."},
        sparse_manifest,
    )
    assert uncertain
    assert uncertain[0]["manifest_status"] == "not_confidently_located"


def test_derivative_false_missing_findings_are_removed_before_external_review() -> None:
    manifest = build_document_manifest(seven_chapter_rows(), summary={"academic_level": "PhD"})
    rows = [
        {"comment": "Chapter Three is entirely missing."},
        {"comment": "The theoretical synthesis needs sharper comparison of competing explanations."},
    ]
    kept, rejected = filter_contradicted_rows(rows, manifest)
    assert kept == [rows[1]]
    assert len(rejected) == 1


def test_numeric_claims_must_appear_in_the_cited_source_evidence() -> None:
    source_rows = seven_chapter_rows()
    unsupported = {
        "results_or_findings": {
            "evidence_ids": ["P8"],
            "assessment": "The reported path coefficient is 0.709 and the usable sample is 315.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        }
    }
    issues = find_unsupported_numeric_claims(unsupported, source_rows)
    assert {item["token"] for item in issues} == {"0.709", "315"}

    supported = {
        "results_or_findings": {
            "evidence_ids": ["P20"],
            "assessment": "The reported path coefficient is 0.709 and the usable sample is 315.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        }
    }
    assert find_unsupported_numeric_claims(supported, source_rows) == []


def test_high_risk_reference_allegations_require_reference_list_evidence() -> None:
    source_rows = seven_chapter_rows()
    manifest = build_document_manifest(source_rows, summary={"academic_level": "PhD"})
    unsupported = {
        "academic_writing_and_presentation": {
            "evidence_ids": ["P20"],
            "assessment": "The thesis contains a fabricated reference.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        }
    }
    assert find_unsupported_reference_risk_claims(unsupported, manifest)

    supported = {
        "academic_writing_and_presentation": {
            "evidence_ids": ["P27"],
            "assessment": "The cited reference entry requires verification as a potentially fabricated reference.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        }
    }
    assert find_unsupported_reference_risk_claims(supported, manifest) == []


def test_stage_validator_rejects_irrelevant_and_invented_evidence() -> None:
    source_rows = seven_chapter_rows()
    manifest = build_document_manifest(source_rows, summary={"academic_level": "PhD"})
    data = {
        "results_or_findings": {
            "domain": "Results",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P8", "P999"],
            "assessment": "The path coefficient is 0.709.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "discussion_and_interpretation": {
            "domain": "Discussion",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P23"],
            "assessment": "The discussion links findings to theory.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "conclusions_recommendations_and_contribution": {
            "domain": "Conclusions",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P26"],
            "assessment": "The conclusions follow from the findings.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "structural_coherence_and_alignment": {
            "domain": "Coherence",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P20"],
            "assessment": "The chapters are broadly aligned.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "academic_writing_and_presentation": {
            "domain": "Writing",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P20"],
            "assessment": "The writing is generally clear.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "ethics_and_research_integrity": {
            "domain": "Ethics",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P31"],
            "assessment": "Ethical clearance is documented.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "originality_and_contribution": {
            "domain": "Originality",
            "judgement": "appropriate_with_minor_refinement",
            "coverage_status": "fully_assessed",
            "evidence_ids": ["P26"],
            "assessment": "The contribution is stated.",
            "strengths": [],
            "concerns": [],
            "required_corrections": [],
        },
        "major_strengths": [],
        "publication_potential": "Potential after revision.",
    }
    feedback = _validate_stage_output(
        "evidence",
        data,
        manifest=manifest,
        metadata={},
        allowed_ids=manifest["valid_evidence_ids"],
    )
    joined = " ".join(feedback)
    assert "P999" in joined
    assert "outside the manifest's relevant research function" in joined
    assert "0.709" in joined


def test_docx_parser_and_manifest_retain_a_seven_chapter_architecture() -> None:
    doc = Document()
    doc.add_paragraph("FORENSIC INVESTIGATION TECHNIQUES AND PROSECUTION OUTCOMES IN GHANA")
    doc.add_paragraph("BY")
    doc.add_paragraph("EDMUND ANIMLEY")
    doc.add_paragraph("Index No. 10292830")
    toc_style = doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("TABLE OF CONTENTS")
    toc_words = ("ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT")
    for number, words in enumerate(toc_words, start=1):
        paragraph = doc.add_paragraph(
            f"CHAPTER {words} ........................ {number}",
            style=toc_style,
        )
    for number, title in (
        (1, "INTRODUCTION"),
        (2, "CONTEXT OF FINANCIAL CRIME"),
        (3, "LITERATURE REVIEW"),
        (4, "RESEARCH METHODOLOGY"),
        (5, "RESULTS"),
        (6, "DISCUSSION"),
        (7, "CONCLUSIONS AND RECOMMENDATIONS"),
    ):
        words = ("ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN")[number - 1]
        doc.add_paragraph(f"CHAPTER {words}", style="Heading 1")
        doc.add_paragraph(title, style="Heading 2")
        doc.add_paragraph(
            {
                1: "Statement of the problem, objectives and research questions.",
                2: "Institutional context and legal framework.",
                3: "Literature review, theoretical framework, conceptual model and hypotheses.",
                4: "Research philosophy, design, sampling, questionnaire and data collection.",
                5: "Measurement model, structural model, HTMT, bootstrapping and path coefficients.",
                6: "Discussion of findings and interpretation against theory.",
                7: "Conclusions, recommendations, contribution, limitations and future research.",
            }[number]
        )
    doc.add_paragraph("APPENDIX A: QUESTIONNAIRE", style="Heading 1")
    doc.add_paragraph("APPENDIX B: ETHICAL CLEARANCE", style="Heading 1")
    buffer = BytesIO()
    doc.save(buffer)

    parsed = parse_document(buffer.getvalue(), "Edmund Animley Thesis.docx")
    manifest = build_document_manifest(parsed, summary={"academic_level": "PhD"})
    assert manifest["detected_chapters"] == [1, 2, 3, 4, 5, 6, 7]
    assert manifest["toc_reconciliation"]["status"] == "mismatch"
    assert manifest["toc_reconciliation"]["toc_only_chapters"] == [8]
    assert manifest["presence_signals"]["measurement_model"]["status"] == "present"
    assert manifest["presence_signals"]["ethics"]["status"] == "present"


def test_prompt_manifest_does_not_expose_off_stage_evidence_ids() -> None:
    from app.external_assessment import _stage_prompt

    rows = seven_chapter_rows()
    manifest = build_document_manifest(rows, summary={"academic_level": "PhD"})
    # Simulate a manifest-only identifier that is not among the stage excerpts.
    manifest["role_presence"]["foundation"]["evidence_ids"].append("P91")
    manifest["presence_signals"]["chapter_1"]["evidence_ids"].append("P91")
    manifest["toc_reconciliation"]["toc_evidence_ids"].append("P91")
    manifest["inferred_metadata"]["evidence"]["candidate_name"] = "P91"
    prompt = _stage_prompt(
        "foundation",
        {"summary": {"academic_level": "PhD"}},
        {"current_paragraphs": rows},
        {},
        manifest,
    )

    # P91 exists only in manifest bookkeeping and has no supplied source excerpt.
    assert '"P91"' not in prompt
    assert '"toc_evidence_ids"' not in prompt
    assert '"evidence_available"' in prompt


def test_grounding_retry_redacts_bad_token_and_adds_valid_source_excerpt() -> None:
    from app.external_assessment import (
        _feedback_for_prompt,
        _selected_evidence_for_stage,
    )

    rows = seven_chapter_rows()
    manifest = build_document_manifest(rows, summary={"academic_level": "PhD"})

    safe_feedback = _feedback_for_prompt(
        ["Unsupported evidence IDs were used: P31"]
    )
    assert "P31" not in " ".join(safe_feedback)
    assert "allowed_evidence_ids" in safe_feedback[0]

    selected = _selected_evidence_for_stage(
        "foundation",
        {"current_paragraphs": rows},
        manifest,
        concise_retry=False,
        additional_evidence_ids=("P31",),
    )
    by_id = {item["id"]: item for item in selected}
    assert "P31" in by_id
    assert "ethical clearance" in by_id["P31"]["text"].lower()
