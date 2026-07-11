from __future__ import annotations

import io

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.professional_review_pipeline import (
    attach_professional_review_package,
    professional_scope_profile,
)
from app.report_exporter import build_docx_report
from app.review_enrichment import context_specific_example


def _finding(chapter: int, paragraph: int, severity: str = "major") -> dict:
    return {
        "finding_id": f"F-{chapter}-{paragraph}",
        "status": "does_not_meet_requirement",
        "status_label": "Major revision",
        "category": "results_and_interpretation" if chapter == 4 else "methodological_rigour",
        "chapter_number": chapter,
        "section": "Results" if chapter == 4 else "Research Methods",
        "section_reference": "Results" if chapter == 4 else "Research Methods",
        "item": "The reported analysis requires verification",
        "severity": severity,
        "confidence": 0.9,
        "comment": "The reported statistic and interpretation are not sufficiently reconciled.",
        "required_action": "Return to the original software output and rebuild the table from one model run.",
        "illustrative_guidance": "Report the estimate, standard error, test statistic, p-value and confidence interval from the same output.",
        "problematic_quote": "The coefficient was significant",
        "evidence": [{
            "chapter_number": chapter,
            "paragraph": paragraph,
            "section_reference": "Results" if chapter == 4 else "Research Methods",
            "document_role": "current",
            "text": "The coefficient was significant.",
        }],
        "annotation_eligible": True,
        "verification_status": "independent_ai_audit",
    }


def test_scope_profiles_are_professional_and_scope_specific():
    assert professional_scope_profile({"review_scope": "chapter"})["role"] == "Professional chapter supervisor"
    assert professional_scope_profile({"review_scope": "chapter_range"})["role"] == "Senior supervisor and cross-chapter reviewer"
    assert professional_scope_profile({"review_scope": "full_thesis"})["role"] == "Professional thesis examiner"


def test_canonical_finding_ledger_is_sequential_and_updates_rows():
    review = {
        "summary": {"review_scope": "full_thesis", "current_chapters_detected": [3, 4]},
        "academic_findings": [_finding(4, 12, "critical"), _finding(3, 4, "major")],
        "academic_section_reviews": [
            {"chapter_number": 3, "heading": "Research Methods", "section_score": 62},
            {"chapter_number": 4, "heading": "Results", "section_score": 48},
        ],
        "academic_strengths": [],
        "alignment_results": [],
    }
    attach_professional_review_package(review)
    ledger = review["finding_ledger"]
    assert [row["number"] for row in ledger] == [1, 2]
    assert [row["chapter_number"] for row in ledger] == [3, 4]
    assert sorted(row["finding_number"] for row in review["academic_findings"]) == [1, 2]
    assert review["summary"]["professional_reviewer_role"] == "Professional thesis examiner"
    assert review["professional_review"]["quality_controls"]["finding_quota_used"] is False


def test_context_example_does_not_leak_another_study():
    row = {
        "category": "theoretical_grounding",
        "section": "Theoretical Framework",
        "item": "The constructs are not linked to the theory",
        "evidence": [{"text": "Classroom incivility and academic entitlement are expected to influence academic engagement."}],
    }
    example = context_specific_example(row).lower()
    assert "assinman" not in example
    assert "rural bank" not in example
    assert "fraud" not in example
    assert "academic" in example or "incivility" in example or "entitlement" in example


def test_professional_report_uses_examiner_title_and_detailed_findings():
    review = {
        "summary": {
            "review_scope": "full_thesis",
            "filename": "Thesis.docx",
            "document_label": "Complete thesis",
            "academic_level": "PhD",
            "review_depth": "advanced",
            "current_chapters_detected": [3, 4],
            "readiness_label": "Major revision required",
        },
        "academic_findings": [_finding(3, 2), _finding(4, 4, "critical")],
        "academic_section_reviews": [
            {"chapter_number": 3, "heading": "Research Methods", "section_score": 60, "section_assessment": "Methods need strengthening."},
            {"chapter_number": 4, "heading": "Results", "section_score": 45, "section_assessment": "Results need verification."},
        ],
        "academic_strengths": [],
        "alignment_results": [],
        "overall_academic_assessment": "The thesis has a relevant focus but requires major methodological and analytical revision.",
    }
    attach_professional_review_package(review)
    output = build_docx_report(review)
    text = "\n".join(p.text for p in Document(io.BytesIO(output)).paragraphs)
    assert "PROFESSIONAL THESIS EXAMINER’S REPORT" in text
    assert "Detailed Professional Findings and Required Corrections" in text
    assert "Methods, Results and Discussion Accuracy Audit" in text
    assert "Evidence Required for Verification" in text


def test_native_comment_uses_canonical_finding_number():
    document = Document()
    document.add_heading("Results", level=1)
    document.add_paragraph("The coefficient was significant.")
    source = io.BytesIO()
    document.save(source)
    row = _finding(4, 2)
    row["finding_number"] = 7
    review = {"summary": {"reviewer_name": "Examiner"}, "academic_findings": [row]}
    output = build_annotated_docx(source.getvalue(), review)
    comments = list(Document(io.BytesIO(output)).comments)
    assert len(comments) == 1
    assert comments[0].text.startswith("7.")
