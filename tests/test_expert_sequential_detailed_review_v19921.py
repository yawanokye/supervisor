from __future__ import annotations

import io

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.inline_annotated_exporter import build_inline_annotated_docx
from app.professional_review_pipeline import attach_professional_review_package
from app.statistical_review import audit_statistical_consistency
from app.document_parser import extract_docx


def _docx_bytes(doc: Document) -> bytes:
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _finding(paragraph: int, severity: str, quote: str, item: str) -> dict:
    return {
        "status": "does_not_meet_requirement",
        "severity": severity,
        "chapter_number": 1,
        "section": "Background to the Study",
        "section_reference": "Background to the Study",
        "item": item,
        "comment": "The uploaded document does not explain the point at the selected academic level.",
        "academic_consequence": "The weakness prevents the reader from tracing the claim to evidence.",
        "required_action": "Revise the sentence and support it with evidence from the study.",
        "illustrative_guidance": "For example, state the exact construct, population and evidence used in the work.",
        "problematic_quote": quote,
        "evidence": [{
            "chapter_number": 1,
            "paragraph": paragraph,
            "document_role": "current",
            "section_reference": "Background to the Study",
            "text": quote,
        }],
        "annotation_eligible": True,
    }


def test_ledger_numbering_follows_thesis_position_not_severity():
    early = _finding(2, "moderate", "Early sentence.", "Early correction")
    late = _finding(20, "critical", "Late sentence.", "Late correction")
    review = {
        "summary": {"review_scope": "full_thesis", "academic_level": "PhD"},
        "academic_findings": [late, early],
        "alignment_results": [],
        "revision_results": [],
    }
    attach_professional_review_package(review)
    assert [(row["number"], row["issue"].rstrip(".")) for row in review["finding_ledger"]] == [
        (1, "Early correction"),
        (2, "Late correction"),
    ]


def test_native_markers_are_sequential_and_never_split_words():
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_paragraph("Teacher education requires strong academic engagement.")
    doc.add_paragraph("Academic entitlement may weaken participation.")
    source = _docx_bytes(doc)
    rows = extract_docx(source)
    first = next(row for row in rows if "Teacher education" in row.get("text", ""))
    second = next(row for row in rows if "Academic entitlement" in row.get("text", ""))
    review = {
        "summary": {"reviewer_name": "Examiner", "academic_level": "PhD"},
        "academic_findings": [
            {**_finding(second["paragraph"], "critical", "Academic entitlement", "Clarify the second claim"), "evidence": [{**second, "document_role": "current"}], "finding_number": 45},
            {**_finding(first["paragraph"], "moderate", "Teache", "Clarify the first claim"), "evidence": [{**first, "document_role": "current"}], "finding_number": 6},
        ],
    }
    output = Document(io.BytesIO(build_annotated_docx(source, review)))
    body = "\n".join(p.text for p in output.paragraphs)
    comments = [comment.text for comment in output.comments]
    assert "teache [1]r" not in body.lower()
    assert "academi [2]c" not in body.lower()
    assert "[1]" not in output.paragraphs[1].text
    assert "[2]" not in output.paragraphs[2].text
    assert comments[0].startswith("1.")
    assert comments[1].startswith("2.")
    assert "w:commentRangeStart" in output.element.body.xml


def test_public_comment_uses_study_and_phd_level_language():
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_paragraph("The problem requires stronger evidence.")
    source = _docx_bytes(doc)
    evidence = next(row for row in extract_docx(source) if "requires stronger" in row.get("text", ""))
    row = _finding(evidence["paragraph"], "major", "requires stronger evidence", "Strengthen the evidence")
    row["evidence"] = [{**evidence, "document_role": "current"}]
    review = {"summary": {"reviewer_name": "Examiner", "academic_level": "PhD"}, "academic_findings": [row]}
    output = Document(io.BytesIO(build_annotated_docx(source, review)))
    text = list(output.comments)[0].text
    assert "uploaded document" not in text.lower()
    assert "selected academic level" not in text.lower()
    assert "the study" in text.lower()
    assert "At PhD level" not in text


def test_inline_annotation_is_detailed_and_uses_same_number():
    doc = Document()
    doc.add_heading("CHAPTER FOUR", level=1)
    doc.add_paragraph("The coefficient was statistically significant.")
    source = _docx_bytes(doc)
    evidence = next(row for row in extract_docx(source) if "coefficient" in row.get("text", ""))
    row = _finding(evidence["paragraph"], "critical", "The coefficient was statistically significant.", "The result is incompletely reported")
    row.update({
        "chapter_number": 4,
        "section": "Regression Results",
        "section_reference": "Regression Results",
        "comment": "The statement gives significance but omits the estimate, standard error, confidence interval and model context.",
        "academic_consequence": "The reader cannot verify the magnitude, precision or substantive importance of the result.",
        "required_action": "Report and interpret the complete model result from the same software output.",
        "illustrative_guidance": "For example, report B, SE, t, p, the confidence interval, R² and the applicable diagnostic evidence.",
        "evidence": [{**evidence, "chapter_number": 4, "document_role": "current"}],
    })
    review = {"summary": {"academic_level": "PhD"}, "academic_findings": [row]}
    output = Document(io.BytesIO(build_inline_annotated_docx(source, review)))
    body = "\n".join(p.text for p in output.paragraphs)
    assert "[1]" in output.paragraphs[1].text
    assert "Detailed supervisor comment:" in body
    assert "magnitude, precision" in body
    assert "Report and interpret" in body
    assert "For example," in body
    assert "At PhD level" not in body
    assert 'w:val="C00000"' in output.element.xml


def test_statistical_audit_checks_b_se_t_and_r2_f_df_accuracy():
    rows = [
        {
            "text": "B = 0.40, SE = 0.10, t = 2.00, p = .001, 95% CI [.20, .60].",
            "paragraph": 1,
            "chapter_number": 4,
            "heading": "Regression Results",
        },
        {
            "text": "The model reported R² = .41 and F(2, 347) = 48.92, p < .001.",
            "paragraph": 2,
            "chapter_number": 4,
            "heading": "Model Summary",
        },
    ]
    warnings = audit_statistical_consistency(rows)
    kinds = {item["kind"] for item in warnings}
    assert "coefficient_se_t_mismatch" in kinds
    assert "r2_f_df_mismatch" in kinds


def test_export_fallback_placeholder_joins_canonical_sequence_and_report_ledger():
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_paragraph("The study covered [insert actual data-collection period].")
    doc.add_paragraph("The final paragraph needs a clearer contribution statement.")
    source = _docx_bytes(doc)
    evidence = next(row for row in extract_docx(source) if "final paragraph" in row.get("text", ""))
    late = _finding(evidence["paragraph"], "major", "The final paragraph", "Clarify the contribution")
    late["evidence"] = [{**evidence, "chapter_number": 1, "document_role": "current"}]
    review = {
        "summary": {"review_scope": "chapter", "academic_level": "MPhil"},
        "academic_findings": [late],
        "alignment_results": [],
        "revision_results": [],
        "professional_review": {"finding_ledger": [{"number": 99}]},
    }

    build_annotated_docx(source, review)
    assert "professional_review" not in review
    attach_professional_review_package(review)
    ledger = review["finding_ledger"]
    assert [row["number"] for row in ledger] == [1, 2]
    assert "placeholder" in ledger[0]["issue"].lower()
    assert "contribution" in ledger[1]["issue"].lower()
