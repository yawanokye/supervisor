from __future__ import annotations

import io

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.document_parser import extract_docx
from app.final_review_quality import build_canonical_finding_rows
from app.professional_review_pipeline import attach_professional_review_package
from app.report_exporter import build_docx_report


def _finding(fid: str, paragraph: int, *, chapter: int = 1, section: str = "Background to the Study", item: str, comment: str, action: str, quote: str = "", category: str = "other", table_number: str = "", table_title: str = ""):
    evidence = {
        "document_role": "current",
        "chapter_number": chapter,
        "paragraph": paragraph,
        "section_reference": section,
        "section_path": [f"Chapter {chapter}", section],
        "text": quote or item,
    }
    if table_number:
        evidence.update({"table_number": table_number, "table_title": table_title, "table_index": 1})
    return {
        "finding_id": fid,
        "status": "does_not_meet_requirement",
        "severity": "major",
        "confidence": 0.95,
        "chapter_number": chapter,
        "section": section,
        "section_reference": section,
        "category": category,
        "item": item,
        "comment": comment,
        "required_action": action,
        "problematic_quote": quote,
        "evidence": [evidence],
        "annotation_eligible": True,
    }


def test_global_manifest_removes_false_missing_reference_finding():
    review = {
        "summary": {"academic_level": "PhD", "supervisory_document_manifest": {"exact_section_and_subsection_headings": ["REFERENCES"]}},
        "academic_findings": [_finding(
            "F1", 10, section="References", item="References section is missing",
            comment="No reference list is evident in the work.", action="Add a reference list.",
        )],
    }
    assert build_canonical_finding_rows(review, force=True) == []


def test_brevity_alone_does_not_create_purpose_or_question_comment():
    review = {
        "summary": {"academic_level": "MPhil"},
        "academic_findings": [
            _finding("F1", 2, section="Purpose of the Study", item="The purpose section is too brief to perform its purpose adequately", comment="Expand the purpose.", action="Expand the purpose."),
            _finding("F2", 3, section="Research Questions", item="The research questions section needs further development", comment="The section is brief.", action="Expand the research questions."),
        ],
    }
    assert build_canonical_finding_rows(review, force=True) == []


def test_mechanical_level_and_traceability_language_is_removed():
    review = {
        "summary": {"academic_level": "PhD"},
        "academic_findings": [_finding(
            "F1", 2, item="Method alignment is unclear",
            comment="At PhD level, every conclusion should be traceable to the correct table, estimate, uncertainty measure, diagnostic evidence and decision rule.",
            action="Create an audit trail from objective to analysis, result, hypothesis decision and conclusion.",
        )],
    }
    row = build_canonical_finding_rows(review, force=True)[0]
    text = " ".join(str(row.get(k, "")) for k in ("comment", "required_action"))
    assert "At PhD level" not in text
    assert "traceable" not in text.lower()
    assert "linked" in text.lower() or "connection" in text.lower()


def test_chapter_one_background_is_not_held_to_chapter_two_synthesis_standard():
    review = {
        "summary": {"academic_level": "PhD", "study_title": "Influence of Classroom Incivility and Academic Entitlement on Academic Engagement: Moderating Role of Perceived Academic Support"},
        "academic_findings": [_finding(
            "F1", 2, chapter=1, section="Background to the Study",
            item="The background lacks deep critical synthesis",
            comment="The studies are not critically synthesised by differences in method and context.",
            action="Compare every study by methods, contradictions and limitations.",
        )],
    }
    row = build_canonical_finding_rows(review, force=True)[0]
    assert row["item"] == "The background needs a clearer evidence-led progression"
    assert "does not require the full study-by-study critical synthesis" in row["comment"]
    assert "Reserve detailed comparison" in row["required_action"]
    assert "classroom incivility" in row.get("illustrative_guidance", "").lower()


def test_chapter_two_retains_deep_synthesis_requirement():
    review = {
        "summary": {"academic_level": "PhD"},
        "academic_findings": [_finding(
            "F1", 20, chapter=2, section="Empirical Review",
            item="The empirical review is descriptive",
            comment="The studies are summarised one by one and are not critically synthesised.",
            action="Compare the studies by context, design, measures, findings, contradictions and limitations.",
        )],
    }
    row = build_canonical_finding_rows(review, force=True)[0]
    assert row["item"] == "The empirical review is descriptive"
    assert "context, design, measures" in row["required_action"]


def test_same_table_statistical_findings_are_consolidated_before_numbering():
    review = {
        "summary": {"academic_level": "PhD"},
        "academic_findings": [
            _finding("F1", 40, chapter=4, section="Regression Results", table_number="6", table_title="Regression Model", category="statistical_accuracy", item="R squared and F do not reconcile", comment="The printed R squared and F statistic cannot come from the same model.", action="Re-run the model and reproduce one complete model summary."),
            _finding("F2", 40, chapter=4, section="Regression Results", table_number="6", table_title="Regression Model", category="analysis_appropriateness", item="The model omits confidence intervals", comment="The coefficient table is incomplete.", action="Report coefficients, standard errors, confidence intervals and diagnostics from the same output."),
        ],
    }
    rows = build_canonical_finding_rows(review, force=True)
    assert len(rows) == 1
    assert rows[0]["finding_number"] == 1
    assert rows[0]["section_reference"] == "Regression Results, Table 6: Regression Model"
    assert "same output" in rows[0]["required_action"].lower() or "same model" in rows[0]["required_action"].lower()


def test_final_numbers_are_assigned_after_filtering_and_exported_beside_passages(monkeypatch):
    monkeypatch.setenv("VPROF_COMMENT_MERGE_BY_SECTION", "false")
    monkeypatch.setenv("VPROF_EXPORT_ONE_COMMENT_PER_FINDING", "true")
    monkeypatch.setenv("VPROF_NATIVE_GROUP_LOCATION_MARKERS", "true")
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_paragraph("Teacher education requires strong evidence.")
    doc.add_paragraph("Academic entitlement may weaken engagement.")
    source = io.BytesIO(); doc.save(source)
    parsed = extract_docx(source.getvalue())
    first = next(r for r in parsed if "Teacher education" in r.get("text", ""))
    second = next(r for r in parsed if "Academic entitlement" in r.get("text", ""))
    review = {
        "summary": {"academic_level": "PhD", "reviewer_name": "Supervisor"},
        "academic_findings": [
            _finding("LATE", second["paragraph"], item="Clarify the second claim", comment="The relationship is asserted without adequate evidence.", action="Support or qualify the relationship.", quote="Academic entitlement",),
            _finding("EARLY", first["paragraph"], item="Clarify the first claim", comment="The claim is too broad.", action="State the precise evidence required.", quote="Teacher",),
        ],
    }
    review["academic_findings"][0]["evidence"] = [{**second, "document_role": "current"}]
    review["academic_findings"][1]["evidence"] = [{**first, "document_role": "current"}]
    reviewed = Document(io.BytesIO(build_annotated_docx(source.getvalue(), review)))
    assert "[1]" in reviewed.paragraphs[1].text
    assert "[2]" in reviewed.paragraphs[2].text
    assert "teache [1]r" not in "\n".join(p.text for p in reviewed.paragraphs).lower()
    comments = [c.text for c in reviewed.comments]
    assert comments[0].startswith("1.")
    assert comments[1].startswith("2.")


def test_spec_report_is_concise_and_does_not_duplicate_all_findings(monkeypatch):
    monkeypatch.setenv("VPROF_SPEC_ALIGNED_SUPERVISORY_REPORT", "true")
    monkeypatch.setenv("VPROF_REPORT_INCLUDE_DETAILED_FINDINGS", "false")
    review = {
        "summary": {"review_scope": "full_thesis", "academic_level": "PhD", "filename": "Study.docx", "current_chapters_detected": [1, 4]},
        "academic_findings": [
            _finding("F1", 2, item="The problem needs local evidence", comment="The claim is broad.", action="Add verified local evidence."),
            _finding("F2", 20, chapter=4, section="Regression Results", category="statistical_accuracy", item="The model statistics do not reconcile", comment="The values cannot come from one model.", action="Re-run the model."),
        ],
        "academic_strengths": [],
        "academic_section_reviews": [],
        "alignment_results": [],
        "statistical_review": {"consistency_warnings": []},
    }
    attach_professional_review_package(review)
    output = build_docx_report(review)
    text = "\n".join(p.text for p in Document(io.BytesIO(output)).paragraphs)
    assert "Numbered comments and detailed corrections" in text
    assert "deliberately summarises" in text
    assert "Detailed Professional Findings and Required Corrections" not in text
    assert "At PhD level" not in text
