from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.comment_quality import (
    prepare_public_issues,
    public_text,
    sentence_safe_trim,
)
from app.supervisory_accuracy_guard import deterministic_expert_issues


def _docx_bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_student_comment_removes_internal_status_and_generated_placeholders():
    document = Document()
    document.add_heading("CHAPTER ONE", level=1)
    document.add_heading("Statement of the Problem", level=2)
    document.add_paragraph("The regional evidence remains limited.")
    review = {
        "summary": {"reviewer_name": "Supervisor"},
        "academic_findings": [{
            "status": "partly_meets_requirement",
            "category": "citations_and_sources",
            "section": "Statement of the Problem",
            "section_reference": "Statement of the Problem",
            "reference_label": "Statement of the Problem",
            "item": "Regional evidence requires verification",
            "comment": "The claim needs stronger evidence.",
            "required_action": "Verify that [verified scholarly source] supports [X]% of firms.",
            "illustrative_guidance": "EPA reports that [X]% of firms comply.",
            "problematic_quote": "regional evidence",
            "evidence": [{
                "paragraph": 3,
                "text": "The regional evidence remains limited.",
                "section_reference": "Statement of the Problem",
                "document_role": "current",
            }],
            "annotation_eligible": True,
            "manual_confirmation_required": True,
        }],
    }
    output = build_annotated_docx(_docx_bytes(document), review)
    comments = list(Document(io.BytesIO(output)).comments)
    assert len(comments) == 1
    text = comments[0].text
    assert "Manual confirmation recommended" not in text
    assert "[verified" not in text
    assert "[X]" not in text
    assert "Verify the relevant claim" in text


def test_sentence_trimming_never_releases_a_cut_fragment():
    text = (
        "Revise the purpose so it covers both environmental sustainability and operational performance. "
        "Then align the problem statement, objectives and research questions with the revised purpose. "
        "This final sentence should be omitted when the limit is reached."
    )
    shortened = sentence_safe_trim(text, 150)
    assert shortened.endswith(".")
    assert "This final sentence" not in shortened
    assert len(shortened) <= 150


def test_current_year_reference_is_not_labelled_future_dated():
    year = datetime.now(timezone.utc).year
    issues = [{
        "category": "citations_and_sources",
        "section": "Background to the Study",
        "issue_title": "Future-dated references require replacement",
        "severity": "moderate",
        "confidence": 0.9,
        "assessment": f"The source dated {year} is future-dated.",
        "academic_consequence": "The citation may be unreliable.",
        "required_action": "Replace each future-dated reference.",
        "illustrative_guidance": "",
        "evidence_paragraph_ids": ["P2"],
    }]
    cleaned, stats = prepare_public_issues(issues)
    assert cleaned == []
    assert stats["dropped"] == 1


def test_repeated_hypothesis_comments_are_consolidated_and_conditional():
    base = {
        "category": "objectives_questions_hypotheses",
        "section": "Research Objectives",
        "severity": "moderate",
        "confidence": 0.86,
        "assessment": "Objectives 2 and 4 are inferential but no hypotheses are stated.",
        "academic_consequence": "The inferential expectations are not explicit.",
        "illustrative_guidance": "",
    }
    issues = [
        {
            **base,
            "issue_title": "Hypotheses are absent for inferential objectives",
            "required_action": "Develop formal hypotheses corresponding to Objectives 2 and 4.",
            "evidence_paragraph_ids": ["P20"],
        },
        {
            **base,
            "issue_title": "Testable hypotheses should accompany the objectives",
            "required_action": "Formulate testable hypotheses for Objectives 2 and 4 and state them after the research questions.",
            "evidence_paragraph_ids": ["P21"],
        },
    ]
    cleaned, _ = prepare_public_issues(issues)
    assert len(cleaned) == 1
    assert "Where required by the programme's thesis format" in cleaned[0]["required_action"]
    assert cleaned[0]["evidence_paragraph_ids"] == ["P20", "P21"]


def test_document_placeholders_and_stage_tense_are_detected_deterministically():
    paragraphs = [
        {
            "document_role": "current", "paragraph": 1, "chapter_number": 1,
            "heading": "Limitations of the Study", "section_path": ["Limitations of the Study"],
            "text": "The data will be obtained at a single point in time.",
        },
        {
            "document_role": "current", "paragraph": 2, "chapter_number": 1,
            "heading": "Limitations of the Study", "section_path": ["Limitations of the Study"],
            "text": "The study faced practical financial and time constraints.",
        },
        {
            "document_role": "current", "paragraph": 3, "chapter_number": 1,
            "heading": "Delimitation of the Study", "section_path": ["Delimitation of the Study"],
            "text": "Data will be collected between [insert start month/year] and [insert end month/year].",
        },
    ]
    issues = deterministic_expert_issues(paragraphs)
    ids = {item["finding_id"] for item in issues}
    assert "DET-UNRESOLVED-DRAFT-PLACEHOLDERS" in ids
    assert "DET-CHAPTER-ONE-TENSE-MISMATCH" in ids


def test_public_text_rejects_any_unresolved_bracket_prompt():
    assert public_text("Use [insert verified source] to support the claim.") == ""


def test_incomplete_illustrative_guidance_is_omitted_not_exported():
    issues = [{
        "category": "cross_section_coherence",
        "section": "Purpose of the Study",
        "issue_title": "Purpose and objective require alignment",
        "severity": "major",
        "confidence": 0.92,
        "assessment": "The purpose omits operational performance although it appears in Objective 4.",
        "academic_consequence": "The chapter does not present one coherent scope.",
        "required_action": "Revise the purpose and Objective 4 together so the intended outcomes are stated consistently.",
        "illustrative_guidance": "The purpose of this study is to examine the effect of green procurement practices on",
        "evidence_paragraph_ids": ["P19", "P23"],
    }]
    cleaned, _ = prepare_public_issues(issues)
    assert len(cleaned) == 1
    assert cleaned[0]["illustrative_guidance"] == ""
    assert cleaned[0]["required_action"].endswith(".")


def test_internal_audit_coverage_warning_is_removed_from_public_text():
    warning = (
        "One or more independent accuracy-audit batches were unavailable; "
        "the displayed comments passed exact evidence and placement checks "
        "but are marked for manual confirmation."
    )
    assert public_text(warning) == ""


def test_native_export_consolidates_duplicate_alignment_comments():
    document = Document()
    document.add_heading("Purpose of the Study", level=2)
    document.add_paragraph(
        "The purpose of this study is to examine the effect of green procurement practices on environmental sustainability."
    )
    base = {
        "status": "partly_meets_requirement",
        "category": "cross_section_coherence",
        "section": "Purpose of the Study",
        "section_reference": "Purpose of the Study",
        "reference_label": "Purpose of the Study",
        "severity": "major",
        "confidence": 0.9,
        "annotation_eligible": True,
        "evidence": [{
            "paragraph": 2,
            "text": "The purpose of this study is to examine the effect of green procurement practices on environmental sustainability.",
            "section_reference": "Purpose of the Study",
            "document_role": "current",
        }],
    }
    review = {
        "summary": {"reviewer_name": "Supervisor"},
        "academic_findings": [{
            **base,
            "item": "Purpose and Objective 4 are misaligned",
            "comment": "Objective 4 introduces operational performance, which the purpose omits.",
            "required_action": "Revise the purpose to include operational performance or remove Objective 4.",
        }],
        "alignment_results": [{
            **base,
            "item": "Operational performance is absent from the purpose",
            "comment": "The purpose omits the outcome introduced in Objective 4.",
            "required_action": "Revise the purpose to include operational performance or remove Objective 4 to maintain alignment.",
        }],
    }
    output = build_annotated_docx(_docx_bytes(document), review)
    comments = list(Document(io.BytesIO(output)).comments)
    assert len(comments) == 1
