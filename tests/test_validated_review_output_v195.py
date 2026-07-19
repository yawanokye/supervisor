from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from app.annotated_exporter import build_annotated_docx, native_comment_count
from app.ai_config import HybridAIConfig


def _docx_bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_manual_confirmation_fallback_still_exports_native_comment():
    document = Document()
    document.add_heading("CHAPTER ONE", level=1)
    document.add_heading("Statement of the Problem", level=2)
    document.add_paragraph("The study will prove that green procurement causes environmental sustainability.")
    source = _docx_bytes(document)
    review = {
        "summary": {"reviewer_name": "Anokye Mohammed Adam"},
        "academic_findings": [{
            "status": "does_not_meet_requirement",
            "section": "Statement of the Problem",
            "section_reference": "Statement of the Problem",
            "reference_label": "Statement of the Problem",
            "required_action": "Replace the unsupported causal claim with wording consistent with the research design.",
            "problematic_quote": "causes environmental sustainability",
            "evidence": [{
                "paragraph": 3,
                "text": "The study will prove that green procurement causes environmental sustainability.",
                "section_reference": "Statement of the Problem",
                "document_role": "current",
            }],
            "annotation_eligible": True,
            "manual_confirmation_required": True,
        }],
    }
    output = build_annotated_docx(source, review)
    assert native_comment_count(output) == 1
    comments = list(Document(io.BytesIO(output)).comments)
    assert comments[0].author == "Anokye Mohammed Adam"
    assert "Manual confirmation recommended" not in comments[0].text


def test_smaller_verification_batches_are_default(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    monkeypatch.delenv("AI_VERIFICATION_BATCH_SIZE", raising=False)
    assert HybridAIConfig.from_env().verification_batch_size == 12


def test_limited_completed_review_has_rebuild_controls():
    portal = Path("app/templates/portal.html").read_text(encoding="utf-8")
    detail = Path("app/templates/review_detail.html").read_text(encoding="utf-8")
    main = Path("app/main.py").read_text(encoding="utf-8")
    assert "Rebuild review" in portal
    assert "Rebuild review and comments" in detail
    assert "Rebuild requested for the limited review output" in main
    assert "review-pipeline-v2.1.1-provider-recovery-hotfix" in main


def test_empty_annotated_output_is_rejected():
    main = Path("app/main.py").read_text(encoding="utf-8")
    engine = Path("app/academic_ai_engine.py").read_text(encoding="utf-8")
    assert "actual_comments = native_comment_count(annotated_data)" in main
    assert "if actual_comments < 1" in main
    assert "ReviewOutputValidationError" in engine
    assert "fresh automatic expert" in engine
