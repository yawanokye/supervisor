from __future__ import annotations

import io

from docx import Document

from app.academic_ai_engine import _section_groups
from app.annotated_exporter import build_annotated_docx
from app.document_parser import extract_docx
from app.report_exporter import _summary_units
from app.supervisory_accuracy_guard import (
    build_factual_index,
    guard_issue,
    guard_strength,
    paragraph_id,
)


def _row(number, text, heading, chapter, *, is_heading=False, table_number="", table_title=""):
    return {
        "paragraph": number,
        "text": text,
        "heading": heading,
        "section_path": [heading],
        "is_heading": is_heading,
        "chapter_number": chapter,
        "document_role": "current",
        "source_kind": "table_row" if table_number else "paragraph",
        "table_index": int(table_number) if table_number and str(table_number).isdigit() else None,
        "table_row": 1 if table_number else None,
        "table_number": table_number,
        "table_title": table_title,
    }


def _issue(section, ids, action, *, title="Issue", category="chapter_structure"):
    return {
        "finding_id": "F1",
        "category": category,
        "section": section,
        "issue_title": title,
        "severity": "major",
        "confidence": 0.95,
        "evidence_paragraph_ids": ids,
        "problematic_quote": "",
        "assessment": action,
        "academic_consequence": "The review would be inaccurate.",
        "required_action": action,
        "illustrative_guidance": "",
        "guidance_type": "direct_correction",
        "source_verification_required": False,
        "context_guard_adjusted": False,
    }


def test_bare_chapter_markers_and_titles_are_not_review_units():
    rows = [
        _row(1, "CHAPTER THREE", "CHAPTER THREE", 3, is_heading=True),
        _row(2, "RESEARCH METHODS", "RESEARCH METHODS", 3, is_heading=True),
        _row(3, "Introduction", "Introduction", 3, is_heading=True),
        _row(4, "This chapter explains the research approach, design, sampling and ethics.", "Introduction", 3),
        _row(5, "Research Approach", "Research Approach", 3, is_heading=True),
        _row(6, "The study adopted a quantitative approach.", "Research Approach", 3),
    ]
    groups = _section_groups(rows)
    headings = [group["heading"] for group in groups]
    assert "CHAPTER THREE" not in headings
    assert "RESEARCH METHODS" not in headings
    assert headings == ["Introduction", "Research Approach"]
    assert groups[0]["chapter_number"] == 3


def test_populate_chapter_three_comment_is_rejected_when_methods_exist():
    rows = [
        _row(1, "CHAPTER THREE", "CHAPTER THREE", 3, is_heading=True),
        _row(2, "Introduction", "Introduction", 3, is_heading=True),
        _row(3, "This chapter outlines the methodological procedures.", "Introduction", 3),
        _row(4, "Research Approach", "Research Approach", 3, is_heading=True),
        _row(5, "A quantitative approach was used.", "Research Approach", 3),
        _row(6, "Research Design", "Research Design", 3, is_heading=True),
        _row(7, "A cross-sectional design was used.", "Research Design", 3),
        _row(8, "Ethical Considerations", "Ethical Considerations", 3, is_heading=True),
        _row(9, "Consent and confidentiality were addressed.", "Ethical Considerations", 3),
    ]
    index = {paragraph_id(row): row for row in rows}
    issue = _issue(
        "CHAPTER THREE",
        ["P1"],
        "Populate Chapter Three with the full methodological details from research approach through ethical considerations.",
        title="Populate Chapter Three",
    )
    assert guard_issue(issue, index, build_factual_index(rows)) is None


def test_no_extra_intro_is_requested_under_chapter_title_when_introduction_exists():
    rows = [
        _row(1, "RESEARCH METHODS", "RESEARCH METHODS", 3, is_heading=True),
        _row(2, "Introduction", "Introduction", 3, is_heading=True),
        _row(3, "This chapter outlines the approach, design, population, sampling, analysis and ethics.", "Introduction", 3),
        _row(4, "Research Approach", "Research Approach", 3, is_heading=True),
        _row(5, "The study used a quantitative approach.", "Research Approach", 3),
    ]
    index = {paragraph_id(row): row for row in rows}
    issue = _issue(
        "RESEARCH METHODS",
        ["P1"],
        "Add a short introductory paragraph under the chapter heading to introduce the chapter purpose and structure.",
    )
    assert guard_issue(issue, index, build_factual_index(rows)) is None


def test_strength_about_anova_is_rejected_without_anova_evidence():
    rows = [
        _row(1, "Introduction", "Introduction", 3, is_heading=True),
        _row(2, "This chapter outlines the research approach and sampling procedure.", "Introduction", 3),
    ]
    index = {paragraph_id(row): row for row in rows}
    strength = {
        "category": "results_and_interpretation",
        "section": "Introduction",
        "evidence_paragraph_ids": ["P2"],
        "observation": "The ANOVA results are presented with a clear table and appropriate interpretation.",
    }
    assert guard_strength(strength, index, build_factual_index(rows), "Introduction") is None


def test_cross_section_comment_cannot_be_placed_under_section_without_evidence():
    rows = [
        _row(1, "Introduction", "Introduction", 3, is_heading=True),
        _row(2, "This chapter outlines the sampling procedure.", "Introduction", 3),
        _row(3, "Table 9: ANOVA", "ANOVA Results for E-Sourcing", 4, is_heading=True, table_number="9", table_title="ANOVA"),
    ]
    index = {paragraph_id(row): row for row in rows}
    issue = _issue(
        "ANOVA Results for E-Sourcing",
        ["P2"],
        "Estimate one combined multiple-regression model.",
        category="cross_section_coherence",
    )
    assert guard_issue(issue, index, build_factual_index(rows)) is None


def test_multi_table_comment_is_filtered_to_the_named_table():
    rows = [
        _row(1, "Table 20: Model Summary", "E-Invoicing Results", 4, table_number="20", table_title="Model Summary"),
        _row(2, "Table 21: ANOVA", "E-Invoicing Results", 4, table_number="21", table_title="ANOVA"),
    ]
    index = {paragraph_id(row): row for row in rows}
    issue = _issue(
        "E-Invoicing Results",
        ["P1", "P2"],
        "Revise Table 21 and explain the ANOVA result.",
        title="Table 21 requires clearer interpretation",
        category="results_and_interpretation",
    )
    guarded = guard_issue(issue, index, build_factual_index(rows))
    assert guarded is not None
    assert guarded["evidence_paragraph_ids"] == ["P2"]
    assert guarded["canonical_table_number"] == "21"


def test_heading_evidence_anchors_to_exact_heading_not_keyword_occurrence():
    doc = Document()
    doc.add_paragraph("PROCUREMENT EFFICIENCY")
    doc.add_heading("CHAPTER FOUR", level=1)
    target = doc.add_heading("Objective Four: Effects of E-Invoicing on Procurement Efficiency", level=2)
    doc.add_paragraph("The analysis is presented below.")
    stream = io.BytesIO()
    doc.save(stream)
    source = stream.getvalue()
    rows = extract_docx(source)
    heading_evidence = next(row for row in rows if row.get("text", "").startswith("Objective Four"))
    review = {
        "academic_findings": [{
            "status": "partly_meets_requirement",
            "section": heading_evidence["text"],
            "section_reference": heading_evidence["text"],
            "chapter_number": 4,
            "required_action": "Add one sentence explaining the analyses used for this objective.",
            "problematic_quote": "",
            "evidence": [{**heading_evidence, "document_role": "current"}],
            "annotation_eligible": True,
        }]
    }
    out = Document(io.BytesIO(build_annotated_docx(source, review)))
    comments = list(out.comments)
    assert len(comments) == 1
    body_xml = out.element.body.xml
    target_index = body_xml.index("Objective Four")
    keyword_index = body_xml.index("PROCUREMENT EFFICIENCY")
    comment_index = body_xml.index("commentRangeStart")
    assert keyword_index < target_index
    assert comment_index > keyword_index
    assert target.text in [p.text for p in out.paragraphs]


def test_summary_keeps_same_named_introductions_in_their_correct_chapters():
    review = {
        "summary": {"current_chapters_detected": [2, 3], "review_scope": "full_thesis"},
        "academic_section_reviews": [
            {"heading": "Introduction", "chapter_number": 2, "section_assessment": "The literature chapter is clearly introduced.", "section_score": 80},
            {"heading": "Introduction", "chapter_number": 3, "section_assessment": "The methods chapter introduction accurately outlines its contents.", "section_score": 82},
        ],
        "academic_strengths": [
            {"section": "Introduction", "observation": "The literature chapter is clearly introduced.", "evidence": [{"chapter_number": 2}]},
            {"section": "Introduction", "observation": "The methods chapter introduction accurately outlines its contents.", "evidence": [{"chapter_number": 3}]},
        ],
        "academic_findings": [],
    }
    units = _summary_units(review)
    by_chapter = {unit["chapter_number"]: unit for unit in units}
    assert "literature chapter" in " ".join(by_chapter[2]["strengths"]).lower()
    assert "methods chapter" not in " ".join(by_chapter[2]["strengths"]).lower()
    assert "methods chapter" in " ".join(by_chapter[3]["strengths"]).lower()


def test_section_assessment_drops_unsupported_anova_claim():
    from app.supervisory_accuracy_guard import guard_section_assessment

    rows = [
        {"text": "This chapter presents the research approach, design, population and ethical considerations.", "heading": "Introduction"},
    ]
    result = guard_section_assessment(
        "The introduction outlines the chapter. The ANOVA results are presented with a clear table and appropriate interpretation.",
        rows,
    )
    assert "outlines the chapter" in result
    assert "ANOVA" not in result


def test_report_strengths_do_not_infer_strength_from_unverified_assessment():
    from app.report_exporter import _section_strengths

    assert _section_strengths(
        "Introduction",
        [],
        ["The ANOVA results are presented with a clear table and appropriate interpretation."],
        3,
    ) == []


def test_reference_entries_and_keywords_are_not_split_into_false_sections(tmp_path):
    from docx import Document
    from app.document_parser import extract_docx
    from app.academic_ai_engine import _section_groups

    doc = Document()
    doc.add_heading("KEYWORDS", level=1)
    doc.add_paragraph("PROCUREMENT EFFICIENCY", style="Heading 2")
    doc.add_paragraph("E-SOURCING", style="Heading 2")
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("INTRODUCTION", level=1)
    doc.add_heading("Background to the Study", level=2)
    doc.add_paragraph("Substantive background text.")
    doc.add_heading("CHAPTER FIVE", level=1)
    doc.add_heading("SUMMARY, CONCLUSION AND RECOMMENDATIONS", level=1)
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("A substantive conclusion.")
    doc.add_heading("References", level=1)
    doc.add_paragraph("Author, A. (2024). A reference title.", style="Heading 2")
    doc.add_paragraph("Another, B. (2023). Another source.", style="Heading 2")
    doc.add_heading("APPENDIX", level=1)
    doc.add_heading("QUESTIONNAIRE", level=1)
    doc.add_heading("INTRODUCTION", level=2)
    doc.add_paragraph("This questionnaire is for academic purposes.")
    path = tmp_path / "sample.docx"
    doc.save(path)

    rows = extract_docx(path.read_bytes())
    groups = _section_groups(rows)
    headings = [group["heading"] for group in groups]
    assert headings.count("KEYWORDS") == 1
    assert "PROCUREMENT EFFICIENCY" not in headings
    assert "E-SOURCING" not in headings
    assert headings.count("References") == 1
    assert "Author, A. (2024). A reference title." not in headings
    assert "QUESTIONNAIRE" not in headings
    appendix_rows = [row for row in rows if row.get("heading") in {"APPENDIX", "QUESTIONNAIRE", "INTRODUCTION"} and row.get("paragraph", 0) > 10]
    assert all(row.get("chapter_number") is None for row in appendix_rows)


def test_heading_only_parent_of_subsection_is_not_reviewed_as_missing():
    from app.academic_ai_engine import _section_groups

    rows = [
        {"text": "CHAPTER FOUR", "is_heading": True, "chapter_number": 4, "section_path": ["CHAPTER FOUR"]},
        {"text": "Objective One: E-Sourcing", "is_heading": True, "chapter_number": 4, "section_path": ["Objective One: E-Sourcing"]},
        {"text": "Correlation Analysis", "is_heading": True, "chapter_number": 4, "section_path": ["Objective One: E-Sourcing", "Correlation Analysis"]},
        {"text": "The relationship was assessed.", "is_heading": False, "heading": "Correlation Analysis", "chapter_number": 4, "section_path": ["Objective One: E-Sourcing", "Correlation Analysis"]},
    ]
    headings = [group["heading"] for group in _section_groups(rows)]
    assert "Objective One: E-Sourcing" not in headings
    assert "Correlation Analysis" in headings


def test_cross_section_table_issue_keeps_method_evidence_and_exact_table():
    from app.supervisory_accuracy_guard import build_factual_index, guard_issue, paragraph_id

    rows = [
        {"text": "Multiple regression entered all predictors simultaneously.", "paragraph": 1, "heading": "Data Processing and Analysis", "section_path": ["Data Processing and Analysis"], "chapter_number": 3},
        {"text": "Table 20: Model Summary", "paragraph": 2, "heading": "E-Invoicing Results", "section_path": ["E-Invoicing Results"], "chapter_number": 4, "source_kind": "table_caption", "table_index": 20, "table_number": "20", "table_title": "Model Summary"},
        {"text": "Table 21: ANOVA", "paragraph": 3, "heading": "E-Invoicing Results", "section_path": ["E-Invoicing Results"], "chapter_number": 4, "source_kind": "table_caption", "table_index": 21, "table_number": "21", "table_title": "ANOVA"},
    ]
    index = {paragraph_id(row): row for row in rows}
    issue = {
        "category": "cross_section_coherence",
        "section": "E-Invoicing Results",
        "issue_title": "Table 21 does not implement the stated joint model",
        "assessment": "Chapter Three promises a joint model, but Table 21 belongs to a separate single-predictor model.",
        "academic_consequence": "The analysis does not match the method.",
        "required_action": "Estimate the stated combined model.",
        "evidence_paragraph_ids": ["P1", "P2", "P3"],
        "severity": "major",
        "confidence": 0.95,
    }
    result = guard_issue(issue, index, build_factual_index(rows))
    assert result is not None
    assert "P1" in result["evidence_paragraph_ids"]
    assert "P3" in result["evidence_paragraph_ids"]
    assert "P2" not in result["evidence_paragraph_ids"]
    assert result["canonical_table_number"] == "21"
