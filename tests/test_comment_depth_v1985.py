from __future__ import annotations

import io

from docx import Document

from app.academic_ai_engine import _degree_comment_floor
from app.ai_config import HybridAIConfig
from app.annotated_exporter import build_annotated_docx


def _docx_bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_standard_depth_floor_orders_non_research_below_mphil(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test")
    monkeypatch.setenv("DEEPSEEK_API_KEY", "test")
    config = HybridAIConfig.from_env()

    assert _degree_comment_floor("Non-Research Masters", "standard", config) == 18
    assert _degree_comment_floor("Research Masters / MPhil", "standard", config) == 24
    assert _degree_comment_floor("Professional Doctorate", "standard", config) == 28
    assert _degree_comment_floor("PhD", "standard", config) == 32
    assert _degree_comment_floor("Non-Research Masters", "standard", config) < _degree_comment_floor("Research Masters / MPhil", "standard", config)


def test_native_comments_export_developmental_detail(monkeypatch):
    monkeypatch.setenv("VPROF_COMMENT_MAX_CHARS", "1100")
    document = Document()
    document.add_heading("Purpose of the Study", level=2)
    document.add_paragraph(
        "The purpose of this study is to examine green procurement and environmental sustainability."
    )
    review = {
        "summary": {"reviewer_name": "Supervisor"},
        "academic_findings": [{
            "status": "partly_meets_requirement",
            "category": "cross_section_coherence",
            "section": "Purpose of the Study",
            "section_reference": "Purpose of the Study",
            "reference_label": "Purpose of the Study",
            "item": "Purpose and objectives are not fully aligned",
            "comment": "Objective 4 introduces operational performance, but the purpose is limited to environmental sustainability.",
            "required_action": "Revise the purpose so it covers every principal outcome, or remove objectives that fall outside the stated scope.",
            "academic_consequence": "The methodology and conclusions may not be traceable to one coherent purpose.",
            "problematic_quote": "The purpose of this study",
            "evidence": [{
                "paragraph": 2,
                "text": "The purpose of this study is to examine green procurement and environmental sustainability.",
                "section_reference": "Purpose of the Study",
                "document_role": "current",
            }],
            "annotation_eligible": True,
        }],
    }
    output = build_annotated_docx(_docx_bytes(document), review)
    comments = list(Document(io.BytesIO(output)).comments)
    assert len(comments) == 1
    text = comments[0].text
    assert "Issue:" not in text
    assert "Why this matters:" not in text
    assert "Revise by:" not in text
    assert "Objective 4 introduces operational performance" in text
