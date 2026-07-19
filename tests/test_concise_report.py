from io import BytesIO

from docx import Document

from app.report_exporter import build_docx_report


def sample_review():
    return {
        "summary": {
            "filename": "Chapter One.docx",
            "document_label": "Chapter 1",
            "academic_level": "Research Masters / MPhil",
            "research_approach": "quantitative",
            "review_depth": "standard",
            "review_benchmark": "Research Master’s / MPhil",
            "academic_sections_reviewed": 3,
            "readiness_label": "Major revision required",
            "readiness_meaning": "Revise the major academic issues before resubmission.",
            "critical_issues": 0,
            "major_issues": 3,
            "moderate_issues": 2,
            "minor_issues": 1,
            "selected_chapter": 1,
            "current_chapters_detected": [1],
        },
        "study_context": {
            "confirmed_countries": ["Ghana"],
            "confirmed_locations": ["Western Region"],
            "confirmed_sectors": ["mining"],
            "title_or_opening_focus": "Stakeholder engagement and project outcomes",
        },
        "overall_academic_assessment": (
            "The chapter establishes a relevant research topic and includes the "
            "main introductory sections. The problem, variables and study context "
            "need sharper alignment before the chapter is ready for approval. "
            "Language and source verification also require attention."
        ),
        "academic_strengths": [
            {
                "section": "Background to the Study",
                "observation": "The section introduces stakeholder engagement as a relevant project-management concern.",
            },
            {
                "section": "Statement of the Problem",
                "observation": "The section identifies delays, cost overruns and stakeholder conflict as practical concerns.",
            },
        ],
        "academic_section_reviews": [
            {
                "heading": "Background to the Study",
                "section_assessment": "The background introduces the topic but needs stronger empirical support.",
                "section_score": 62,
                "coverage_warning": "",
            },
            {
                "heading": "Statement of the Problem",
                "section_assessment": "The problem is relevant but the empirical and contextual gap is not demonstrated sufficiently.",
                "section_score": 55,
                "coverage_warning": "",
            },
            {
                "heading": "Research Objectives",
                "section_assessment": "The objectives cover relevant engagement practices but require clearer construct separation.",
                "section_score": 58,
                "coverage_warning": "",
            },
        ],
        "academic_findings": [
            {
                "finding_id": "F1",
                "section": "Background to the Study",
                "item": "Insufficient empirical context",
                "severity": "major",
                "status": "partly_meets_requirement",
                "required_action": "Add verified empirical evidence that establishes the study context and leads logically to the problem.",
                "comment": "The discussion remains largely descriptive.",
                "evidence": [{"chapter_number": 1, "paragraph": 3}],
            },
            {
                "finding_id": "F2",
                "section": "Statement of the Problem",
                "item": "Problem not demonstrated",
                "severity": "major",
                "status": "partly_meets_requirement",
                "required_action": "Demonstrate the practical and knowledge problem using verified evidence from the actual Ghanaian mining-sector context.",
                "comment": "The problem is asserted rather than demonstrated.",
                "evidence": [{"chapter_number": 1, "paragraph": 8}],
            },
            {
                "finding_id": "F3",
                "section": "Statement of the Problem",
                "item": "Research gap unclear",
                "severity": "major",
                "status": "partly_meets_requirement",
                "required_action": "Explain what previous studies have not established and end with the precise focus of this study.",
                "comment": "The gap needs clearer derivation.",
                "evidence": [{"chapter_number": 1, "paragraph": 9}],
            },
            {
                "finding_id": "F4",
                "section": "Research Objectives",
                "item": "Composite objective",
                "severity": "moderate",
                "status": "partly_meets_requirement",
                "required_action": "Separate decision-making involvement, collaboration and monitoring, or justify them as dimensions of one construct.",
                "comment": "The objective combines several predictors.",
                "evidence": [{"chapter_number": 1, "paragraph": 14}],
            },
        ],
        "priority_actions": [
            {
                "section": "Statement of the Problem",
                "severity": "major",
                "action": "Demonstrate the problem and research gap with verified context-specific evidence.",
            },
            {
                "section": "Research Objectives",
                "severity": "moderate",
                "action": "Align the constructs, objectives and research questions.",
            },
        ],
        "alignment_results": [],
        "revision_results": [],
    }


def test_report_is_concise_and_uses_professional_action_sections():
    data = build_docx_report(sample_review())
    document = Document(BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "SUPERVISORY REVIEW REPORT" in text
    assert "Overall supervisory assessment" in text
    assert "Main strengths" in text
    assert "Actions Required Before Supervisor Approval or Submission" in text
    assert "Professional recommendation" in text
    assert "Review point 1" not in text
    assert "Illustrative guidance" not in text


def test_report_action_table_lists_each_material_correction_with_verification():
    data = build_docx_report(sample_review())
    document = Document(BytesIO(data))
    action_tables = [table for table in document.tables if len(table.columns) == 5]
    assert action_tables
    headers = [cell.text for cell in action_tables[0].rows[0].cells]
    assert headers == [
        "Priority",
        "Location",
        "Specific action required",
        "Why this must be corrected",
        "How to verify completion",
    ]
    problem_rows = [
        row for row in action_tables[0].rows[1:]
        if "Statement of the Problem" in row.cells[1].text
    ]
    assert len(problem_rows) >= 2
    assert all(row.cells[2].text.strip() for row in problem_rows)
    assert all(row.cells[4].text.strip() for row in problem_rows)
