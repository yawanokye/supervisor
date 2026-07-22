from __future__ import annotations

import io
import re

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.document_parser import extract_docx
from app.final_review_quality import build_canonical_finding_rows


def _row(fid: str, paragraph: int, *, section: str, item: str, comment: str, action: str, category: str = "other", quote: str = "", severity: str = "major"):
    return {
        "finding_id": fid,
        "status": "does_not_meet_requirement",
        "severity": severity,
        "confidence": 0.95,
        "chapter_number": 1,
        "section": section,
        "section_reference": section,
        "category": category,
        "item": item,
        "comment": comment,
        "required_action": action,
        "problematic_quote": quote,
        "evidence": [{
            "document_role": "current",
            "chapter_number": 1,
            "paragraph": paragraph,
            "heading": section,
            "section_reference": section,
            "section_path": ["Chapter One", section],
            "text": quote or item,
        }],
        "annotation_eligible": True,
    }


def test_significance_root_causes_are_consolidated_and_not_repeated():
    review = {
        "summary": {"academic_level": "MPhil"},
        "academic_findings": [
            _row("S1", 20, section="Significance of the Study", item="The scholarly contribution is not explicit", comment="The section lists beneficiaries.", action="State the scholarly contribution."),
            _row("S2", 20, section="Significance of the Study", item="Theory practice and policy are not separated", comment="The contribution is not organised.", action="Separate theory, practice and policy."),
            _row("S3", 20, section="Significance of the Study", item="The research contribution is unclear", comment="The study does not say what it adds.", action="Explain what the study adds."),
            _row("S4", 20, section="Significance of the Study", item="The research gap is placed inside significance", comment="The gap belongs in the problem logic.", action="Move the research gap to the problem statement."),
        ],
    }
    rows = build_canonical_finding_rows(review, force=True)
    assert len(rows) == 2
    titles = " | ".join(row["item"] for row in rows).lower()
    assert "scholarly, practical and policy contribution" in titles
    assert "research gap" in titles
    assert all(row.get("student_comment") for row in rows)


def test_chapter_one_comment_rejects_statistical_example():
    finding = _row(
        "P1", 10, section="Statement of the Problem",
        item="The problem statement lacks local evidence",
        comment="The problem is described generally.",
        action="Add evidence from the study context.",
    )
    finding["illustrative_guidance"] = "Report the coefficient, standard error, p-value and confidence interval."
    review = {"summary": {"academic_level": "MPhil"}, "academic_findings": [finding]}
    row = build_canonical_finding_rows(review, force=True)[0]
    text = row["student_comment"].lower()
    assert "standard error" not in text
    assert "confidence interval" not in text
    assert "study context" in text or "institutional" in text or "regulatory" in text


def test_scope_inconsistency_is_added_conservatively_from_current_work():
    current = [
        {"paragraph": 1, "chapter_number": 1, "heading": "Background of the Study", "text": "This study examines internal controls among commercial banks in Ghana."},
        {"paragraph": 2, "chapter_number": 1, "heading": "Background of the Study", "text": "Rural banks play an important role in local finance."},
        {"paragraph": 3, "chapter_number": 1, "heading": "Research Objectives", "text": "The study specifically uses Assinman Rural Bank PLC as the case study."},
    ]
    review = {
        "summary": {"academic_level": "MPhil"},
        "academic_findings": [],
        "_runtime_context": {"current_paragraphs": current},
    }
    rows = build_canonical_finding_rows(review, force=True)
    assert len(rows) == 1
    assert "population and institutional scope" in rows[0]["item"].lower()
    assert "Assinman Rural Bank" in rows[0]["student_comment"]


def test_verified_missing_section_uses_insertion_anchor_for_numbering():
    missing = _row(
        "M1", 5, section="Chapter One", item="Purpose of the Study is missing from Chapter One",
        comment="Purpose of the Study is missing from Chapter One.", action="Add the section after the problem statement.",
    )
    missing.update({
        "missing_section_label": "Purpose of the Study",
        "section_contract_label": "Purpose of the Study",
        "section_contract_verified": True,
        "section_status": "missing",
    })
    later = _row(
        "L1", 10, section="Research Objectives", item="The objectives are inconsistent",
        comment="The objectives use different populations.", action="Use one population consistently.",
    )
    review = {"summary": {"academic_level": "MPhil"}, "academic_findings": [later, missing]}
    rows = build_canonical_finding_rows(review, force=True)
    assert rows[0]["missing_section_label"] == "Purpose of the Study"
    assert [row["finding_number"] for row in rows] == [1, 2]


def test_docx_markers_are_never_inserted_inside_words(monkeypatch):
    monkeypatch.setenv("VPROF_HUMAN_SUPERVISORY_EDITOR", "true")
    monkeypatch.setenv("VPROF_NATIVE_GROUP_LOCATION_MARKERS", "true")
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_paragraph("Fraudulent activities pose significant challenges to the banking sector.")
    source = io.BytesIO(); doc.save(source)
    parsed = extract_docx(source.getvalue())
    target = next(row for row in parsed if "significant challenges" in row.get("text", ""))
    findings = [
        _row("A1", target["paragraph"], section="Problem Statement", item="The claim needs evidence", comment="The claim is broad.", action="Add evidence.", category="evidence", quote="challen"),
        _row("A2", target["paragraph"], section="Problem Statement", item="The wording is imprecise", comment="The wording is vague.", action="Clarify the wording.", category="academic_writing", quote="challen"),
    ]
    for finding in findings:
        finding["evidence"] = [{**target, "document_role": "current"}]
    review = {"summary": {"academic_level": "MPhil", "reviewer_name": "Supervisor"}, "academic_findings": findings}
    output = build_annotated_docx(source.getvalue(), review)
    reviewed = Document(io.BytesIO(output))
    body = "\n".join(paragraph.text for paragraph in reviewed.paragraphs)
    assert not re.search(r"[A-Za-z]\s*\[\d+\]\s*[A-Za-z]", body)
    assert "challen [" not in body.lower()
    assert "[1]" in body and "[2]" in body


def test_minor_human_comment_does_not_force_an_example():
    finding = _row(
        "C1", 2, section="Introduction", item="A duplicate citation should be removed",
        comment="The same citation appears twice in one sentence.", action="Retain the citation once.", category="citation_and_reference_integrity", severity="minor",
    )
    review = {"summary": {"academic_level": "Bachelor"}, "academic_findings": [finding]}
    row = build_canonical_finding_rows(review, force=True)[0]
    assert "For example" not in row["student_comment"]


def test_missing_section_comment_keeps_context_specific_example():
    missing = _row(
        "M2", 8, section="Problem Statement",
        item="Purpose of the Study is missing from Chapter One",
        comment="Purpose of the Study is missing from Chapter One.",
        action="Add a clearly labelled Purpose of the Study section after the problem statement.",
    )
    missing.update({
        "missing_section_label": "Purpose of the Study",
        "section_contract_label": "Purpose of the Study",
        "section_contract_verified": True,
        "section_status": "missing",
    })
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Effect of Internal Controls on Fraud Detection and Prevention at Assinman Rural Bank PLC"},
        "academic_findings": [missing],
    }
    row = build_canonical_finding_rows(review, force=True)[0]
    assert "For example" in row["student_comment"]
    assert "internal controls" in row["student_comment"].lower()
    assert "fraud detection and prevention" in row["student_comment"].lower()
