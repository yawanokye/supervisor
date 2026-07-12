from __future__ import annotations

import io
import zipfile

from docx import Document

from app.annotated_exporter import ANNOTATION_EXPORT_VERSION, build_annotated_docx, native_comment_count
from app.document_parser import extract_docx


def _docx_bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _visible_content(document: Document):
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    return paragraphs, tables


def test_annotations_are_native_word_comments_and_body_is_unchanged():
    document = Document()
    document.add_heading("CHAPTER FOUR", level=1)
    document.add_heading("4.2 Regression Results", level=2)
    target = document.add_paragraph(
        "The coefficient is statistically significant, holding all other variables constant."
    )
    document.add_paragraph("Table 4.1 Regression estimates")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Variable"
    table.cell(0, 1).text = "Coefficient"
    table.cell(1, 0).text = "E-sourcing"
    table.cell(1, 1).text = "0.42"
    source = _docx_bytes(document)

    rows = extract_docx(source)
    paragraph_evidence = next(
        row for row in rows
        if "holding all other variables constant" in row.get("text", "")
    )
    table_evidence = next(
        row for row in rows
        if row.get("source_kind") == "table_row" and row.get("table_row") == 2
    )
    review = {
        "summary": {"reviewer_name": "Anokye Mohammed Adam"},
        "academic_findings": [
            {
                "status": "partly_meets_requirement",
                "section": "4.2 Regression Results",
                "section_reference": "4.2 Regression Results",
                "required_action": (
                    "Remove the claim that other variables were held constant because "
                    "the reported model contains only one predictor."
                ),
                "problematic_quote": "holding all other variables constant",
                "evidence": [{**paragraph_evidence, "document_role": "current"}],
                "annotation_eligible": True,
            },
            {
                "status": "partly_meets_requirement",
                "section": "4.2 Regression Results",
                "section_reference": "4.2 Regression Results",
                "required_action": "Interpret the coefficient and report its uncertainty.",
                "problematic_quote": "",
                "evidence": [{**table_evidence, "document_role": "current"}],
                "annotation_eligible": True,
            },
        ]
    }

    before = Document(io.BytesIO(source))
    annotated_bytes = build_annotated_docx(source, review)
    after = Document(io.BytesIO(annotated_bytes))

    assert ANNOTATION_EXPORT_VERSION == "1.9.9.28-context-guidance-editor"
    after_paragraphs, after_tables = _visible_content(after)
    before_paragraphs, before_tables = _visible_content(before)
    assert after_tables == before_tables
    assert target.text.replace(".", "") in after_paragraphs[2].replace(" [1]", "").replace(".", "")
    # Evidence-anchored grouping keeps different locations as separate comments:
    # one comment on the exact sentence and one comment on the table evidence.
    assert len(list(after.comments)) == 2
    comments = list(after.comments)
    comment_text = "\n".join(comment.text for comment in comments)
    assert all(comment.author == "Anokye Mohammed Adam" for comment in comments)
    assert all(comment.initials == "AMA" for comment in comments)
    assert "1. " in comment_text or len(comments) >= 2
    assert "Remove the claim" in comment_text
    assert "Interpret the coefficient" in comment_text
    assert all("Supervisor comment" not in paragraph.text for paragraph in after.paragraphs)
    assert "SUPERVISOR REVIEW NOTES" not in "\n".join(paragraph.text for paragraph in after.paragraphs)
    assert 'w:val="C00000"' in after.element.xml
    assert 'w:val="008000"' not in after.element.xml
    assert "w:commentRangeStart" in after.element.body.xml

    with zipfile.ZipFile(io.BytesIO(annotated_bytes)) as package:
        names = set(package.namelist())
        assert "word/comments.xml" in names
        assert "word/_rels/document.xml.rels" in names


def test_missing_section_feedback_is_added_as_blue_inline_bottom_note_not_native_comment():
    document = Document()
    document.add_paragraph("Original title")
    document.add_paragraph("Original body text remains unchanged.")
    source = _docx_bytes(document)

    review = {
        "summary": {"reviewer_name": "Dr Priscilla Boafowaa Oppong"},
        "academic_findings": [
            {
                "status": "does_not_meet_requirement",
                "section": "Definition of Terms",
                "section_reference": "Definition of Terms",
                "item": "Expected UCC thesis section is not evident: Definition of Terms",
                "comment": "The chapter does not make the Definition of Terms section evident.",
                "required_action": "Add the required section and explain its purpose.",
                "problematic_quote": "",
                "evidence": [],
                "headings": ["Definition of Terms"],
                "annotation_eligible": True,
            }
        ]
    }

    after = Document(io.BytesIO(build_annotated_docx(source, review)))

    paragraphs = [p.text for p in after.paragraphs]
    assert paragraphs[:2] == ["Original title", "Original body text remains unchanged."]
    assert "Specific corrections required" in paragraphs
    assert any("Definition of Terms" in text for text in paragraphs)
    assert any("Add a clearly labelled Definition of Terms section" in text for text in paragraphs)
    assert any(text.startswith("1. ") for text in paragraphs)
    comments = list(after.comments)
    assert len(comments) == 0
    assert "SUPERVISOR REVIEW NOTES" not in "\n".join(p.text for p in after.paragraphs)


def test_explicit_comment_author_overrides_review_metadata():
    document = Document()
    document.add_paragraph("Text requiring review.")
    source = _docx_bytes(document)
    rows = extract_docx(source)
    evidence = next(row for row in rows if "requiring review" in row.get("text", ""))
    review = {
        "summary": {"reviewer_name": "Stored Reviewer"},
        "academic_findings": [{
            "status": "partly_meets_requirement",
            "section": "Body",
            "section_reference": "Body",
            "required_action": "Clarify this statement.",
            "problematic_quote": "requiring review",
            "evidence": [{**evidence, "document_role": "current"}],
            "annotation_eligible": True,
        }],
    }
    output = Document(io.BytesIO(build_annotated_docx(
        source, review, "Anokye Mohammed Adam"
    )))
    comments = list(output.comments)
    assert len(comments) == 1
    assert comments[0].author == "Anokye Mohammed Adam"
    assert comments[0].initials == "AMA"


def test_grouped_native_comment_numbers_related_findings_and_keeps_one_context_example():
    document = Document()
    document.add_heading("1.3 Problem Statement", level=2)
    document.add_paragraph("Commercial banks and rural banks are both discussed without a clear boundary.")
    source = _docx_bytes(document)
    rows = extract_docx(source)
    evidence = next(row for row in rows if "Commercial banks" in row.get("text", ""))
    review = {
        "summary": {"reviewer_name": "Dr Priscilla Boafowaa Oppong"},
        "academic_findings": [
            {
                "status": "does_not_meet_requirement",
                "section": "Problem Statement",
                "section_reference": "1.3 Problem Statement",
                "item": "The study population and case setting are not consistently stated",
                "comment": "The chapter alternates between commercial banks, rural banks and Assinman Rural Bank PLC.",
                "required_action": "State the target population and case setting consistently across the title, problem, objectives, questions, scope and significance.",
                "illustrative_guidance": "state whether the study is a case study of Assinman Rural Bank PLC within Ghana's rural banking sector",
                "evidence": [{**evidence, "document_role": "current"}],
                "headings": ["1.3 Problem Statement"],
                "annotation_eligible": True,
            },
            {
                "status": "partly_meets_requirement",
                "section": "Problem Statement",
                "section_reference": "1.3 Problem Statement",
                "item": "The problem statement lacks local evidence",
                "comment": "The problem is argued generally without concrete Ghana rural banking evidence.",
                "required_action": "Support the problem with verifiable sector, policy or institutional evidence.",
                "illustrative_guidance": "use Bank of Ghana, ARB Apex Bank or permitted institutional records as evidence types without inventing figures",
                "evidence": [{**evidence, "document_role": "current"}],
                "headings": ["1.3 Problem Statement"],
                "annotation_eligible": True,
            },
        ],
    }
    output = Document(io.BytesIO(build_annotated_docx(source, review)))
    comments = list(output.comments)
    assert len(comments) == 2
    assert comments[0].text.startswith("1. ")
    assert comments[1].text.startswith("2. ")
    text = " ".join(comment.text for comment in comments)
    assert "Applies to" not in text
    assert "For example," in text
    assert output.paragraphs[1].text.endswith("[1] [2]")
    stream = io.BytesIO()
    output.save(stream)
    with zipfile.ZipFile(io.BytesIO(stream.getvalue())) as package:
        assert 'w:val="C00000"' in package.read("word/document.xml").decode("utf-8")
