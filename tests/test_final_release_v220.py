from __future__ import annotations

import io
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree

from app.academic_ai_engine import _requires_paid_comment_audit
from app.ai_config import HybridAIConfig
from app.annotated_exporter import build_annotated_docx, native_comment_count
from app.deterministic_supervisory_checklist import (
    deterministic_supervisory_checklist_issues,
    hard_chapter_one_supervisory_issues,
)
from app.document_parser import parse_document
from app.final_review_quality import build_canonical_finding_rows
from app.inline_annotated_exporter import build_inline_annotated_docx
from app.model_router import CostAwareAIProvider, ReviewStage


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _fixture_docx() -> bytes:
    document = Document()
    document.add_paragraph("IMPACT OF EFFECTIVE PLANNING AND CONTROL IN PROCUREMENT ACTIVITIES OF A MANUFACTURING FIRM")
    document.add_heading("CHAPTER 1", level=1)
    document.add_heading("INTRODUCTION", level=1)
    document.add_paragraph(
        "This part of the work consists of background to the study, statement of the problem, purpose of the study and objectives. "
        "Also included in this part are research questions, significance of the study, limitation and organization of the study."
    )
    document.add_heading("1.1 Background to the Study", level=2)
    document.add_paragraph("Delete all numbering from the work. This is not UCC style.")
    document.add_paragraph(
        "Prior to 1900, procurement was recognized as an independent function by many railroad organizations. "
        "In1970s & 1980s, procurement strategy became important. In the 1990s, software helped automate the process (Smith,"
    )
    document.add_paragraph(
        "According to White et al (2016), procurement planning supports performance."
    )
    document.add_paragraph(
        "The concept of control procurement emerged from industry. Control Procurement is used by numerous companies."
    )
    document.add_heading("1.2 Statement of the Problem", level=2)
    document.add_paragraph(
        "Organisations should have its procurement function as a vital section. Evidence from the Kenya Financial Sector identifies poor planning (Nyakundi, et al., 2012)."
    )
    document.add_paragraph(
        "Despite the significant impact of planning and control, researches on these constructs have not received much attention."
    )
    document.add_heading("1.3 Purpose of the Study", level=2)
    document.add_paragraph(
        "The purpose is to assess the impact of planning and control with regards to a manufacturing firm and suggest measures manufacturing firms can adopt to maximise the benefits that comes with planning and control."
    )
    document.add_heading("1.4 Objectives of the Study", level=2)
    document.add_paragraph("1. To assess current practices.\n2. To identify effects.\n3. To examine challenges.")
    document.add_heading("1.5 Research Questions", level=2)
    document.add_paragraph("What is the current practices of planning and control in a manufacturing firm?")
    document.add_heading("1.6 Significance of the Study", level=2)
    document.add_paragraph("The findings will contribute to research, policy and practice and serve as a reference to students.")
    document.add_heading("1.7 Limitation of the Study", level=2)
    document.add_paragraph("The study will be limited to manufacturing firms and may not be generalised to other industries.")
    document.add_heading("1.8 Organization of the Study", level=2)
    document.add_paragraph(
        "Chapter one introduces the study. Chapter two reviews the literature, chapter three presents methods, chapter four presents results and chapter five concludes the work."
    )
    buffer = io.BytesIO()
    document.save(buffer)

    # Wrap the embedded supervisor instruction in a tracked insertion so the
    # parser must read revision-visible text rather than paragraph.text alone.
    source = ZipFile(io.BytesIO(buffer.getvalue()))
    output = io.BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                for paragraph in root.xpath(".//w:p", namespaces={"w": W}):
                    visible = "".join(paragraph.xpath(".//w:t/text()", namespaces={"w": W}))
                    if visible == "Delete all numbering from the work. This is not UCC style.":
                        insertion = etree.Element(f"{{{W}}}ins")
                        insertion.set(f"{{{W}}}author", "Supervisor")
                        insertion.set(f"{{{W}}}id", "101")
                        for run in list(paragraph.xpath("./w:r", namespaces={"w": W})):
                            paragraph.remove(run)
                            insertion.append(run)
                        paragraph.append(insertion)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            target.writestr(info, data)
    return output.getvalue()


def _review(source: bytes):
    rows = parse_document(source, "chapter-one.docx")
    issues = deterministic_supervisory_checklist_issues(
        rows,
        academic_level="Bachelors",
        submission_scope="chapter",
        max_issues=100,
    )
    review = {
        "summary": {
            "academic_level": "Bachelors",
            "review_depth": "standard",
            "submission_scope": "chapter",
            "review_scope": "chapter",
            "overall_decision": "MAJOR REVISION REQUIRED",
            "study_title": "Impact of Effective Planning and Control in Procurement Activities of a Manufacturing Firm",
        },
        "academic_findings": issues,
        "_runtime_context": {"current_paragraphs": rows},
    }
    canonical = build_canonical_finding_rows(review, force=True)
    return rows, review, canonical


def test_parser_keeps_revision_text_and_does_not_invent_chapter_two():
    rows = parse_document(_fixture_docx(), "chapter-one.docx")
    instruction = next(row for row in rows if "Delete all numbering" in row["text"])
    assert instruction["contains_tracked_changes"] is True
    assert {row.get("chapter_number") for row in rows if row.get("chapter_number")} == {1}
    organisation = next(row for row in rows if row["text"].startswith("Chapter one introduces"))
    assert organisation["is_heading"] is False
    assert organisation["chapter_number"] == 1


def test_known_quality_defects_are_detected_and_false_generic_rules_are_absent():
    rows, _, _ = _review(_fixture_docx())
    findings = hard_chapter_one_supervisory_issues(
        rows, academic_level="Bachelors", submission_scope="chapter"
    )
    ids = {item["finding_id"] for item in findings}
    required = {
        "DSC-HARD-DOC-UNRESOLVED-SUPERVISOR-INSTRUCTION",
        "DSC-HARD-CIT-INCOMPLETE-PARENTHETICAL",
        "DSC-HARD-B2-CONTEXT-EVIDENCE-MISMATCH",
        "DSC-HARD-B2-UNSUPPORTED-LITERATURE-GAP",
        "DSC-HARD-B3-PURPOSE-OBJECTIVE-CONTENT-MISMATCH",
        "DSC-HARD-B3-UNIT-OF-ANALYSIS-SINGULAR-PLURAL",
        "DSC-HARD-RQ-SUBJECT-VERB-AGREEMENT",
        "DSC-HARD-B4-LIMITATION-DELIMITATION-CONFUSION",
    }
    assert required <= ids
    assert not any(item["finding_id"].startswith("DSC-C1.") for item in findings)


def test_canonical_ledger_deduplicates_and_numbers_only_after_final_filtering():
    _, _, canonical = _review(_fixture_docx())
    assert [row["finding_number"] for row in canonical] == list(range(1, len(canonical) + 1))
    titles = [row["issue_title"] for row in canonical]
    assert sum("purpose" in title.lower() and "object" in title.lower() for title in titles) == 1
    assert sum("significance" in title.lower() or "contribution" in title.lower() for title in titles) == 1
    assert len(canonical) >= 15


def test_same_anchor_findings_share_one_sorted_native_comment_box(monkeypatch):
    monkeypatch.setenv("VPROF_NATIVE_COMMENT_STYLE", "exact_anchor_grouped")
    monkeypatch.setenv("VPROF_GROUP_SAME_ANCHOR_COMMENTS", "true")
    monkeypatch.setenv("VPROF_EXPORT_ONE_COMMENT_PER_FINDING", "false")
    source = _fixture_docx()
    _, review, canonical = _review(source)
    output = build_annotated_docx(source, review, comment_author="V-Professor")
    document = Document(io.BytesIO(output))
    comment_texts = [" ".join(p.text for p in comment.paragraphs) for comment in document.comments]
    assert native_comment_count(output) < len(canonical)
    purpose_comment = next(text for text in comment_texts if "benefits that come" in text)
    purpose_numbers = [int(token[:-1]) for token in purpose_comment.split() if token[:-1].isdigit() and token.endswith(".")]
    assert purpose_numbers == sorted(purpose_numbers)
    assert "unit and scope" in purpose_comment.lower()
    assert any("incomplete parenthetical citation" in text.lower() for text in comment_texts)
    assert any("unresolved supervisor" in text.lower() for text in comment_texts)


def test_inline_annotations_are_concise_and_come_from_same_numbered_ledger(monkeypatch):
    monkeypatch.setenv("VPROF_INLINE_COMMENT_MAX_CHARS", "480")
    source = _fixture_docx()
    _, review, canonical = _review(source)
    output = build_inline_annotated_docx(source, review, comment_author="V-Professor")
    document = Document(io.BytesIO(output))
    notes = [p.text for p in document.paragraphs if p.text.startswith("Detailed supervisor comment:")]
    assert notes
    assert all(len(note) <= 2300 for note in notes)
    summary = "\n".join(p.text for p in document.paragraphs[-50:])
    assert f"identified {len(canonical)} actionable matters" in summary
    assert f"{len(canonical)}. [" in summary


def test_standard_review_avoids_extra_paid_audits_but_keeps_high_risk_checks():
    editorial = {
        "severity": "minor",
        "confidence": 0.99,
        "category": "academic writing",
        "issue_title": "A spelling inconsistency is present",
    }
    causal = {
        "severity": "major",
        "confidence": 0.98,
        "category": "methodological rigour",
        "issue_title": "Causal language is not supported by the design",
    }
    assert _requires_paid_comment_audit(editorial, depth="standard", academic_level="Bachelors") is False
    assert _requires_paid_comment_audit(causal, depth="standard", academic_level="Bachelors") is True


def test_standard_route_stays_on_terra_and_advanced_can_selectively_escalate(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    monkeypatch.setenv("VPROF_ENABLE_OPENAI", "true")
    monkeypatch.setenv("VPROF_ENABLE_DEEPSEEK", "false")
    monkeypatch.setenv("VPROF_COMBINED_APP_PIPELINE", "true")
    monkeypatch.setenv("OPENAI_SECTION_ANALYSIS_MODEL", "gpt-5.6-terra")
    monkeypatch.setenv("OPENAI_FINAL_SYNTHESIS_MODEL", "gpt-5.6-terra")
    monkeypatch.setenv("OPENAI_PHD_FINAL_SYNTHESIS_MODEL", "gpt-5.6-sol")
    config = HybridAIConfig.from_env()
    router = CostAwareAIProvider(config)
    standard = router.plan(stage=ReviewStage.STANDARD_REVIEW, review_depth="standard")
    advanced = router.plan(stage=ReviewStage.ADVANCED_REVIEW, review_depth="advanced")
    final = router.plan(stage=ReviewStage.FINAL_AUDIT, review_depth="standard")
    assert standard.primary.model == "gpt-5.6-terra"
    assert standard.allow_escalation is False
    assert advanced.allow_escalation is True
    assert final.primary.model == "gpt-5.6-terra"
