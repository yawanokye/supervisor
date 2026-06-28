from io import BytesIO

from docx import Document

from app.document_parser import (
    detect_document_chapter_profile,
    extract_docx,
)


def test_docx_unnumbered_chapter_titles_create_boundaries():
    document = Document()

    p = document.add_paragraph("INTRODUCTION")
    p.style = document.styles["Heading 1"]
    document.add_paragraph("Background to the Study").style = document.styles["Heading 2"]
    document.add_paragraph("Statement of the Problem").style = document.styles["Heading 2"]

    p = document.add_paragraph("LITERATURE REVIEW")
    p.style = document.styles["Heading 1"]
    document.add_paragraph("Conceptual Review").style = document.styles["Heading 2"]

    p = document.add_paragraph("RESEARCH METHODS")
    p.style = document.styles["Heading 1"]
    document.add_paragraph("Research Design").style = document.styles["Heading 2"]

    p = document.add_paragraph("RESULTS AND DISCUSSION")
    p.style = document.styles["Heading 1"]
    document.add_paragraph("Introduction").style = document.styles["Heading 2"]
    document.add_paragraph("Results").style = document.styles["Heading 2"]

    p = document.add_paragraph(
        "SUMMARY, CONCLUSIONS AND RECOMMENDATIONS"
    )
    p.style = document.styles["Heading 1"]
    document.add_paragraph("Summary of Findings").style = document.styles["Heading 2"]

    buffer = BytesIO()
    document.save(buffer)

    rows = extract_docx(buffer.getvalue())
    profile = detect_document_chapter_profile(rows)

    assert profile["detected_chapters"] == [1, 2, 3, 4, 5]
    assert profile["numbering_used"] is False

    chapter_four_introduction = next(
        row
        for row in rows
        if row["text"] == "Introduction"
        and row["chapter_number"] == 4
    )
    assert chapter_four_introduction["chapter_number"] == 4


def test_sentence_case_heading_one_titles_are_accepted():
    document = Document()
    for title in (
        "Introduction",
        "Literature Review",
        "Research Methods",
        "Results and Discussion",
        "Summary, Conclusions and Recommendations",
    ):
        paragraph = document.add_paragraph(title)
        paragraph.style = document.styles["Heading 1"]
        document.add_paragraph("Section content.")

    buffer = BytesIO()
    document.save(buffer)
    rows = extract_docx(buffer.getvalue())
    profile = detect_document_chapter_profile(rows)

    assert profile["detected_chapters"] == [1, 2, 3, 4, 5]
