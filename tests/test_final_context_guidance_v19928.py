from __future__ import annotations

import io
import re

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.document_parser import extract_docx
from app.final_review_quality import build_canonical_finding_rows


def _row(fid: str, paragraph: int, *, section: str, item: str, comment: str, action: str, severity: str = "major"):
    return {
        "finding_id": fid,
        "status": "does_not_meet_requirement",
        "severity": severity,
        "confidence": 0.96,
        "chapter_number": 1,
        "section": section,
        "section_reference": section,
        "category": "other",
        "item": item,
        "comment": comment,
        "required_action": action,
        "evidence": [{
            "document_role": "current",
            "chapter_number": 1,
            "paragraph": paragraph,
            "heading": section,
            "section_reference": section,
            "section_path": ["CHAPTER ONE", section],
            "text": comment,
        }],
        "annotation_eligible": True,
    }


def _sample_runtime():
    return [
        {"paragraph": 1, "chapter_number": 1, "heading": "CHAPTER ONE", "is_heading": True, "text": "CHAPTER ONE", "section_path": ["CHAPTER ONE"]},
        {"paragraph": 2, "chapter_number": 1, "heading": "Introduction", "is_heading": True, "text": "Introduction", "section_path": ["CHAPTER ONE", "Introduction"]},
        {"paragraph": 3, "chapter_number": 1, "heading": "Introduction", "text": "Internal controls support fraud prevention (Amanamah, 2024), (Amanamah, 2024).", "section_path": ["CHAPTER ONE", "Introduction"]},
        {"paragraph": 4, "chapter_number": 1, "heading": "Background of the Study", "is_heading": True, "text": "Background of the Study", "section_path": ["CHAPTER ONE", "Background of the Study"]},
        {"paragraph": 5, "chapter_number": 1, "heading": "Background of the Study", "text": "The study moves between commercial banks and rural banks in Ghana.", "section_path": ["CHAPTER ONE", "Background of the Study"]},
        {"paragraph": 6, "chapter_number": 1, "heading": "Problem Statement", "is_heading": True, "text": "Problem Statement", "section_path": ["CHAPTER ONE", "Problem Statement"]},
        {"paragraph": 7, "chapter_number": 1, "heading": "Problem Statement", "text": "Fraud continues, but the scale of the problem is not stated.", "section_path": ["CHAPTER ONE", "Problem Statement"]},
        {"paragraph": 8, "chapter_number": 1, "heading": "General Objective", "is_heading": True, "text": "General Objective", "section_path": ["CHAPTER ONE", "Research Objectives", "General Objective"]},
        {"paragraph": 9, "chapter_number": 1, "heading": "General Objective", "text": "The primary objective of this study is to examine the effect of internal controls on fraud detection and prevention among rural banks in Ghana, specifically Assinman Rural Bank PLC.", "section_path": ["CHAPTER ONE", "Research Objectives", "General Objective"]},
        {"paragraph": 10, "chapter_number": 1, "heading": "Specific Objectives", "is_heading": True, "text": "Specific Objectives", "section_path": ["CHAPTER ONE", "Research Objectives", "Specific Objectives"]},
        {"paragraph": 11, "chapter_number": 1, "heading": "Specific Objectives", "text": "To assess the internal control mechanisms used for fraud detection. To determine the effectiveness of the controls. To examine the impact of behavioural and institutional factors (Pressure, Opportunity and Rationalisation) on fraud incidence at Assinman Rural Bank PLC.", "section_path": ["CHAPTER ONE", "Research Objectives", "Specific Objectives"]},
        {"paragraph": 12, "chapter_number": 1, "heading": "Research Questions", "is_heading": True, "text": "Research Questions", "section_path": ["CHAPTER ONE", "Research Questions"]},
        {"paragraph": 13, "chapter_number": 1, "heading": "Research Questions", "text": "What is the rate of successful fraud detection and prevention?", "section_path": ["CHAPTER ONE", "Research Questions"]},
        {"paragraph": 14, "chapter_number": 1, "heading": "Scope of the Study", "is_heading": True, "text": "Scope of the Study", "section_path": ["CHAPTER ONE", "Scope of the Study"]},
        {"paragraph": 15, "chapter_number": 1, "heading": "Scope of the Study", "text": "The study focuses on internal controls and fraud prevention at Assinman Rural Bank PLC.", "section_path": ["CHAPTER ONE", "Scope of the Study"]},
        {"paragraph": 16, "chapter_number": 1, "heading": "Limitations of the Study", "is_heading": True, "text": "Limitations of the Study", "section_path": ["CHAPTER ONE", "Limitations of the Study"]},
        {"paragraph": 17, "chapter_number": 1, "heading": "Limitations of the Study", "text": "The findings will be sufficient for generalization to all rural banks.", "section_path": ["CHAPTER ONE", "Limitations of the Study"]},
    ]


def test_semantic_gate_removes_statistical_output_guidance_from_problem_and_limitations():
    rows = _sample_runtime()
    problem = _row(
        "P1", 7, section="Problem Statement",
        item="The problem statement needs evidence",
        comment="The problem is described generally. Report the coefficient, standard error and confidence interval.",
        action="Add current evidence. Also report model diagnostics and degrees of freedom.",
    )
    problem["illustrative_guidance"] = "Report the factor loading, reliability coefficient and p-value."
    limitation = _row(
        "L1", 17, section="Limitations of the Study",
        item="Generalisation is overstated",
        comment="The single case cannot support broad generalisation. Report the regression coefficient and standard error.",
        action="Qualify the generalisation and provide the model diagnostic.",
    )
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Effect of Internal Controls on Fraud Detection and Prevention at Assinman Rural Bank PLC"},
        "academic_findings": [problem, limitation],
        "_runtime_context": {"current_paragraphs": rows},
    }
    result = build_canonical_finding_rows(review, force=True)
    combined = " ".join(row["student_comment"] for row in result).lower()
    for forbidden in ("standard error", "confidence interval", "factor loading", "degrees of freedom", "model diagnostic"):
        assert forbidden not in combined
    assert "assiman" not in combined  # guard against accidental corruption of the setting
    assert "assinman rural bank plc" in combined


def test_scope_completeness_audit_names_missing_boundaries_and_gives_context_guidance():
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Effect of Internal Controls on Fraud Detection and Prevention at Assinman Rural Bank PLC"},
        "academic_findings": [],
        "_runtime_context": {"current_paragraphs": _sample_runtime()},
    }
    rows = build_canonical_finding_rows(review, force=True)
    scope = next(row for row in rows if row.get("finding_id") == "HUMAN-SCOPE-COMPLETENESS")
    text = scope["student_comment"].lower()
    assert "unit of analysis" in text
    assert "study period" in text or "period covered" in text
    assert "staff" in text and "institutional records" in text
    assert "assinman rural bank plc" in text


def test_missing_purpose_hypothesis_and_definition_examples_use_current_study():
    runtime = _sample_runtime()
    findings = []
    for fid, label, paragraph, section, action in (
        ("M1", "Purpose of the Study", 7, "Problem Statement", "Add a Purpose of the Study section after the problem statement."),
        ("M2", "Research Hypotheses", 13, "Research Questions", "Add a Research Hypotheses section after the research questions."),
        ("M3", "Definition of Terms", 15, "Scope of the Study", "Add a Definition of Terms section before the organisation of the study."),
    ):
        row = _row(fid, paragraph, section=section, item=f"{label} is missing from Chapter One", comment=f"{label} is missing from Chapter One.", action=action)
        row.update({"missing_section_label": label, "section_contract_label": label, "section_contract_verified": True, "section_status": "missing"})
        findings.append(row)
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Effect of Internal Controls on Fraud Detection and Prevention at Assinman Rural Bank PLC"},
        "academic_findings": findings,
        "_runtime_context": {"current_paragraphs": runtime},
    }
    rows = build_canonical_finding_rows(review, force=True)
    comments = {row["missing_section_label"]: row["student_comment"] for row in rows if row.get("missing_section_label")}
    assert "internal controls" in comments["Purpose of the Study"].lower()
    assert "assinman rural bank plc" in comments["Purpose of the Study"].lower()
    assert "pressure" in comments["Research Hypotheses"].lower()
    assert "fraud incidence" in comments["Research Hypotheses"].lower()
    assert "internal controls" in comments["Definition of Terms"].lower()
    assert "rationalisation" in comments["Definition of Terms"].lower()
    assert not any(text.endswith(".\”.") or text.endswith('.".') for text in comments.values())


def test_objective_alignment_example_distinguishes_descriptive_measurement_and_causal_claims():
    finding = _row(
        "A1", 11, section="Research Objectives",
        item="The objectives and questions are not aligned",
        comment="The study mixes descriptive and inferential aims.",
        action="Classify each objective and match it to a question or hypothesis and analysis.",
    )
    review = {
        "summary": {"academic_level": "MPhil", "study_title": "Effect of Internal Controls on Fraud Detection and Prevention at Assinman Rural Bank PLC"},
        "academic_findings": [finding],
        "_runtime_context": {"current_paragraphs": _sample_runtime()},
    }
    row = next(row for row in build_canonical_finding_rows(review, force=True) if row.get("finding_id") == "A1")
    text = row["student_comment"].lower()
    assert "descriptive" in text
    assert "rate" in text and "calculated" in text
    assert "pressure" in text and "predictive" in text
    assert "causal" in text


def test_every_numbered_actionable_finding_has_a_current_anchor_and_visible_numbering_begins_at_one(monkeypatch):
    monkeypatch.setenv("VPROF_COMMENT_MERGE_BY_SECTION", "false")
    monkeypatch.setenv("VPROF_EXPORT_ONE_COMMENT_PER_FINDING", "true")
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("Internal controls support fraud prevention (Amanamah, 2024), (Amanamah, 2024).")
    doc.add_heading("Problem Statement", level=2)
    doc.add_paragraph("The problem is stated generally without local evidence.")
    source = io.BytesIO(); doc.save(source)
    parsed = extract_docx(source.getvalue())
    unanchored = {
        "finding_id": "U1", "status": "does_not_meet_requirement", "severity": "major", "confidence": 0.9,
        "chapter_number": 1, "section": "Introduction", "section_reference": "Introduction",
        "category": "citation_and_reference_integrity", "item": "Source attribution requires correction",
        "comment": "Several claims need source verification.", "required_action": "Verify every citation and correct the citation presentation.",
        "annotation_eligible": False,
    }
    anchored = _row("U2", 5, section="Problem Statement", item="The problem needs evidence", comment="The problem is stated generally.", action="Add local evidence.")
    target = next(row for row in parsed if "without local evidence" in row.get("text", ""))
    anchored["evidence"] = [{**target, "document_role": "current"}]
    review = {
        "summary": {"academic_level": "MPhil", "reviewer_name": "Supervisor", "study_title": "Internal Controls and Fraud Prevention"},
        "academic_findings": [anchored, unanchored],
        "_runtime_context": {"current_paragraphs": parsed},
    }
    canonical = build_canonical_finding_rows(review, force=True)
    assert [row["finding_number"] for row in canonical] == list(range(1, len(canonical) + 1))
    assert all(row.get("annotation_eligible") is not False for row in canonical)
    assert all(row.get("export_anchor_verified") for row in canonical)
    output = build_annotated_docx(source.getvalue(), review)
    reviewed = Document(io.BytesIO(output))
    body = "\n".join(paragraph.text for paragraph in reviewed.paragraphs)
    assert "[1]" not in body
    assert not re.search(r"[A-Za-z]\s*\[\d+\]\s*[A-Za-z]", body)
    assert len(list(reviewed.comments)) >= 1
    assert "w:commentRangeStart" in reviewed.element.body.xml


def test_problematic_quote_is_clipped_at_a_word_boundary():
    long_quote = "This objective is intended to examine the relationship between internal controls and fraud prevention among rural banks in Ghana and must be aligned with the research question and analysis plan before the study proceeds further into the methodology chapter."
    finding = _row("Q1", 9, section="Research Objectives", item="The objective needs alignment", comment="The objective does not match the question.", action="Align the objective and question.")
    finding["problematic_quote"] = long_quote * 3
    review = {"summary": {"academic_level": "MPhil"}, "academic_findings": [finding]}
    row = build_canonical_finding_rows(review, force=True)[0]
    quote = row["problematic_quote"]
    assert len(quote) <= 261
    assert not quote.endswith(" amon")
    assert quote.endswith((".", "?", "!", "…"))
