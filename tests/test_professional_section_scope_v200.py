from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.document_parser import extract_docx, parse_document
from app.inline_annotated_exporter import build_inline_annotated_docx
from app.report_exporter import build_docx_report
from app.review_scope import (
    apply_review_scope_filter,
    apply_selected_section_scope,
    build_document_outline,
)
from app.submission_readiness import build_supervisory_readiness


def _bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _chapter_source() -> bytes:
    document = Document()
    document.add_heading("CHAPTER ONE", level=0)
    document.add_heading("1.1 Background of the Study", level=1)
    document.add_paragraph("The background establishes the broad context.")
    document.add_heading("1.2 Problem Statement", level=1)
    document.add_paragraph("The problem is important but the evidence is limited. A second sentence states the proposed response.")
    document.add_heading("1.3 Research Objectives", level=1)
    document.add_paragraph("The study examines the stated problem.")
    return _bytes(document)


def test_outline_detects_real_sections_and_scope_isolates_selected_heading():
    rows = parse_document(_chapter_source(), "chapter-one.docx")
    outline = build_document_outline(rows)
    chapter = outline["chapters"][0]
    titles = [item["section_title"] for item in chapter["sections"]]
    assert titles == [
        "1.1 Background of the Study",
        "1.2 Problem Statement",
        "1.3 Research Objectives",
    ]

    selected_key = chapter["sections"][1]["section_key"]
    scoped, metadata = apply_selected_section_scope(
        rows,
        selected_chapter=1,
        section_scope_mode="selected_sections",
        selected_sections=[selected_key],
    )
    text = "\n".join(row["text"] for row in scoped)
    assert "1.2 Problem Statement" in text
    assert "The problem is important" in text
    assert "broad context" not in text
    assert "Research Objectives" not in text
    assert metadata["selection_precision"] == "detected_heading_boundaries"


def test_scope_filter_removes_out_of_scope_findings_and_rebuilds_derived_packages():
    review = {
        "summary": {
            "selected_section_scope": {
                "mode": "selected_sections",
                "selected_paragraph_numbers": [4, 5],
                "selected_section_titles": ["1.2 Problem Statement"],
            }
        },
        "academic_findings": [
            {
                "status": "partly_meets_requirement",
                "severity": "major",
                "section": "1.2 Problem Statement",
                "required_action": "Add evidence.",
                "evidence": [{"paragraph": 5}],
            },
            {
                "status": "partly_meets_requirement",
                "severity": "major",
                "section": "1.1 Background",
                "required_action": "Rewrite the background.",
                "evidence": [{"paragraph": 3}],
            },
        ],
        "professional_review": {"finding_ledger": ["stale"]},
        "articleready_quality_audit": {"audit_rows": ["stale"]},
    }
    filtered = apply_review_scope_filter(review)
    assert len(filtered["academic_findings"]) == 1
    assert filtered["academic_findings"][0]["section"] == "1.2 Problem Statement"
    assert "professional_review" not in filtered
    assert "articleready_quality_audit" not in filtered


def _anchored_review(source: bytes):
    rows = extract_docx(source)
    evidence = next(row for row in rows if "problem is important" in row.get("text", ""))
    base = {
        "status": "partly_meets_requirement",
        "severity": "major",
        "section": "1.2 Problem Statement",
        "section_reference": "1.2 Problem Statement",
        "category": "research_gap_and_problem",
        "evidence": [{**evidence, "document_role": "current"}],
        "annotation_eligible": True,
    }
    return {
        "summary": {"reviewer_name": "Prof Anokye Mohammed Adam"},
        "academic_findings": [
            {
                **base,
                "finding_id": "F-1",
                "issue_title": "Evidence is not demonstrated",
                "required_action": "Add verifiable evidence showing the scale of the problem.",
                "problematic_quote": "problem is important",
            },
            {
                **base,
                "finding_id": "F-2",
                "issue_title": "The unresolved issue is unclear",
                "required_action": "Explain what remains unresolved after the evidence is presented.",
                "problematic_quote": "evidence is limited",
            },
            {
                **base,
                "finding_id": "F-3",
                "issue_title": "The proposed response is disconnected",
                "required_action": "Link the second sentence directly to the unresolved problem.",
                "problematic_quote": "A second sentence states the proposed response",
            },
        ],
    }


def test_native_comments_group_same_sentence_but_keep_different_sentence_separate(monkeypatch):
    monkeypatch.setenv("VPROF_HUMAN_ROOT_CAUSE_CONSOLIDATION", "false")
    monkeypatch.setenv("VPROF_NATIVE_COMMENT_STYLE", "exact_anchor_grouped")
    monkeypatch.setenv("VPROF_EXPORT_ONE_COMMENT_PER_FINDING", "false")
    monkeypatch.setenv("VPROF_NATIVE_GROUP_LOCATION_MARKERS", "false")
    source = _chapter_source()
    output = Document(io.BytesIO(build_annotated_docx(source, _anchored_review(source))))
    comments = list(output.comments)
    assert len(comments) == 2
    text = "\n".join(comment.text for comment in comments)
    assert "Evidence is not demonstrated" in text
    assert "The unresolved issue is unclear" in text
    assert "The proposed response is disconnected" in text
    assert any(sum(f"{number}. " in comment.text for number in (1, 2, 3)) >= 2 for comment in comments)


def test_inline_comments_follow_affected_paragraph_and_group_same_sentence(monkeypatch):
    monkeypatch.setenv("VPROF_HUMAN_ROOT_CAUSE_CONSOLIDATION", "false")
    source = _chapter_source()
    output = Document(io.BytesIO(build_inline_annotated_docx(source, _anchored_review(source))))
    notes = [p.text for p in output.paragraphs if p.text.startswith("Detailed supervisor comment:")]
    assert len(notes) == 2
    assert any("Evidence is not demonstrated" in note and "The unresolved issue is unclear" in note for note in notes)
    target_index = next(i for i, p in enumerate(output.paragraphs) if "problem is important" in p.text)
    assert output.paragraphs[target_index + 1].text.startswith("Detailed supervisor comment:")


def test_submission_readiness_and_report_give_direct_action_schedule(monkeypatch):
    review = {
        "summary": {"document_label": "Chapter One", "academic_level": "Research Masters / MPhil"},
        "academic_findings": [
            {
                "finding_number": 1,
                "status": "does_not_meet_requirement",
                "severity": "major",
                "category": "statistical_accuracy",
                "issue_title": "The reported p-value conflicts with the significance claim",
                "required_action": "Correct the p-value or revise the significance interpretation using the original output.",
                "academic_consequence": "The present claim may lead to an incorrect hypothesis decision.",
                "evidence": [{"paragraph": 22, "section_reference": "4.2 Regression Results", "page": 12}],
            }
        ],
        "statistical_review": {
            "verified_inconsistency_count": 1,
            "reporting_omission_count": 0,
            "consistency_warnings": [],
        },
    }
    readiness = build_supervisory_readiness(review)
    assert readiness["status"] == "Not ready for supervisor approval"
    assert readiness["actions"][0]["priority"] == "Essential before approval"
    assert "original dataset" in readiness["actions"][0]["verification"]

    monkeypatch.setenv("VPROF_SPEC_ALIGNED_SUPERVISORY_REPORT", "true")
    report = Document(io.BytesIO(build_docx_report(review)))
    report_text = "\n".join(p.text for p in report.paragraphs)
    report_text += "\n" + "\n".join(cell.text for table in report.tables for row in table.rows for cell in row.cells)
    assert "Actions Required Before Supervisor Approval or Submission" in report_text
    assert "Correct the p-value" in report_text
    assert "How to verify completion" in report_text


def test_frontend_contains_section_scope_controls_and_outline_endpoint_call():
    root = Path(__file__).resolve().parents[1]
    html = (root / "app/templates/index.html").read_text(encoding="utf-8")
    js = (root / "app/static/app.js").read_text(encoding="utf-8")
    assert 'id="sectionScopeCard"' in html
    assert 'name="section_scope_mode"' in html
    assert 'id="selectedSectionsJson"' in html
    assert 'fetch("/api/review/outline"' in js
    assert "selected_sections_json" in js
