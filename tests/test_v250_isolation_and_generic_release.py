from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.deterministic_supervisory_checklist import hard_chapter_one_supervisory_issues
from app.document_parser import parse_document
from app.natural_supervisor_comment import natural_supervisor_comment
from app.review_isolation import enforce_current_submission_isolation
from app.review_release_guard import filter_and_rewrite_release_findings


def _finding(title: str, section: str, paragraph: int = 3, action: str = "") -> dict:
    return {
        "finding_id": title,
        "chapter_number": 1,
        "section": section,
        "section_reference": section,
        "item": title,
        "issue_title": title,
        "assessment": title,
        "comment": title,
        "academic_consequence": "The weakness affects the coherence of the study.",
        "required_action": action or "State the missing information directly in the relevant section using the actual design, evidence and terminology of the study.",
        "verification_test": "Confirm that that Use an alignment table.",
        "severity": "major",
        "status": "does_not_meet_requirement",
        "confidence": 0.95,
        "evidence": [{
            "paragraph": paragraph,
            "paragraph_id": f"P{paragraph}",
            "text": "This is the marked passage.",
            "document_role": "current",
            "section_reference": section,
            "chapter_number": 1,
        }],
    }


def _runtime_review() -> dict:
    rows = [
        {"chapter_number": 1, "is_heading": True, "text": "CHAPTER ONE", "section_reference": "CHAPTER ONE", "paragraph": 1},
        {"chapter_number": 1, "is_heading": True, "text": "INTRODUCTION", "section_reference": "INTRODUCTION", "paragraph": 2},
        {"chapter_number": 1, "text": "This chapter introduces digital inventory practices in the selected market and outlines the background, problem, objectives, questions, significance, scope, limitations and organisation of the study.", "section_reference": "INTRODUCTION", "paragraph": 3},
        {"chapter_number": 1, "is_heading": True, "text": "Background to the Study", "section_reference": "Background to the Study", "paragraph": 4},
        {"chapter_number": 1, "text": "Digital inventory systems support stock control in many settings (Example, 2024). The present study focuses on the selected market and its order-processing constraints.", "section_reference": "Background to the Study", "paragraph": 5},
        {"chapter_number": 1, "is_heading": True, "text": "Organization of the Study", "section_reference": "Organization of the Study", "paragraph": 6},
        {"chapter_number": 1, "text": "Chapter Two reviews literature and the conceptual framework. Chapter Three presents the methodology. Chapter Four reports the findings. Chapter Five presents conclusions and recommendations.", "section_reference": "Organization of the Study", "paragraph": 7},
    ]
    return {
        "summary": {"submission_scope": "chapter", "academic_level": "Bachelors"},
        "_runtime_context": {"current_paragraphs": rows},
    }


def test_example_content_is_not_persisted_or_used_as_production_rule():
    app_text = "\n".join(
        path.read_text(encoding="utf-8", errors="ignore")
        for path in Path("app").glob("*.py")
    ).lower()
    for literal in ("aboabo", "tamale", "assinman", "rickardo", "gladson"):
        assert literal not in app_text

    review = {
        "summary": {},
        "example_review": {"topic": "discard me"},
        "learned_rules": ["discard me"],
    }
    runtime = {
        "current_paragraphs": [{"text": "Current study evidence"}],
        "previous_paragraphs": [{"text": "Earlier chapter in the same submission"}],
        "sample_context": {"topic": "discard me"},
    }
    isolated_review, isolated_runtime = enforce_current_submission_isolation(
        review, runtime, document_hash="abcdef1234567890"
    )
    assert "example_review" not in isolated_review
    assert "learned_rules" not in isolated_review
    assert "sample_context" not in isolated_runtime
    assert isolated_runtime["current_paragraphs"]
    assert isolated_runtime["previous_paragraphs"]
    assert isolated_review["summary"]["cross_submission_learning"] is False
    assert isolated_review["summary"]["example_content_persisted"] is False


def test_present_sections_are_not_described_as_missing():
    review = _runtime_review()
    findings = [
        _finding("Bare chapter and introduction headings", "INTRODUCTION"),
        _finding("Missing content in Organization of the Study", "Organization of the Study", 7),
        _finding("Missing theoretical framework outline", "Organization of the Study", 7),
    ]
    released = filter_and_rewrite_release_findings(findings, review)
    assert len(released) == 1
    assert released[0]["item"] == "The chapter introduction is present but functions mainly as an outline"
    assert "missing" not in released[0]["item"].lower()


def test_hypotheses_are_not_required_mechanically_for_chapter_only_review():
    review = _runtime_review()
    findings = [
        _finding(
            "Relational objectives are stated without corresponding hypotheses",
            "Research Questions",
            action="Formulate hypotheses aligned to the relational objectives.",
        )
    ]
    assert filter_and_rewrite_release_findings(findings, review) == []


def _setting_drift_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("ROLE OF DIGITAL TOOLS IN INVENTORY CONTROL AT KETA CENTRAL MARKET")
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("INTRODUCTION", level=1)
    doc.add_paragraph("This chapter introduces the study.")
    doc.add_heading("Background to the Study", level=2)
    doc.add_paragraph("Digital tools are increasingly used in inventory control.")
    doc.add_heading("Purpose of the Study", level=2)
    doc.add_paragraph("The purpose is to examine digital tools at Keta Central Market.")
    doc.add_heading("Objectives of the Study", level=2)
    doc.add_paragraph("1. To assess digital inventory tools at Ho Central Market.")
    doc.add_heading("Research Questions", level=2)
    doc.add_paragraph("How are digital inventory tools used at Ho Central Market?")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_setting_drift_is_detected_generically_from_current_document():
    rows = parse_document(_setting_drift_docx(), "setting-drift.docx")
    findings = hard_chapter_one_supervisory_issues(
        rows, academic_level="Bachelors", submission_scope="chapter"
    )
    row = next(item for item in findings if item["finding_id"] == "DSC-HARD-B3-STUDY-SETTING-DRIFT")
    assert "KETA CENTRAL MARKET" in row["assessment"].upper()
    assert "HO CENTRAL MARKET" in row["assessment"].upper()
    assert "wider geographical or institutional context" in row["required_action"].lower()


def test_natural_comment_omits_field_labels_generic_actions_and_broken_verification():
    row = _finding("The problem statement is too broad", "Problem Statement")
    text = natural_supervisor_comment(row)
    for phrase in (
        "Issue:", "Problem identified:", "Action required:", "Verification:",
        "State the missing information directly", "Check that that", "Check that Use",
    ):
        assert phrase not in text
    assert text.endswith(".")


def test_existing_source_comments_are_labelled_separately_from_new_findings(monkeypatch):
    monkeypatch.setenv("VPROF_NATIVE_COMMENT_STYLE", "exact_anchor_grouped")
    doc = Document()
    paragraph = doc.add_paragraph("The purpose of this study is to examine digital inventory practices.")
    doc.add_comment(paragraph.runs[0], text="1. Old supervisor advice", author="Supervisor", initials="SP")
    source = io.BytesIO()
    doc.save(source)

    rows = parse_document(source.getvalue(), "commented.docx")
    target = next(row for row in rows if "purpose of this study" in row.get("text", "").lower())
    findings = []
    specifications = (
        ("A source citation in the paragraph is incomplete", "citations_and_sources", "Complete the citation and verify the reference entry."),
        ("The paragraph contains a verb agreement error", "academic_writing", "Correct the verb agreement without changing the intended meaning."),
    )
    for index, (title, category, action) in enumerate(specifications, start=1):
        row = _finding(title, "Purpose of the Study", target["paragraph"], action=action)
        row["finding_id"] = f"CURRENT-{index}"
        row["category"] = category
        row["evidence"] = [{**target, "document_role": "current"}]
        row["problematic_quote"] = target["text"]
        findings.append(row)
    review = {
        "summary": {"submission_scope": "chapter", "academic_level": "Bachelors", "reviewer_name": "V-Professor"},
        "academic_findings": findings,
        "_runtime_context": {"current_paragraphs": rows},
    }
    output = build_annotated_docx(source.getvalue(), review)
    reviewed = Document(io.BytesIO(output))
    comments = list(reviewed.comments)
    previous = [comment for comment in comments if comment.text.startswith("[Previous comment from source document]")]
    current = [comment for comment in comments if comment.author == "V-Professor" and not comment.text.startswith("[Previous comment")]
    assert len(previous) == 1
    assert current
    combined = " ".join(comment.text for comment in current)
    expected = len(review.get("canonical_findings") or [])
    assert expected >= 2
    assert all(f"{number}. " in combined for number in range(1, expected + 1))
