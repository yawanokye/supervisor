from io import BytesIO

from docx import Document

from app.annotated_exporter import build_annotated_docx, native_comment_count


def _source_doc() -> bytes:
    doc = Document()
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("Background to the Study", level=2)
    doc.add_paragraph("Green procurement is discussed in general terms.")
    doc.add_heading("Statement of the Problem", level=2)
    doc.add_paragraph("The problem is not yet fully localised.")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def test_routine_section_assessments_do_not_clutter_native_comments():
    review = {
        "summary": {"reviewer_name": "Supervisor"},
        "academic_section_reviews": [
            {
                "chapter_number": 1,
                "heading": "Background to the Study",
                "section_path": ["Background to the Study"],
                "section_assessment": "This section has been reviewed against the selected academic level.",
            },
            {
                "chapter_number": 1,
                "heading": "Statement of the Problem",
                "section_path": ["Statement of the Problem"],
                "section_assessment": "No major issues were found.",
            },
        ],
        "academic_findings": [],
        "alignment_results": [],
        "revision_results": [],
    }
    exported = build_annotated_docx(_source_doc(), review)
    assert native_comment_count(exported) == 0


def test_section_assessment_alone_is_kept_in_report_not_native_review_pane():
    review = {
        "summary": {"reviewer_name": "Supervisor"},
        "academic_section_reviews": [
            {
                "chapter_number": 1,
                "heading": "Purpose of the study",
                "section_path": ["Purpose of the study"],
                "section_assessment": "Issue: The purpose is narrower than the objectives and should be aligned.",
            },
        ],
        "academic_findings": [],
        "alignment_results": [],
        "revision_results": [],
    }
    exported = build_annotated_docx(_source_doc(), review)
    assert native_comment_count(exported) == 0

