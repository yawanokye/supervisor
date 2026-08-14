import base64
import io

from docx import Document

from app.academic_ai_engine import _batch_model_route
from app.ai_config import HybridAIConfig
from app.conceptual_framework_review import (
    build_quantitative_framework_audit,
    is_quantitative_study,
)
from app.coverage_review import select_ai_review_units
from app.document_parser import extract_docx


def _row(paragraph, text, heading, chapter, **extra):
    return {
        "paragraph": paragraph,
        "text": text,
        "heading": heading,
        "section_path": [heading],
        "chapter_number": chapter,
        "document_role": "current",
        **extra,
    }


def test_quantitative_framework_audit_joins_research_logic_and_diagram():
    rows = [
        _row(1, "To estimate the effect of service quality on loyalty.", "Specific Objectives", 1),
        _row(2, "What is the effect of service quality on loyalty?", "Research Questions", 1),
        _row(3, "H1: Service quality has a positive effect on loyalty.", "Research Hypotheses", 1),
        _row(4, "Conceptual Framework", "Conceptual Framework", 2),
        _row(
            5,
            "Figure 2.1 Conceptual framework: Service quality → loyalty",
            "Conceptual Framework",
            2,
            contains_drawing=True,
            drawing_count=1,
            drawing_descriptions=["Conceptual framework"],
        ),
        _row(6, "The effect will be estimated with OLS regression.", "Data Analysis", 3),
    ]

    audit = build_quantitative_framework_audit(
        rows, [], research_approach="Quantitative"
    )

    assert audit is not None
    assert audit["conceptual_framework_audit"] is True
    assert audit["extra_context"]["embedded_drawing_count"] == 1
    assert {"objective", "question", "hypothesis", "framework", "diagram"}.issubset(
        set(audit["extra_context"]["audited_components"])
    )
    instruction = audit["extra_context"]["instruction"].lower()
    assert "do not require every study to use ols" in instruction
    assert "arrow direction" in instruction


def test_framework_audit_is_not_added_to_a_qualitative_study():
    rows = [
        _row(1, "Conceptual framework", "Conceptual Framework", 2),
        _row(2, "The framework organises the interview themes.", "Conceptual Framework", 2),
    ]
    assert not is_quantitative_study(rows, "Qualitative")
    assert build_quantitative_framework_audit(
        rows, [], research_approach="Qualitative"
    ) is None


def test_selective_preflight_keeps_decisive_units_and_locally_passes_clean_prose():
    units = [
        {
            "section_key": "S001P01",
            "heading": "Opening note",
            "coverage_unit_kind": "prose",
            "target_paragraph_ids": ["P1"],
            "paragraphs": [_row(1, "This chapter begins with an overview.", "Opening note", 1)],
        },
        {
            "section_key": "S002P01",
            "heading": "Conceptual Framework",
            "coverage_unit_kind": "prose",
            "target_paragraph_ids": ["P2"],
            "paragraphs": [_row(2, "The conceptual framework links X and Y.", "Conceptual Framework", 2)],
        },
        {
            "section_key": "S003P01",
            "heading": "Table 4.1 Regression Results",
            "coverage_unit_kind": "table",
            "target_paragraph_ids": ["P3"],
            "paragraphs": [_row(3, "X | 0.42 | 0.01", "Regression Results", 4)],
        },
    ]

    selected, local, stats = select_ai_review_units(
        units,
        academic_level="Bachelor's",
        clean_sample_rate=0.0,
    )

    assert [row["section_key"] for row in selected] == ["S002P01", "S003P01"]
    assert [row["section_key"] for row in local] == ["S001P01"]
    assert stats["local_preflight_units"] == 1


def test_mandatory_framework_audit_routes_to_terra_for_every_degree(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test")
    config = HybridAIConfig.from_env()
    model, effort = _batch_model_route(
        [{"conceptual_framework_audit": True, "heading": "Framework audit"}],
        "Bachelor's",
        config,
    )
    assert model == "gpt-5.6-terra"
    assert effort == "high"


def test_drawing_only_framework_paragraph_is_retained_for_visual_verification(tmp_path):
    # A valid one-pixel PNG is enough to exercise Word's drawing-only paragraph.
    image_path = tmp_path / "framework.png"
    image_path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9Z7E8AAAAASUVORK5CYII="
    ))
    document = Document()
    document.add_heading("CHAPTER TWO", level=1)
    document.add_heading("Conceptual Framework Diagram", level=2)
    document.add_picture(str(image_path))
    output = io.BytesIO()
    document.save(output)

    rows = extract_docx(output.getvalue())
    drawings = [row for row in rows if row.get("contains_drawing")]

    assert len(drawings) == 1
    assert drawings[0]["source_kind"] == "figure"
    assert drawings[0]["text"].startswith("Embedded figure or diagram")
    audit = build_quantitative_framework_audit(
        rows, [], research_approach="Quantitative"
    )
    assert audit is not None
    assert "diagram" in audit["extra_context"]["audited_components"]
    assert audit["extra_context"]["diagram_images_supplied_to_expert"] == 1
