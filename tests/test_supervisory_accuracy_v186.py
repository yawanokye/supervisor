from __future__ import annotations

import io

from docx import Document

from app.document_parser import extract_docx
from app.supervisory_accuracy_guard import (
    apply_accuracy_gate,
    build_factual_index,
    deterministic_expert_issues,
    guard_issue,
    paragraph_id,
)


def _docx_bytes(document: Document) -> bytes:
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _base_issue(**updates):
    value = {
        "finding_id": "F1",
        "category": "results_and_interpretation",
        "section": "Whole-chapter coherence and consistency audit",
        "issue_title": "Objective Two: Effects of E-tendering on Procurement Efficiency",
        "severity": "major",
        "confidence": 0.94,
        "evidence_paragraph_ids": ["P2"],
        "problematic_quote": "",
        "assessment": "The full results are missing.",
        "academic_consequence": "The objective is not answered.",
        "required_action": "Provide a full results section for this objective.",
        "illustrative_guidance": "",
        "guidance_type": "direct_correction",
        "source_verification_required": False,
        "context_guard_adjusted": False,
    }
    value.update(updates)
    return value


def test_false_missing_results_comment_is_rejected_when_section_exists():
    paragraphs = [
        {
            "paragraph": 1,
            "text": "KEYWORDS",
            "heading": "KEYWORDS",
            "section_path": ["KEYWORDS"],
            "is_heading": True,
            "chapter_number": None,
            "document_role": "current",
        },
        {
            "paragraph": 2,
            "text": "PROCUREMENT EFFICIENCY",
            "heading": "KEYWORDS",
            "section_path": ["KEYWORDS"],
            "is_heading": False,
            "chapter_number": None,
            "document_role": "current",
        },
        {
            "paragraph": 3,
            "text": "Effects of E-tendering on Procurement Efficiency",
            "heading": "Effects of E-tendering on Procurement Efficiency",
            "section_path": ["CHAPTER FOUR", "Effects of E-tendering on Procurement Efficiency"],
            "is_heading": True,
            "chapter_number": 4,
            "document_role": "current",
        },
        {
            "paragraph": 4,
            "text": "The correlation was positive and statistically significant.",
            "heading": "Effects of E-tendering on Procurement Efficiency",
            "section_path": ["CHAPTER FOUR", "Effects of E-tendering on Procurement Efficiency"],
            "is_heading": False,
            "chapter_number": 4,
            "document_role": "current",
        },
        {
            "paragraph": 5,
            "text": "Table 10: E-Tendering Correlations",
            "heading": "Effects of E-tendering on Procurement Efficiency",
            "section_path": ["CHAPTER FOUR", "Effects of E-tendering on Procurement Efficiency"],
            "is_heading": False,
            "chapter_number": 4,
            "document_role": "current",
            "source_kind": "table_row",
            "table_index": 1,
            "table_row": 1,
            "table_number": "10",
            "table_title": "E-Tendering Correlations",
        },
    ]
    index = {paragraph_id(row): row for row in paragraphs}
    assert guard_issue(_base_issue(), index, build_factual_index(paragraphs)) is None


def test_synthetic_audit_finding_is_reanchored_to_exact_source_section():
    paragraphs = [
        {
            "paragraph": 1,
            "text": "The sampling technique was convenience sampling.",
            "heading": "Sampling Procedures and Sample Size",
            "section_path": ["CHAPTER THREE", "Sampling Procedures and Sample Size"],
            "is_heading": False,
            "chapter_number": 3,
            "document_role": "current",
        }
    ]
    issue = _base_issue(
        category="methodological_rigour",
        issue_title="Sampling generalisability requires clarification",
        evidence_paragraph_ids=["P1"],
        assessment="The non-probability selection limits representativeness.",
        required_action="State the inferential limitation and restrict generalisation.",
    )
    index = {"P1": paragraphs[0]}
    guarded = guard_issue(issue, index, build_factual_index(paragraphs))
    assert guarded is not None
    assert guarded["section"] == "Sampling Procedures and Sample Size"


def test_wrong_table_number_is_replaced_from_evidence_metadata():
    paragraph = {
        "paragraph": 1,
        "text": "E-Invoicing | 1 | .658 | 0.000",
        "heading": "Correlation Analysis between E-Invoicing and Procurement Efficiency",
        "section_path": ["CHAPTER FOUR", "Correlation Analysis between E-Invoicing and Procurement Efficiency"],
        "is_heading": False,
        "chapter_number": 4,
        "document_role": "current",
        "source_kind": "table_row",
        "table_index": 17,
        "table_row": 3,
        "table_number": "18",
        "table_title": "E-Invoicing Correlations",
    }
    issue = _base_issue(
        section=paragraph["heading"],
        issue_title="Table 17 requires correction",
        evidence_paragraph_ids=["P1"],
        assessment="Table 17 reports p = .000.",
        required_action="Correct Table 17 and report p < .001.",
    )
    guarded = guard_issue(issue, {"P1": paragraph}, build_factual_index([paragraph]))
    assert guarded is not None
    assert "Table 18" in guarded["issue_title"]
    assert "Table 17" not in guarded["required_action"]
    assert guarded["canonical_table_number"] == "18"


def test_embedded_table_caption_overrides_physical_table_order():
    document = Document()
    document.add_heading("CHAPTER FOUR", level=1)
    document.add_heading("E-Invoicing Results", level=2)
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Table 18: E-Invoicing Correlations"
    table.cell(0, 1).text = "Table 18: E-Invoicing Correlations"
    table.cell(1, 0).text = "Variable"
    table.cell(1, 1).text = "Coefficient"
    table.cell(2, 0).text = "E-Invoicing"
    table.cell(2, 1).text = ".658"

    rows = extract_docx(_docx_bytes(document))
    table_rows = [row for row in rows if row.get("source_kind") == "table_row"]
    assert table_rows
    assert all(row.get("table_number") == "18" for row in table_rows)
    assert all(row.get("table_title") == "E-Invoicing Correlations" for row in table_rows)


def test_deterministic_expert_checks_identify_method_result_mismatch():
    paragraphs = [
        {
            "paragraph": 1,
            "text": "Multiple regression was used to estimate the unique effect of each predictor while controlling for the others.",
            "heading": "Data Processing and Analysis",
            "section_path": ["CHAPTER THREE", "Data Processing and Analysis"],
            "chapter_number": 3,
            "document_role": "current",
        },
        {
            "paragraph": 2,
            "text": "The predictors were entered simultaneously in the model.",
            "heading": "Data Processing and Analysis",
            "section_path": ["CHAPTER THREE", "Data Processing and Analysis"],
            "chapter_number": 3,
            "document_role": "current",
        },
        {
            "paragraph": 3,
            "text": "b. Predictors: (Constant), E-Sourcing",
            "heading": "Effect of E-Sourcing on Procurement Efficiency",
            "section_path": ["CHAPTER FOUR", "Effect of E-Sourcing on Procurement Efficiency"],
            "chapter_number": 4,
            "document_role": "current",
        },
        {
            "paragraph": 4,
            "text": "b. Predictors: (Constant), E-Tendering",
            "heading": "Effect of E-Tendering on Procurement Efficiency",
            "section_path": ["CHAPTER FOUR", "Effect of E-Tendering on Procurement Efficiency"],
            "chapter_number": 4,
            "document_role": "current",
        },
    ]
    issues = deterministic_expert_issues(paragraphs)
    assert any(issue["finding_id"] == "DET-MULTIPLE-REGRESSION-MISMATCH" for issue in issues)


def test_accuracy_gate_reports_dropped_false_findings():
    paragraphs = [
        {
            "paragraph": 1,
            "text": "CHAPTER THREE",
            "heading": "CHAPTER THREE",
            "section_path": ["CHAPTER THREE"],
            "is_heading": True,
            "chapter_number": 3,
            "document_role": "current",
        },
        *[
            {
                "paragraph": index,
                "text": f"Methodology content {index}",
                "heading": "Research Methods",
                "section_path": ["CHAPTER THREE", "Research Methods"],
                "chapter_number": 3,
                "document_role": "current",
            }
            for index in range(2, 8)
        ],
    ]
    index = {paragraph_id(row): row for row in paragraphs}
    issue = _base_issue(
        evidence_paragraph_ids=["P1"],
        issue_title="Write a complete methodology chapter",
        assessment="The methodology chapter is missing.",
        required_action="Write a complete methodology chapter.",
    )
    kept, stats = apply_accuracy_gate([issue], index, paragraphs)
    assert kept == []
    assert stats["dropped"] == 1
