from __future__ import annotations

import io

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.deterministic_supervisory_checklist import hard_chapter_one_supervisory_issues
from app.document_parser import extract_docx
from app.evidence_ledger import evidence_ledger_rows, sentence_spans
from app.submission_readiness import build_supervisory_readiness


def _bytes(document: Document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _chapter_one_source(citations: int = 2) -> bytes:
    document = Document()
    document.add_paragraph(
        "IMPACT OF EFFECTIVE PLANNING AND CONTROL IN PROCUREMENT ACTIVITIES OF A MANUFACTURING FIRM"
    )
    document.add_heading("CHAPTER ONE", level=0)
    document.add_heading("INTRODUCTION", level=1)
    document.add_paragraph("This chapter presents the background and problem.")
    document.add_heading("1.1 Background to the Study", level=1)
    document.add_paragraph("Delete all numbering from the work. This is not UCC style.")
    citation_text = " ".join(f"Evidence was reported (Author{i}, 20{10+i})." for i in range(citations))
    document.add_paragraph(
        "Prior work recognized procurement in organizations (Edwards, 2013). "
        "Control Procurements is the process of monitoring contracts. "
        "A sentence ends with an incomplete source (Smith, " + citation_text
    )
    document.add_heading("1.2 Statement of the Problem", level=1)
    document.add_paragraph(
        "Rotich (2011) reports failures in the Kenya financial sector. "
        "Despite the significant impact of planning and control, researches on these constructs "
        "have not received much attention and the literature is scanty."
    )
    document.add_heading("1.3 Purpose of the Study", level=1)
    document.add_paragraph(
        "The purpose is to assess the impact of planning and control in a manufacturing firm "
        "and suggest measures firms can adopt."
    )
    document.add_heading("1.4 Objectives of the Study", level=1)
    document.add_paragraph(
        "1. To assess current practices. 2. To identify the effects of planning and control. "
        "3. To examine challenges."
    )
    document.add_heading("1.5 Research Questions", level=1)
    document.add_paragraph(
        "What is the current practices of planning and control in manufacturing firms?"
    )
    document.add_heading("1.6 Significance of the Study", level=1)
    document.add_paragraph("The organisation will benefit students, policy makers and managers.")
    document.add_heading("1.7 Limitation of the Study", level=1)
    document.add_paragraph(
        "The study will be limited to manufacturing firms and may not be generalized to other industries."
    )
    return _bytes(document)


def test_chapter_one_contract_detects_material_supplied_weaknesses():
    rows = extract_docx(_chapter_one_source())
    issues = hard_chapter_one_supervisory_issues(
        rows, academic_level="Bachelors", submission_scope="chapter"
    )
    codes = {issue["checklist_code"] for issue in issues}
    expected = {
        "DOC-UNRESOLVED-SUPERVISOR-INSTRUCTION",
        "CIT-INCOMPLETE-PARENTHETICAL",
        "B1-CONSTRUCT-TERMINOLOGY-SHIFT",
        "B3-UNIT-OF-ANALYSIS-SINGULAR-PLURAL",
        "B2-CONTEXT-EVIDENCE-MISMATCH",
        "B2-UNSUPPORTED-LITERATURE-GAP",
        "B3-PURPOSE-OBJECTIVE-CONTENT-MISMATCH",
        "B3-CAUSAL-CLAIM-STRENGTH",
        "RQ-SUBJECT-VERB-AGREEMENT",
        "B4-LIMITATION-DELIMITATION-CONFUSION",
        "STYLE-BRITISH-AMERICAN",
    }
    assert expected <= codes
    combined = " ".join(str(value) for issue in issues for value in issue.values()).lower()
    assert "behaviour" not in combined
    assert "labor" not in combined
    assert "recognized" in combined or "organizations" in combined


def test_reference_list_missing_is_scope_aware():
    rows = extract_docx(_chapter_one_source(citations=6))
    chapter_codes = {
        issue["checklist_code"]
        for issue in hard_chapter_one_supervisory_issues(
            rows, academic_level="Bachelors", submission_scope="chapter"
        )
    }
    full_codes = {
        issue["checklist_code"]
        for issue in hard_chapter_one_supervisory_issues(
            rows, academic_level="Bachelors", submission_scope="full_thesis"
        )
    }
    assert "REF-MISSING-LIST" not in chapter_codes
    assert "REF-MISSING-LIST" in full_codes


def test_evidence_ledger_rejects_invented_examples_and_false_scope_claims():
    paragraphs = [{
        "paragraph": 1,
        "paragraph_id": "P1",
        "text": "The organization uses recognised procurement procedures.",
        "document_role": "current",
    }]
    findings = [
        {
            "finding_id": "F1",
            "status": "partly_meets_requirement",
            "severity": "minor",
            "issue_title": "British and American spelling are mixed",
            "assessment": "The paragraph uses examples such as behaviour, behavior and labor.",
            "required_action": "Apply British English consistently.",
            "evidence_paragraph_ids": ["P1"],
            "problematic_quote": "organization",
        },
        {
            "finding_id": "F2",
            "status": "does_not_meet_requirement",
            "severity": "major",
            "issue_title": "The reference list is missing despite visible in-text citations",
            "required_action": "Add a reference list.",
            "evidence_paragraph_ids": ["P1"],
        },
    ]
    rows = evidence_ledger_rows(findings, paragraphs, {"review_scope": "chapter"})
    assert [row["finding_id"] for row in rows] == ["F1"]
    assert "behaviour" not in rows[0].get("assessment", "").lower()
    assert "behavior" not in rows[0].get("assessment", "").lower()
    assert rows[0]["exact_source_text"] == paragraphs[0]["text"]
    assert rows[0]["verification_test"]


def test_same_exact_sentence_uses_one_numbered_native_comment(monkeypatch):
    monkeypatch.setenv("VPROF_GROUP_SAME_ANCHOR_COMMENTS", "true")
    monkeypatch.setenv("VPROF_HUMAN_ROOT_CAUSE_CONSOLIDATION", "false")
    document = Document()
    document.add_paragraph(
        "The problem is important, but evidence from the selected manufacturing firm is not provided."
    )
    source = _bytes(document)
    evidence = extract_docx(source)[0]
    review = {
        "summary": {"review_scope": "chapter", "reviewer_name": "Supervisor"},
        "_runtime_context": {"current_paragraphs": [{**evidence, "document_role": "current"}]},
        "academic_findings": [
            {
                "finding_id": "F1",
                "status": "does_not_meet_requirement",
                "severity": "major",
                "category": "research_gap_and_problem",
                "section": "Problem Statement",
                "issue_title": "The practical problem is asserted without evidence",
                "assessment": "The sentence states that the problem is important without showing its scale.",
                "academic_consequence": "The need for the study is not demonstrated.",
                "required_action": "Add verifiable evidence showing the scale and consequences of the problem.",
                "problematic_quote": evidence["text"],
                "evidence": [{**evidence, "document_role": "current"}],
                "evidence_paragraph_ids": [evidence.get("paragraph_id") or "P1"],
            },
            {
                "finding_id": "F2",
                "status": "partly_meets_requirement",
                "severity": "major",
                "category": "research_gap_and_problem",
                "section": "Problem Statement",
                "issue_title": "The declared case setting is unsupported",
                "assessment": "No evidence is supplied from the selected manufacturing firm.",
                "academic_consequence": "The contextual gap is not established.",
                "required_action": "Add firm-specific, sector or policy evidence and explain what remains unresolved.",
                "problematic_quote": evidence["text"],
                "evidence": [{**evidence, "document_role": "current"}],
                "evidence_paragraph_ids": [evidence.get("paragraph_id") or "P1"],
            },
        ],
    }
    output = Document(io.BytesIO(build_annotated_docx(source, review)))
    comments = list(output.comments)
    assert len(comments) == 1
    assert "1. " in comments[0].text
    assert "2. " in comments[0].text
    assert "Problem identified" not in comments[0].text
    assert "Action required" not in comments[0].text
    assert output.paragraphs[0].text.endswith("provided.")


def test_sentence_splitter_preserves_academic_numbers_and_dois():
    text = (
        "Funding was $11.4 billion and warming reached 1.2°C. "
        "The source has DOI 10.1108/JICES-09-2025-0264. The result was significant."
    )
    sentences = [text[start:end].strip() for start, end in sentence_spans(text)]
    assert len(sentences) == 3
    assert "$11.4" in sentences[0]
    assert "1.2°C" in sentences[0]
    assert "10.1108" in sentences[1]


def test_readiness_action_contains_exact_text_and_verification():
    review = {
        "summary": {"review_scope": "chapter", "document_label": "Chapter One"},
        "canonical_findings": [{
            "finding_number": 1,
            "status": "does_not_meet_requirement",
            "severity": "major",
            "section": "Problem Statement",
            "section_reference": "1.2 Statement of the Problem",
            "issue_title": "The contextual problem is not demonstrated",
            "exact_source_text": "Researches on these constructs have not received much attention.",
            "required_action": "Replace the general assertion with traceable evidence and a specific unresolved gap.",
            "academic_consequence": "The present wording does not justify the study.",
            "verification_test": "Confirm that the revised paragraph identifies evidence, context, consequence and the unresolved issue.",
            "evidence": [{"paragraph": 9, "page": 4, "section_reference": "1.2 Statement of the Problem"}],
        }],
    }
    readiness = build_supervisory_readiness(review)
    action = readiness["actions"][0]
    assert action["text_requiring_attention"].startswith("Researches on")
    assert "traceable evidence" in action["specific_action"]
    assert "evidence, context" in action["verification"]


def test_professional_review_acceptance_benchmark_reaches_ten_of_ten_gates():
    """Regression gate for the exact weaknesses found in the supplied Chapter One review.

    This is not a claim that every future disciplinary judgement will score 10/10.
    It prevents release when the known defects that caused the 4.8/10 output recur.
    """
    rows = extract_docx(_chapter_one_source(citations=6))
    issues = hard_chapter_one_supervisory_issues(
        rows, academic_level="Bachelors", submission_scope="chapter"
    )
    codes = {issue.get("checklist_code") for issue in issues}
    text = " ".join(str(value) for issue in issues for value in issue.values()).lower()
    gates = {
        "unresolved_supervisor_instruction": "DOC-UNRESOLVED-SUPERVISOR-INSTRUCTION" in codes,
        "broken_citation": "CIT-INCOMPLETE-PARENTHETICAL" in codes,
        "problem_context_evidence": "B2-CONTEXT-EVIDENCE-MISMATCH" in codes,
        "unsupported_gap": "B2-UNSUPPORTED-LITERATURE-GAP" in codes,
        "construct_consistency": "B1-CONSTRUCT-TERMINOLOGY-SHIFT" in codes,
        "unit_of_analysis_consistency": "B3-UNIT-OF-ANALYSIS-SINGULAR-PLURAL" in codes,
        "purpose_objective_alignment": "B3-PURPOSE-OBJECTIVE-CONTENT-MISMATCH" in codes,
        "claim_design_fit": "B3-CAUSAL-CLAIM-STRENGTH" in codes,
        "research_question_grammar": "RQ-SUBJECT-VERB-AGREEMENT" in codes,
        "limitation_delimitation": "B4-LIMITATION-DELIMITATION-CONFUSION" in codes,
    }
    assert sum(gates.values()) == 10, gates
    assert "REF-MISSING-LIST" not in codes
    assert "behaviour" not in text and "labor" not in text
