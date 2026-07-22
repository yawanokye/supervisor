from __future__ import annotations

import io

from docx import Document

from app.annotated_exporter import build_annotated_docx
from app.deterministic_supervisory_checklist import (
    _actual_spelling_variants,
    hard_chapter_one_supervisory_issues,
)
from app.document_parser import parse_document
from app.review_release_guard import filter_and_rewrite_release_findings


def _finding(title: str, section: str, paragraph: int = 3) -> dict:
    return {
        "finding_id": title,
        "chapter_number": 1,
        "section": section,
        "section_reference": section,
        "item": title,
        "issue_title": title,
        "assessment": title,
        "comment": title,
        "academic_consequence": "The weakness reduces the clarity of the study.",
        "required_action": "Revise the marked passage using the actual study evidence.",
        "verification_test": "Confirm that the revision is consistent with the title and objectives.",
        "severity": "major",
        "status": "does_not_meet_requirement",
        "confidence": 0.95,
        "evidence": [{
            "paragraph": paragraph,
            "paragraph_id": f"P{paragraph}",
            "text": "This is the marked passage.",
            "document_role": "current",
            "section_reference": section,
            "chapter_number": 1,
        }],
    }


def _runtime_review() -> dict:
    rows = [
        {"chapter_number": 1, "is_heading": True, "text": "CHAPTER ONE", "section_reference": "CHAPTER ONE", "paragraph": 1},
        {"chapter_number": 1, "is_heading": True, "text": "INTRODUCTION", "section_reference": "INTRODUCTION", "paragraph": 2},
        {
            "chapter_number": 1,
            "text": "This chapter introduces the study context and outlines the background, problem, objectives, questions, significance, scope, limitations and organisation of the study.",
            "section_reference": "INTRODUCTION",
            "paragraph": 3,
        },
        {"chapter_number": 1, "is_heading": True, "text": "Background to the Study", "section_reference": "Background to the Study", "paragraph": 4},
        {
            "chapter_number": 1,
            "text": "Information technology supports supply chains globally and in Ghana, but evidence from Aboabo Market remains limited.",
            "section_reference": "Background to the Study",
            "paragraph": 5,
        },
        {"chapter_number": 1, "is_heading": True, "text": "Objectives of the Study", "section_reference": "Objectives of the Study", "paragraph": 6},
        {
            "chapter_number": 1,
            "text": "1. To assess inventory systems. 2. To examine adoption challenges. 3. To determine supply-chain effectiveness.",
            "section_reference": "Objectives of the Study",
            "paragraph": 7,
        },
        {"chapter_number": 1, "is_heading": True, "text": "Significance of the Study", "section_reference": "Significance of the Study", "paragraph": 8},
        {
            "chapter_number": 1,
            "text": "The findings may help traders and technology providers, inform Ministry policy and support future researchers.",
            "section_reference": "Significance of the Study",
            "paragraph": 9,
        },
        {"chapter_number": 1, "is_heading": True, "text": "Limitations of the Study", "section_reference": "Limitations of the Study", "paragraph": 10},
        {
            "chapter_number": 1,
            "text": "Data access, limited time, financial constraints and respondent reluctance may affect the study.",
            "section_reference": "Limitations of the Study",
            "paragraph": 11,
        },
    ]
    return {
        "summary": {"submission_scope": "chapter", "academic_level": "Bachelors"},
        "_runtime_context": {"current_paragraphs": rows},
    }


def test_whole_section_guard_removes_known_false_missing_claims_and_rewrites_limitations():
    review = _runtime_review()
    findings = [
        _finding("Chapter title is too generic", "CHAPTER ONE"),
        _finding("Missing introduction text", "INTRODUCTION"),
        _finding("Missing third objective", "Problem Statement"),
        _finding("The applied or professional contribution is not explicit", "Significance of the Study"),
        _finding("The limitations and limits of generalisation need clearer explanation", "Limitations of the Study"),
    ]
    released = filter_and_rewrite_release_findings(findings, review)
    assert len(released) == 1
    assert released[0]["item"] == "The limitations are listed without explaining how they may affect the findings"
    assert "generalisation" not in released[0]["required_action"].lower()


def test_related_background_problem_and_significance_findings_are_consolidated():
    review = _runtime_review()
    findings = [
        _finding("Background lacks focus on study variables and context", "Background of the Study"),
        _finding("Gap not signalled", "Background of the Study"),
        _finding("Generic global discussion lacks local relevance", "Background of the Study"),
        _finding("Problem statement does not lead clearly to objectives", "Problem Statement"),
        _finding("Unclear problem-gap articulation", "Problem Statement"),
        _finding("Problem not supported by local evidence", "Problem Statement"),
        _finding("Vague and disconnected significance claims", "Significance of the Study"),
        _finding("The significance section does not distinguish the study's main contributions", "Significance of the Study"),
    ]
    released = filter_and_rewrite_release_findings(findings, review)
    titles = [row["item"] for row in released]
    assert sum("background" in title.lower() for title in titles) == 1
    assert sum("problem statement" in title.lower() for title in titles) == 1
    assert sum("significance" in title.lower() for title in titles) == 1
    assert len(released) == 3


def test_british_spelling_detector_does_not_classify_behavior_as_british():
    variants = _actual_spelling_variants("The organisation examined behavior and labour practices.")
    assert "behavior" in [value.lower() for value in variants["american"]]
    assert "behavior" not in [value.lower() for value in variants["british"]]
    assert "organisation" in [value.lower() for value in variants["british"]]


def _chapter_one_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("CONTRIBUTION OF INFORMATION TECHNOLOGY TO COMMODITY SUPPLY CHAIN MANAGEMENT IN THE ABOABO MARKET OF TAMALE METROPOLIS")
    doc.add_heading("CHAPTER ONE", level=1)
    doc.add_heading("INTRODUCTION", level=1)
    doc.add_paragraph("This chapter introduces the study and outlines its principal sections.")
    doc.add_heading("Background to the Study", level=2)
    doc.add_paragraph("Information technology supports supply chains in Ghana.")
    doc.add_heading("Statement of the Problem", level=2)
    doc.add_paragraph("The problem concerns IT use in informal commodity markets.")
    doc.add_heading("Purpose of the Study", level=2)
    doc.add_paragraph("The purpose of this study is to assess the impact of Information Technology on Supply Chain Management within Aboabo Market in Tamale.")
    doc.add_heading("Objectives of the Study", level=2)
    doc.add_paragraph("1. To assess IT use in Aboabo Market.\n2. To examine challenges in Tamale Market.\n3. To determine effectiveness in Tamale Market.")
    doc.add_heading("Research Questions", level=2)
    doc.add_paragraph("How can the use of IT creates a more effective commodity Supply Chain Management in Tamale Market?")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_deterministic_checks_detect_title_claim_setting_and_modal_verb_drift():
    rows = parse_document(_chapter_one_docx(), "chapter-one.docx")
    findings = hard_chapter_one_supervisory_issues(rows, academic_level="Bachelors", submission_scope="chapter")
    ids = {row["finding_id"] for row in findings}
    assert "DSC-HARD-B3-TITLE-PURPOSE-RELATIONSHIP-TERM" in ids
    assert "DSC-HARD-B3-STUDY-SETTING-DRIFT" in ids
    assert "DSC-HARD-RQ-MODAL-BASE-VERB" in ids


def test_native_comments_are_natural_and_every_number_is_represented(monkeypatch):
    monkeypatch.setenv("VPROF_NATIVE_COMMENT_STYLE", "exact_anchor_grouped")
    source = _chapter_one_docx()
    rows = parse_document(source, "chapter-one.docx")
    target = next(row for row in rows if "purpose of this study" in row.get("text", "").lower())
    findings = []
    specifications = [
        ("A source citation in the paragraph is incomplete", "citations_and_sources", "Complete the source citation and verify the reference entry."),
        ("The paragraph contains a verb agreement error", "academic_writing", "Correct the verb agreement without changing the intended meaning."),
        ("The study period is not stated in the paragraph", "scope_and_context", "State the study period and use it consistently in the methodology."),
    ]
    for index, (title, category, action) in enumerate(specifications, start=1):
        row = _finding(title, "Purpose of the Study", target["paragraph"])
        row["finding_id"] = f"NAT-{index}"
        row["category"] = category
        row["assessment"] = title + "."
        row["required_action"] = action
        row["evidence"] = [{**target, "document_role": "current"}]
        row["problematic_quote"] = target["text"]
        findings.append(row)
    review = {
        "summary": {"submission_scope": "chapter", "academic_level": "Bachelors"},
        "academic_findings": findings,
        "_runtime_context": {"current_paragraphs": rows},
    }
    output = build_annotated_docx(source, review)
    reviewed = Document(io.BytesIO(output))
    comments = list(reviewed.comments)
    assert len(comments) == 1
    text = comments[0].text
    for label in ("Issue:", "Problem identified:", "Action required:", "Verification:"):
        assert label not in text
    expected = len(review.get("canonical_findings") or [])
    assert expected >= 2
    assert all(f"{number}. " in text for number in range(1, expected + 1))
