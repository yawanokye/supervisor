from io import BytesIO

from docx import Document

from app.document_parser import extract_docx


def test_docx_tables_are_extracted_with_chapter_and_table_references():
    document = Document()
    document.add_paragraph("CHAPTER FOUR")
    document.add_paragraph("RESULTS AND DISCUSSION")
    document.add_paragraph("4.1 Introduction")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Statistic"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "R-squared"
    table.cell(1, 1).text = "0.62"
    buffer = BytesIO()
    document.save(buffer)

    rows = extract_docx(buffer.getvalue())
    table_rows = [row for row in rows if row.get("source_kind") == "table_row"]
    assert len(table_rows) == 2
    assert table_rows[1]["chapter_number"] == 4
    assert table_rows[1]["table_index"] == 1
    assert "R-squared" in table_rows[1]["text"]
