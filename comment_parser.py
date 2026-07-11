from __future__ import annotations

import io
import re
import zipfile
from typing import Any, Dict, Iterable, List, Optional
from xml.etree import ElementTree as ET

try:
    from docx import Document
except Exception:
    Document = None

from .document_parser import clean_text, normalised, parse_document

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WORD_NS}}}"
COMMENT_MARKERS = (
    "supervisor comment", "examiner comment", "reviewer comment", "action required",
    "required action", "comment:", "please revise", "please add", "please explain",
    "please clarify", "please justify", "please correct", "needs revision",
)
HEADING_ONLY = {
    "supervisor comments", "supervisor's comments", "comments", "review comments",
    "examiner comments", "corrections", "required corrections", "revision comments",
}
GREEN_VALUES = {"008000", "00B050", "009900", "008A00", "00A000", "006100"}


def _deduplicate(items: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    seen = set()
    for item in items:
        text = clean_text(item.get("text", ""))
        key = normalised(text)
        if len(key) < 6 or key in seen:
            continue
        seen.add(key)
        item = dict(item)
        item["text"] = text
        output.append(item)
    return output


def _strip_label(text: str) -> str:
    value = clean_text(text)
    value = re.sub(r"^\s*[\[(]?(?:rev|comment|c)\s*\d+[\])\.:-]*\s*", "", value, flags=re.I)
    value = re.sub(r"^\s*(?:supervisor|examiner|reviewer)\s+comment\s*[:.-]\s*", "", value, flags=re.I)
    value = re.sub(r"^\s*(?:comment|action required|required action)\s*[:.-]\s*", "", value, flags=re.I)
    return value.strip(" []\t")


def _split_comment_text(text: str) -> List[str]:
    value = clean_text(text)
    if not value:
        return []

    bracketed = [clean_text(x) for x in re.findall(r"\[([^\[\]]{8,})\]", value, flags=re.S)]
    if bracketed:
        return [_strip_label(x) for x in bracketed if _strip_label(x)]

    lines = [clean_text(line) for line in re.split(r"\r?\n+", value) if clean_text(line)]
    output: List[str] = []
    for line in lines:
        if normalised(line) in HEADING_ONLY:
            continue
        parts = re.split(r"(?=(?:^|\s)(?:\d+|[A-Z])?[.)]\s+)", line)
        parts = [clean_text(part) for part in parts if clean_text(part)] or [line]
        for part in parts:
            candidate = _strip_label(part)
            if len(candidate) < 8:
                continue
            if len(candidate) > 650:
                sentences = [clean_text(x) for x in re.split(r"(?<=[.!?])\s+", candidate) if clean_text(x)]
                if len(sentences) > 1:
                    output.extend(sentences)
                    continue
            output.append(candidate)
    return output


def _run_colour(run) -> str:
    try:
        rgb = run.font.color.rgb
        if rgb:
            return str(rgb).upper()
    except Exception:
        pass
    try:
        nodes = run._r.xpath("./w:rPr/w:color")
        if nodes:
            return (nodes[0].get(f"{W}val") or "").upper()
    except Exception:
        pass
    return ""


def _docx_native_comments(data: bytes, filename: str) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            if "word/comments.xml" not in archive.namelist():
                return []
            root = ET.fromstring(archive.read("word/comments.xml"))
            for node in root.findall(f".//{W}comment"):
                text = clean_text(" ".join((t.text or "") for t in node.findall(f".//{W}t")))
                if not text:
                    continue
                output.append({
                    "text": _strip_label(text),
                    "source_filename": filename,
                    "source_type": "native_word_comment",
                    "author": node.get(f"{W}author") or "",
                    "date": node.get(f"{W}date") or "",
                })
    except Exception:
        return []
    return output


def _iter_table_text(document) -> Iterable[str]:
    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                yield " | ".join(cells)


def _docx_body_comments(data: bytes, filename: str) -> List[Dict[str, Any]]:
    if Document is None:
        return []
    document = Document(io.BytesIO(data))
    marked: List[Dict[str, Any]] = []
    fallback_lines: List[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = clean_text(paragraph.text)
        if not paragraph_text:
            continue
        fallback_lines.append(paragraph_text)

        bracketed = re.findall(r"\[([^\[\]]{8,})\]", paragraph_text, flags=re.S)
        for value in bracketed:
            marked.append({
                "text": _strip_label(value),
                "source_filename": filename,
                "source_type": "inline_bracket_comment",
            })

        green_text = clean_text(" ".join(run.text for run in paragraph.runs if _run_colour(run) in GREEN_VALUES))
        if green_text:
            for value in _split_comment_text(green_text):
                marked.append({
                    "text": value,
                    "source_filename": filename,
                    "source_type": "coloured_inline_comment",
                })

        low = normalised(paragraph_text)
        if any(marker in low for marker in COMMENT_MARKERS):
            for value in _split_comment_text(paragraph_text):
                marked.append({
                    "text": value,
                    "source_filename": filename,
                    "source_type": "labelled_comment",
                })

    for text in _iter_table_text(document):
        fallback_lines.append(text)
        low = normalised(text)
        if any(marker in low for marker in COMMENT_MARKERS):
            for value in _split_comment_text(text):
                marked.append({
                    "text": value,
                    "source_filename": filename,
                    "source_type": "table_comment",
                })

    if marked:
        return _deduplicate(marked)

    # When the user explicitly uploads a supervisor-comment document, ordinary body
    # paragraphs are treated as comments if no labelled or coloured comments exist.
    fallback: List[Dict[str, Any]] = []
    for line in fallback_lines:
        if normalised(line) in HEADING_ONLY:
            continue
        for value in _split_comment_text(line):
            fallback.append({
                "text": value,
                "source_filename": filename,
                "source_type": "comment_document_text",
            })
    return _deduplicate(fallback)


def extract_comments_from_document(data: bytes, filename: str) -> List[Dict[str, Any]]:
    low = (filename or "").lower()
    if low.endswith(".docx"):
        comments = _docx_native_comments(data, filename) + _docx_body_comments(data, filename)
        return _deduplicate(comments)
    if low.endswith(".pdf"):
        paragraphs = parse_document(data, filename)
        text = "\n".join(clean_text(p.get("text", "")) for p in paragraphs if clean_text(p.get("text", "")))
        return _deduplicate([
            {
                "text": value,
                "source_filename": filename,
                "source_type": "pdf_comment_text",
            }
            for value in _split_comment_text(text)
        ])
    return []


def extract_supervisor_comments(
    documents: Optional[List[Dict[str, Any]]] = None,
    pasted_text: str = "",
    *,
    max_comments: int = 120,
) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    for document in documents or []:
        data = document.get("data") or b""
        filename = document.get("filename") or "supervisor-comments"
        if not data:
            continue
        output.extend(extract_comments_from_document(data, filename))

    for value in _split_comment_text(pasted_text):
        output.append({
            "text": value,
            "source_filename": "Pasted supervisor comments",
            "source_type": "pasted_text",
        })

    return _deduplicate(output)[:max_comments]
