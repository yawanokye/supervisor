from __future__ import annotations

import io
import os
import re
from collections import Counter, OrderedDict, defaultdict
from typing import Any, Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .comment_quality import public_text, sanitise_finding_rows, sentence_safe_trim
from .articleready_review_bridge import build_articleready_quality_audit
from .professional_review_pipeline import build_professional_review_package
from .reviewer_language import academic_level_label, professionalise_reviewer_language
from .supervisory_review_algorithm import build_supervisory_report_spec
from .submission_readiness import build_supervisory_readiness

INK = "1F2937"
MUTED = "667085"
BRAND = "2F5597"
SOFT = "EEF3F8"
PALE_GREEN = "EAF4EC"
PALE_AMBER = "FFF4E5"
LINE = "D7DEE8"


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _normalised(value: Any) -> str:
    return re.sub(r"[^a-z0-9 ]+", " ", _clean(value).lower()).strip()


def _env_bool(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _env_int(name: str, default: int, minimum: int = 0, maximum: int = 500) -> int:
    try:
        return max(minimum, min(maximum, int(os.getenv(name, str(default)))))
    except (TypeError, ValueError):
        return default


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: Any, bold: bool = False, colour: str = INK, size: float = 9.2) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(_clean(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(colour)


def _set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, colour in (("Title", 18, BRAND), ("Heading 1", 14, BRAND), ("Heading 2", 11.5, INK), ("Heading 3", 10.5, INK)):
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)


def _location_text(evidence: Iterable[Dict[str, Any]]) -> str:
    values: List[str] = []
    for item in evidence:
        parts: List[str] = []
        section = _clean(
            item.get("section_reference")
            or ((item.get("section_path") or [""])[-1] if item.get("section_path") else "")
            or item.get("heading")
        )
        if section:
            parts.append(section)
        if item.get("source_kind") == "table_row" or item.get("table_number"):
            number = _clean(item.get("table_number"))
            title = _clean(item.get("table_title"))
            table_label = f"Table {number}" if number else "Table"
            if title:
                table_label += f": {title}"
            if item.get("table_row") is not None:
                table_label += f", row {item['table_row']}"
            parts.append(table_label)
        if item.get("page") is not None:
            parts.append(f'page {item["page"]}')
        if item.get("paragraph") is not None and not item.get("table_number"):
            parts.append(f'paragraph {item["paragraph"]}')
        if parts:
            values.append(", ".join(parts))
    return "; ".join(dict.fromkeys(values)) or "section-level guidance"


def _unique(values: Iterable[Any], limit: int | None = None) -> List[str]:
    output: List[str] = []
    seen = set()
    for value in values:
        text = _clean(value)
        key = _normalised(text)
        if not text or not key or key in seen:
            continue
        seen.add(key)
        output.append(text)
        if limit is not None and len(output) >= limit:
            break
    return output


def _trim(value: str, limit: int) -> str:
    return sentence_safe_trim(public_text(value, reject_placeholders=True), limit)


def _severity_rank(value: str) -> int:
    return {"critical": 0, "major": 1, "moderate": 2, "minor": 3}.get(str(value).lower(), 9)


def _finding_anchor(row: Dict[str, Any]) -> Tuple[str, str]:
    section = _normalised(row.get("section_reference") or row.get("section") or "chapter-wide review")
    evidence = row.get("evidence") or []
    if evidence:
        best = evidence[0]
        if best.get("source_kind") == "table_row" or best.get("table_index"):
            anchor = f'table:{best.get("table_index", "")}:{best.get("table_row", "")}'
        else:
            anchor = f'{best.get("document_role", "current")}:{best.get("paragraph", "")}'
    else:
        anchor = f'missing:{_normalised(row.get("category", "other"))}'
    return section, anchor


def _group_findings(rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    groups: "OrderedDict[Tuple[str, str], Dict[str, Any]]" = OrderedDict()
    safe_rows = sanitise_finding_rows(rows)
    for row in sorted(safe_rows, key=lambda x: (_severity_rank(x.get("severity", "minor")), _normalised(x.get("item", "")))):
        key = _finding_anchor(row)
        group = groups.setdefault(key, {
            "section": _clean(row.get("section_reference") or row.get("section") or "Chapter-wide review"),
            "reference_label": _clean(row.get("reference_label") or row.get("section_reference") or row.get("section")),
            "severity": row.get("severity", "moderate"),
            "rows": [],
            "evidence": row.get("evidence") or [],
        })
        group["rows"].append(row)
        if _severity_rank(row.get("severity", "minor")) < _severity_rank(group["severity"]):
            group["severity"] = row.get("severity", "moderate")
        if not group["evidence"] and row.get("evidence"):
            group["evidence"] = row.get("evidence")
    return list(groups.values())


def _section_summaries(review: Dict[str, Any]) -> Dict[str, List[str]]:
    summaries: Dict[str, List[str]] = defaultdict(list)
    for row in review.get("academic_section_reviews") or []:
        heading = _clean(row.get("heading") or row.get("section_name"))
        assessment = _clean(row.get("section_assessment"))
        if not heading or not assessment:
            continue
        summaries[_normalised(heading)].append(assessment)
    return {key: _unique(values, 2) for key, values in summaries.items()}


def _matching_summary(section_name: str, summaries: Dict[str, List[str]]) -> str:
    target = _normalised(section_name)
    direct = summaries.get(target)
    if direct:
        return " ".join(direct)
    for key, values in summaries.items():
        if key in target or target in key:
            return " ".join(values)
    return ""


def _is_synthetic_review_section(name: str) -> bool:
    value = _normalised(name)
    return any(term in value for term in (
        "whole chapter coherence", "whole chapter consistency", "cross chapter coherence",
        "cross chapter alignment", "supervisor comment compliance audit",
    ))


def _ordered_section_reviews(review: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Combine split section parts without merging same-named sections across chapters."""
    groups: "OrderedDict[Tuple[int | None, str], Dict[str, Any]]" = OrderedDict()
    for row in review.get("academic_section_reviews") or []:
        heading = _clean(row.get("heading") or row.get("section_name") or "Untitled section")
        chapter_number = row.get("chapter_number")
        try:
            chapter_number = int(chapter_number) if chapter_number is not None else None
        except (TypeError, ValueError):
            chapter_number = None
        key = (chapter_number, _normalised(heading))
        if not key[1]:
            continue
        group = groups.setdefault(key, {
            "heading": heading,
            "chapter_number": chapter_number,
            "section_path": list(row.get("section_path") or []),
            "assessments": [],
            "scores": [],
            "coverage_warnings": [],
            "synthetic": _is_synthetic_review_section(heading),
        })
        assessment = _clean(row.get("section_assessment"))
        if assessment:
            group["assessments"].append(assessment)
        try:
            group["scores"].append(float(row.get("section_score")))
        except (TypeError, ValueError):
            pass
        warning = _clean(row.get("coverage_warning"))
        if warning:
            group["coverage_warnings"].append(warning)
    output = []
    for group in groups.values():
        scores = group.pop("scores")
        group["score"] = round(sum(scores) / len(scores), 1) if scores else None
        group["assessments"] = _unique(group["assessments"], 3)
        group["coverage_warnings"] = _unique(group["coverage_warnings"], 2)
        output.append(group)
    return output


def _row_chapter_number(row: Dict[str, Any]) -> int | None:
    value = row.get("chapter_number")
    try:
        if value is not None:
            return int(value)
    except (TypeError, ValueError):
        pass
    for evidence in row.get("evidence") or []:
        try:
            if evidence.get("chapter_number") is not None:
                return int(evidence.get("chapter_number"))
        except (TypeError, ValueError):
            continue
    return None


def _rows_for_section(
    section_name: str,
    rows: Sequence[Dict[str, Any]],
    chapter_number: int | None = None,
) -> List[Dict[str, Any]]:
    target = _normalised(section_name)
    matched = []
    for row in rows:
        source = _normalised(row.get("section_reference") or row.get("section", ""))
        if source != target:
            continue
        if chapter_number is not None and _row_chapter_number(row) not in {None, chapter_number}:
            continue
        matched.append(row)
    return matched


def _add_location(doc: Document, evidence: Sequence[Dict[str, Any]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Location: {_location_text(evidence)}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def _add_numbered_guidance(doc: Document, label: str, values: Sequence[str]) -> None:
    values = _unique(values)
    if not values:
        return
    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(3)
    lead.paragraph_format.space_after = Pt(2)
    lead.add_run(label).bold = True
    if len(values) == 1:
        lead.add_run(" " + values[0])
        return
    for index, value in enumerate(values, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"{index}. ").bold = True
        p.add_run(value)


def _add_review_point(doc: Document, group: Dict[str, Any], index: int) -> None:
    rows = group["rows"]
    titles = _unique((row.get("item") for row in rows), 3)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(2)
    reference = _clean(group.get("reference_label") or group.get("section"))
    heading.add_run(f"Review point {index}").bold = True
    if reference:
        heading.add_run(f" — {reference}").bold = True
    heading.add_run(": ").bold = True
    heading.add_run(_trim("; ".join(titles) if titles else "Academic revision required", 240)).bold = True

    _add_location(doc, group.get("evidence") or [])

    assessments = _unique((_trim(row.get("comment", ""), 520) for row in rows), 2)
    if assessments:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.add_run("Supervisor’s assessment: ").bold = True
        p.add_run(" ".join(assessments))

    actions = _unique((_trim(row.get("required_action", ""), 460) for row in rows), 4)
    _add_numbered_guidance(doc, "Required revision:", actions)

    examples = _unique((_trim(row.get("illustrative_guidance", ""), 360) for row in rows), 1)
    if examples:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run("Illustrative guidance: ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BRAND)
        example_run = p.add_run(examples[0])
        example_run.italic = True
        example_run.font.color.rgb = RGBColor.from_string(INK)


def _actionable(rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return [row for row in rows if row.get("status") not in {"meets_requirement", "not_applicable"}]


def _add_follow_up_section(doc: Document, title: str, rows: Sequence[Dict[str, Any]]) -> bool:
    rows = _actionable(rows)
    if not rows:
        return False
    doc.add_heading(title, level=1)
    for index, group in enumerate(_group_findings(rows), start=1):
        _add_review_point(doc, group, index)
    return True



def _context_summary(review: Dict[str, Any]) -> List[Tuple[str, str]]:
    context = review.get("study_context") or {}
    rows: List[Tuple[str, str]] = []
    title = _clean(context.get("title_or_opening_focus"))
    if title:
        rows.append(("Study focus recognised", _trim(title, 320)))
    for label, key in (
        ("Country", "confirmed_countries"),
        ("Study location", "confirmed_locations"),
        ("Sector or field", "confirmed_sectors"),
    ):
        values = _unique(context.get(key) or [], 6)
        if values:
            rows.append((label, ", ".join(values)))
    return rows


def _is_source_verification(row: Dict[str, Any]) -> bool:
    return bool(
        row.get("source_verification_required")
        or row.get("guidance_type") == "source_verification"
        or row.get("category") in {"citations_and_sources", "ethics_and_integrity"}
    )


def _add_source_verification_summary(doc: Document, rows: Sequence[Dict[str, Any]], section_number: int) -> bool:
    selected = [row for row in rows if _is_source_verification(row)]
    if not selected:
        return False
    doc.add_heading(f"{section_number}. Evidence and Source Verification", level=1)
    intro = doc.add_paragraph(
        "The following matters require verification against original, credible sources. They are not findings of misconduct."
    )
    intro.paragraph_format.space_after = Pt(5)
    seen = set()
    count = 0
    for row in selected:
        signature = (_normalised(row.get("section", "")), _normalised(row.get("required_action", ""))[:180])
        if signature in seen or not signature[1]:
            continue
        seen.add(signature)
        count += 1
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f'{_clean(row.get("section", "Chapter"))}: ').bold = True
        p.add_run(_trim(row.get("required_action", ""), 520))
        if row.get("evidence"):
            loc = p.add_run(f' ({_location_text(row.get("evidence") or [])})')
            loc.italic = True
            loc.font.color.rgb = RGBColor.from_string(MUTED)
        if count >= 10:
            break
    return True


CHAPTER_NAMES = {
    1: "Chapter One",
    2: "Chapter Two",
    3: "Chapter Three",
    4: "Chapter Four",
    5: "Chapter Five",
    6: "Chapter Six",
}


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _compact_sentence(value: Any, limit: int = 260) -> str:
    text = _clean(value)
    text = re.sub(r"\bExample\s*:\s*", "", text, flags=re.I)
    text = re.sub(r"\bIllustrative guidance\s*:\s*", "", text, flags=re.I)
    return _trim(text, limit)


def _compact_overall_assessment(review: Dict[str, Any]) -> str:
    summary = review.get("summary") or {}
    raw = _clean(
        review.get("overall_academic_assessment")
        or summary.get("readiness_meaning")
        or ""
    )
    if not raw:
        return "The work has been reviewed against the applicable programme and disciplinary standards."
    sentences = re.split(r"(?<=[.!?])\s+", raw)
    selected: List[str] = []
    total = 0
    for sentence in sentences:
        sentence = _clean(sentence)
        if not sentence:
            continue
        if total + len(sentence) > 650 and selected:
            break
        selected.append(sentence)
        total += len(sentence)
        if len(selected) >= 4:
            break
    return professionalise_reviewer_language(" ".join(selected), summary.get("academic_level"))


def _section_strengths(
    section_name: str,
    strengths: Sequence[Dict[str, Any]],
    assessments: Sequence[str],
    chapter_number: int | None = None,
) -> List[str]:
    target = _normalised(section_name)
    matched = []
    for strength in strengths:
        source = _normalised(strength.get("section", ""))
        strength_chapter = _row_chapter_number(strength)
        if source == target and (chapter_number is None or strength_chapter in {None, chapter_number}):
            observation = _compact_sentence(strength.get("observation", ""), 210)
            if observation:
                matched.append(observation)

    return _unique(matched, 2)


def _section_corrections(
    section_name: str,
    findings: Sequence[Dict[str, Any]],
    limit: int | None = None,
    chapter_number: int | None = None,
) -> List[str]:
    rows = _rows_for_section(section_name, sanitise_finding_rows(findings), chapter_number)
    rows = [
        row for row in rows
        if row.get("status") not in {"meets_requirement", "not_applicable"}
    ]
    rows.sort(
        key=lambda row: (
            _severity_rank(row.get("severity", "minor")),
            0 if _is_source_verification(row) else 1,
            _normalised(row.get("item", "")),
        )
    )

    output: List[str] = []
    seen = set()
    for row in rows:
        action = _compact_sentence(
            row.get("required_action")
            or row.get("comment")
            or row.get("item"),
            260,
        )
        signature = _normalised(action)
        if not action or not signature or signature in seen:
            continue
        seen.add(signature)
        output.append(action)
        if limit is not None and len(output) >= limit:
            break
    return output


def _chapter_number_from_rows(
    section_name: str,
    findings: Sequence[Dict[str, Any]],
    default_chapter: int | None,
) -> int | None:
    for row in _rows_for_section(section_name, findings, default_chapter):
        for evidence in row.get("evidence") or []:
            value = evidence.get("chapter_number")
            try:
                return int(value)
            except (TypeError, ValueError):
                continue

    match = re.search(
        r"\bchapter\s+(one|two|three|four|five|six|\d+)\b",
        section_name,
        flags=re.I,
    )
    if match:
        token = match.group(1).lower()
        words = {
            "one": 1, "two": 2, "three": 3,
            "four": 4, "five": 5, "six": 6,
        }
        return words.get(token, int(token) if token.isdigit() else None)
    return default_chapter


def _summary_units(review: Dict[str, Any]) -> List[Dict[str, Any]]:
    summary = review.get("summary") or {}
    findings = sanitise_finding_rows(review.get("academic_findings") or [])
    strengths = review.get("academic_strengths") or []
    section_rows = [
        row for row in _ordered_section_reviews(review)
        if not row.get("synthetic")
    ]

    detected_chapters = summary.get("current_chapters_detected") or []
    multi_chapter = len(detected_chapters) > 1 or str(
        summary.get("review_scope", "")
    ).lower() in {
        "complete_thesis", "complete dissertation", "full_thesis",
        "thesis", "complete",
    }

    default_chapter = summary.get("selected_chapter") or summary.get("inferred_chapter")
    try:
        default_chapter = int(default_chapter) if default_chapter else None
    except (TypeError, ValueError):
        default_chapter = None

    raw_units: List[Dict[str, Any]] = []
    for section in section_rows:
        heading = section["heading"]
        chapter_number = section.get("chapter_number")
        if chapter_number is None:
            chapter_number = _chapter_number_from_rows(
                heading, findings, default_chapter
            )
        raw_units.append({
            "label": heading,
            "chapter_number": chapter_number,
            "strengths": _section_strengths(
                heading,
                strengths,
                section.get("assessments") or [],
                chapter_number,
            ),
            "corrections": _section_corrections(heading, findings, None, chapter_number),
            "warnings": _unique(section.get("coverage_warnings") or [], 1),
            "assessments": section.get("assessments") or [],
        })

    if not multi_chapter:
        return raw_units

    grouped: "OrderedDict[int | None, Dict[str, Any]]" = OrderedDict()
    for unit in raw_units:
        chapter_number = unit["chapter_number"]
        group = grouped.setdefault(chapter_number, {
            "label": CHAPTER_NAMES.get(
                chapter_number,
                f"Chapter {chapter_number}" if chapter_number else "Whole Thesis",
            ),
            "chapter_number": chapter_number,
            "strengths": [],
            "corrections": [],
            "warnings": [],
            "assessments": [],
        })
        group["strengths"].extend(unit["strengths"])
        group["corrections"].extend(unit["corrections"])
        group["warnings"].extend(unit["warnings"])
        group["assessments"].extend(unit["assessments"])

    output = []
    for group in grouped.values():
        group["strengths"] = _unique(group["strengths"], 3)
        group["corrections"] = _unique(group["corrections"], None)
        group["warnings"] = _unique(group["warnings"], 1)
        output.append(group)
    return output


def _add_cell_list(
    cell,
    values: Sequence[str],
    *,
    numbered: bool = False,
    empty_text: str = "",
    colour: str = INK,
    size: float = 8.9,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cleaned = _unique(values)
    if not cleaned:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(empty_text)
        run.italic = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        return

    for index, value in enumerate(cleaned, start=1):
        p = cell.paragraphs[0] if index == 1 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        if numbered:
            lead = p.add_run(f"{index}. ")
            lead.bold = True
            lead.font.size = Pt(size)
            lead.font.color.rgb = RGBColor.from_string(BRAND)
        run = p.add_run(_compact_sentence(value, 280))
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(colour)




def _add_articleready_quality_audit_section(doc: Document, review: Dict[str, Any], number: int) -> bool:
    """Add a substantive ArticleReady-style audit section to the supervisor report."""
    audit = review.get("articleready_quality_audit") or build_articleready_quality_audit(review)
    rows = audit.get("audit_rows") or []
    routes = audit.get("detected_review_routes") or []
    if not rows and not routes:
        return False
    doc.add_heading(f"{number}. Method, Results and Discussion Quality Audit", level=1)
    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(5)
    lead.add_run("Review approach: ").bold = True
    lead.add_run(_clean(audit.get("principle_summary") or "The review is evidence-preserving and method-sensitive."))
    if routes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run("Detected review route(s): ").bold = True
        p.add_run(", ".join(_unique(routes)))
    if not rows:
        doc.add_paragraph(
            "No separate methods, results or discussion audit finding was available from the evidence-gated review output. "
            "The annotated document should still be checked for passage-level comments."
        )
        return True
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    _set_repeat_table_header(header)
    for cell, width in zip(header.cells, (1.45, 1.65, 2.7, 2.7)):
        _set_cell_width(cell, width)
        _set_cell_shading(cell, BRAND)
    _set_cell_text(header.cells[0], "Audit area", True, "FFFFFF", 8.7)
    _set_cell_text(header.cells[1], "Location", True, "FFFFFF", 8.7)
    _set_cell_text(header.cells[2], "Supervisor evaluation", True, "FFFFFF", 8.7)
    _set_cell_text(header.cells[3], "Required correction", True, "FFFFFF", 8.7)
    for idx, row_data in enumerate(rows[:18]):
        row = table.add_row()
        _prevent_row_split(row)
        for cell, width in zip(row.cells, (1.45, 1.65, 2.7, 2.7)):
            _set_cell_width(cell, width)
            if idx % 2:
                _set_cell_shading(cell, "F8FAFC")
        _set_cell_text(row.cells[0], row_data.get("area", "Review"), True, BRAND, 8.2)
        _set_cell_text(row.cells[1], row_data.get("location", "section-level evidence"), False, MUTED, 8.0)
        _set_cell_text(row.cells[2], _compact_sentence(row_data.get("finding", ""), 360), False, INK, 8.0)
        _set_cell_text(row.cells[3], _compact_sentence(row_data.get("required_action", ""), 360), False, INK, 8.0)
    return True


def _add_supervisory_readiness_section(
    doc: Document,
    review: Dict[str, Any],
    heading: str,
) -> bool:
    """Add the direct ArticleReady-style action schedule.

    This is the supervisor-facing decision section. It states exactly what must
    be corrected, where, why and how the supervisor can verify completion.
    """
    readiness = review.get("supervisory_readiness") or build_supervisory_readiness(review)
    actions = list(readiness.get("actions") or [])
    doc.add_heading(heading, level=1)
    status = doc.add_paragraph()
    status.add_run("Current status: ").bold = True
    status.add_run(_clean(readiness.get("status") or "Review completed"))
    meaning = _clean(readiness.get("meaning"))
    if meaning:
        doc.add_paragraph(meaning)

    if not actions:
        doc.add_paragraph(
            "No unresolved material correction was identified in the selected review scope. "
            "The supervisor should still confirm institutional formatting, source files and final submission requirements."
        )
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.autofit = False
        header = table.rows[0]
        _set_repeat_table_header(header)
        widths = (0.9, 1.75, 1.75, 2.15, 1.95, 1.75)
        for cell, width in zip(header.cells, widths):
            _set_cell_width(cell, width)
            _set_cell_shading(cell, BRAND)
        for cell, text in zip(
            header.cells,
            ("Priority", "Exact location and text", "Problem identified", "Action required", "Why it matters", "How to verify completion"),
        ):
            _set_cell_text(cell, text, True, "FFFFFF", 7.8)
        for index, action in enumerate(actions[:100]):
            row = table.add_row()
            _prevent_row_split(row)
            for cell, width in zip(row.cells, widths):
                _set_cell_width(cell, width)
                if index % 2:
                    _set_cell_shading(cell, "F8FAFC")
            exact = _clean(action.get("text_requiring_attention"))
            location_text = _clean(action.get("location"))
            if exact:
                location_text = location_text + "\n\nText: “" + _compact_sentence(exact, 240) + "”"
            _set_cell_text(row.cells[0], action.get("priority"), True, BRAND, 7.3)
            _set_cell_text(row.cells[1], location_text, False, MUTED, 7.3)
            _set_cell_text(row.cells[2], action.get("issue"), False, INK, 7.3)
            _set_cell_text(row.cells[3], action.get("specific_action"), False, INK, 7.3)
            _set_cell_text(row.cells[4], action.get("why_it_matters"), False, INK, 7.3)
            _set_cell_text(row.cells[5], action.get("verification"), False, INK, 7.3)

    statistical = readiness.get("statistical_assurance") or {}
    if statistical:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.add_run("Statistical accuracy and adequacy assurance: ").bold = True
        p.add_run(
            f"{_clean(statistical.get('accuracy_status'))}. "
            f"{_clean(statistical.get('adequacy_status'))}. "
            f"{_clean(statistical.get('limitation'))}"
        )
    analysis_actions = list(readiness.get("additional_analysis_actions") or [])
    if analysis_actions:
        doc.add_heading("Additional analyses or result-verification actions", level=2)
        doc.add_paragraph(
            "The following analyses or checks are recommendations until they are completed from the original data and output. "
            "They must not be described in the thesis as though they have already been performed."
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        header = table.rows[0]
        for cell, text in zip(
            header.cells,
            ("Priority and location", "Rationale", "Data required", "Suitable method and output", "Consequence of omission"),
        ):
            _set_cell_shading(cell, BRAND)
            _set_cell_text(cell, text, True, "FFFFFF", 7.8)
        for index, action in enumerate(analysis_actions[:40]):
            row = table.add_row()
            _prevent_row_split(row)
            if index % 2:
                for cell in row.cells:
                    _set_cell_shading(cell, "F8FAFC")
            _set_cell_text(row.cells[0], f"{action.get('priority')}\n{action.get('location')}", True, BRAND, 7.4)
            _set_cell_text(row.cells[1], action.get("rationale"), False, INK, 7.4)
            _set_cell_text(row.cells[2], action.get("data_required"), False, INK, 7.4)
            _set_cell_text(row.cells[3], f"{action.get('suitable_method')} Output: {action.get('output_to_report')}", False, INK, 7.4)
            _set_cell_text(row.cells[4], action.get("consequence_of_omission"), False, INK, 7.4)
    note = _clean(readiness.get("approval_note"))
    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    return True

def _add_compact_follow_up(
    doc: Document,
    title: str,
    rows: Sequence[Dict[str, Any]],
    limit: int = 4,
) -> bool:
    rows = _actionable(rows)
    if not rows:
        return False
    doc.add_heading(title, level=1)
    added = 0
    seen = set()
    for row in sorted(
        rows,
        key=lambda item: _severity_rank(item.get("severity", "moderate")),
    ):
        action = _compact_sentence(
            row.get("required_action") or row.get("comment"), 300
        )
        signature = _normalised(action)
        if not signature or signature in seen:
            continue
        seen.add(signature)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        section_name = _clean(row.get("section", "Review"))
        p.add_run(section_name + ": ").bold = True
        p.add_run(action)
        added += 1
        if added >= limit:
            break
    return added > 0

def _add_professional_finding_table(doc: Document, ledger: Sequence[Dict[str, Any]], title: str) -> None:
    doc.add_heading(title, level=1)
    if not ledger:
        doc.add_paragraph("No material evidence-anchored correction was identified at the applicable academic level.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    _set_repeat_table_header(header)
    for cell, width in zip(header.cells, (0.45, 1.7, 2.75, 3.25)):
        _set_cell_width(cell, width)
        _set_cell_shading(cell, BRAND)
    _set_cell_text(header.cells[0], "No.", True, "FFFFFF", 8.5)
    _set_cell_text(header.cells[1], "Location", True, "FFFFFF", 8.5)
    _set_cell_text(header.cells[2], "Professional assessment", True, "FFFFFF", 8.5)
    _set_cell_text(header.cells[3], "Specific correction required", True, "FFFFFF", 8.5)
    for idx, item in enumerate(ledger):
        row = table.add_row()
        _prevent_row_split(row)
        for cell, width in zip(row.cells, (0.45, 1.7, 2.75, 3.25)):
            _set_cell_width(cell, width)
            if idx % 2:
                _set_cell_shading(cell, "F8FAFC")
        _set_cell_text(row.cells[0], item.get("number", ""), True, BRAND, 8.2)
        location = item.get("location") or item.get("section") or "Chapter-level finding"
        _set_cell_text(row.cells[1], location, False, MUTED, 7.9)
        assessment = item.get("assessment") or item.get("issue") or "Academic correction required"
        issue = item.get("issue")
        if issue and _normalised(issue) not in _normalised(assessment):
            assessment = f"{issue}. {assessment}"
        verification = item.get("verification")
        if verification:
            assessment = f"{assessment} [{verification}]"
        _set_cell_text(row.cells[2], _compact_sentence(assessment, 520), False, INK, 8.0)
        correction = item.get("required_correction") or "Revise the cited passage in line with the professional assessment."
        example = item.get("example")
        if example:
            correction = f"{correction} For example, {re.sub(r'^(?:for\s+)?example[:,]?\s*', '', str(example), flags=re.I).strip()}"
        _set_cell_text(row.cells[3], _compact_sentence(correction, 620), False, INK, 8.0)


def _add_chapter_judgement_table(doc: Document, judgements: Sequence[Dict[str, Any]], title: str) -> None:
    doc.add_heading(title, level=1)
    if not judgements:
        doc.add_paragraph("No distinct chapter judgement was available from the submitted scope.")
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    _set_repeat_table_header(header)
    for cell, width in zip(header.cells, (1.05, 1.9, 0.65, 1.15, 3.05)):
        _set_cell_width(cell, width)
        _set_cell_shading(cell, BRAND)
    for cell, text in zip(header.cells, ("Chapter", "Specialist lens", "Score", "Issue profile", "Professional judgement")):
        _set_cell_text(cell, text, True, "FFFFFF", 8.4)
    for idx, item in enumerate(judgements):
        row = table.add_row()
        _prevent_row_split(row)
        for cell, width in zip(row.cells, (1.05, 1.9, 0.65, 1.15, 3.05)):
            _set_cell_width(cell, width)
            if idx % 2:
                _set_cell_shading(cell, "F8FAFC")
        counts = item.get("severity_counts") or {}
        profile = f"{counts.get('critical', 0)} critical; {counts.get('major', 0)} major; {counts.get('moderate', 0)} moderate"
        decision = item.get("decision", "")
        strengths = item.get("strengths") or []
        if strengths:
            decision += " Strength to retain: " + _compact_sentence(strengths[0], 180)
        _set_cell_text(row.cells[0], item.get("chapter", "Chapter"), True, BRAND, 8.2)
        _set_cell_text(row.cells[1], item.get("specialist_role", "Professional chapter reviewer"), False, INK, 8.0)
        _set_cell_text(row.cells[2], item.get("score") if item.get("score") is not None else "—", True, INK, 8.1)
        _set_cell_text(row.cells[3], profile, False, MUTED, 7.9)
        _set_cell_text(row.cells[4], decision, False, INK, 8.0)


def _add_methods_results_audit(doc: Document, audit: Dict[str, Any], number: int) -> int:
    doc.add_heading(f"{number}. Methods, Results and Discussion Accuracy Audit", level=1)
    doc.add_paragraph(_clean(audit.get("accuracy_statement")))
    groups = [
        ("Methods and reproducibility", audit.get("methods_findings") or []),
        ("Results and analytical accuracy", audit.get("results_accuracy_findings") or []),
        ("Discussion and interpretation", audit.get("discussion_findings") or []),
    ]
    for label, rows in groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(label).bold = True
        if not rows:
            p.add_run(": no separate material finding was identified from the evidence supplied.")
            continue
        for item in rows[:18]:
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Inches(0.24)
            q.paragraph_format.first_line_indent = Inches(-0.18)
            q.paragraph_format.space_after = Pt(2)
            q.add_run(f"{item.get('number')}. ").bold = True
            q.add_run(_compact_sentence(item.get("issue") or item.get("assessment"), 260))
            loc = q.add_run(f" ({_compact_sentence(item.get('location'), 130)})")
            loc.italic = True
            loc.font.color.rgb = RGBColor.from_string(MUTED)
    return number + 1


def _add_priority_plan(doc: Document, plan: Dict[str, Sequence[Dict[str, Any]]], number: int) -> int:
    doc.add_heading(f"{number}. Prioritised Correction Plan", level=1)
    labels = [
        ("Priority 1: validity and submission blockers", "priority_1_validity_and_submission_blockers"),
        ("Priority 2: major scholarly revision", "priority_2_major_scholarly_revision"),
        ("Priority 3: targeted and editorial revision", "priority_3_targeted_and_editorial_revision"),
    ]
    for label, key in labels:
        rows = list(plan.get(key) or [])
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(label).bold = True
        if not rows:
            p.add_run(": none identified.")
            continue
        for item in rows[:15]:
            q = doc.add_paragraph(style="List Bullet")
            q.paragraph_format.space_after = Pt(2)
            q.add_run(f"Correction {item.get('number')} — {item.get('section')}: ").bold = True
            q.add_run(_compact_sentence(item.get("required_correction"), 360))
    return number + 1



def _add_simple_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(_clean(text), level=level)


def _build_spec_aligned_docx_report(review: Dict[str, Any]) -> bytes:
    """Build a supervisor/examiner report aligned with the supplied review specification.

    The report is intentionally decision-led. It explains scope and limitations,
    identifies barriers to validity, presents the deterministic statistical audit,
    provides a chapter correction plan and then lists the detailed findings from the
    same canonical ledger used by the annotated DOCX.
    """
    doc = Document()
    _set_document_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    summary = review.get("summary") or {}
    package = review.get("professional_review") or build_professional_review_package(review)
    spec = package.get("supervisory_report_spec") or build_supervisory_report_spec(review, package)
    ledger = package.get("finding_ledger") or []
    profile = package.get("profile") or {}

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(_clean(spec.get("report_title") or "SUPERVISORY REVIEW REPORT"))

    study_label = _clean(summary.get("study_title") or summary.get("document_label") or summary.get("filename") or "Study reviewed")
    details = doc.add_table(rows=0, cols=2)
    details.style = "Table Grid"
    for label, value in (
        ("Study reviewed", study_label),
        ("Review mode", _clean(spec.get("review_mode") or profile.get("role") or "Professional academic supervisor")),
        ("Overall decision", _clean(spec.get("overall_decision") or "Review completed")),
        ("Context-specific comments", len(ledger)),
    ):
        cells = details.add_row().cells
        _set_cell_shading(cells[0], SOFT)
        _set_cell_text(cells[0], label, True, BRAND, 8.7)
        _set_cell_text(cells[1], value, False, INK, 8.7)

    _add_simple_heading(doc, "1. Scope and limitation of the review")
    doc.add_paragraph(_clean(spec.get("scope_and_limitation")))

    _add_simple_heading(doc, "2. Overall supervisory assessment")
    overall = _clean(spec.get("overall_assessment"))
    if not overall:
        overall = _clean((package.get("recommendation") or {}).get("meaning"))
    doc.add_paragraph(overall or "The study was reviewed against the academic standard, internal alignment and evidence reported in the work.")

    _add_simple_heading(doc, "2.1 Main strengths", level=2)
    strengths = list(spec.get("strengths") or [])
    if strengths:
        for item in strengths:
            doc.add_paragraph(_clean(item), style="List Bullet")
    else:
        doc.add_paragraph("Retain the parts of the study that are accurate, well supported and clearly aligned while completing the corrections below.")

    _add_simple_heading(doc, "2.2 Critical corrections required before the study can proceed", level=2)
    critical = list(spec.get("critical_corrections") or [])
    if critical:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ("Area", "Required correction")):
            _set_cell_shading(cell, BRAND)
            _set_cell_text(cell, text, True, "FFFFFF", 8.4)
        for item in critical:
            cells = table.add_row().cells
            _set_cell_text(cells[0], item.get("area"), True, BRAND, 8.2)
            refs = ", ".join(str(v) for v in item.get("finding_numbers") or [])
            action = _clean(item.get("required_correction"))
            if refs:
                action += f" [Corrections {refs}]"
            _set_cell_text(cells[1], action, False, INK, 8.2)
    else:
        doc.add_paragraph("No critical or major correction was identified. Address the remaining targeted corrections before resubmission.")

    _add_supervisory_readiness_section(
        doc,
        review,
        "2.3 Actions Required Before Supervisor Approval or Submission",
    )

    _add_simple_heading(doc, "3. Methods, Results and Discussion Accuracy Audit")
    p = doc.add_paragraph()
    p.add_run("Statistical consistency and analysis appropriateness: ").bold = True
    p.add_run("Each reported model is checked for internal numerical consistency, suitability for the research task and agreement between the table and interpretation.")
    stat_rows = list(spec.get("statistical_audit") or [])
    if stat_rows:
        doc.add_paragraph(
            "These checks evaluate the internal consistency and adequacy of the statistics printed in the study. They do not replace re-analysis from the raw data and original software output."
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ("Check", "Finding", "Status", "Action")):
            _set_cell_shading(cell, BRAND)
            _set_cell_text(cell, text, True, "FFFFFF", 8.1)
        for item in stat_rows:
            cells = table.add_row().cells
            _set_cell_text(cells[0], item.get("check"), True, BRAND, 7.9)
            _set_cell_text(cells[1], item.get("finding"), False, INK, 7.9)
            _set_cell_text(cells[2], item.get("status"), True, INK, 7.9)
            _set_cell_text(cells[3], item.get("action"), False, INK, 7.9)
    else:
        doc.add_paragraph(
            "No deterministic statistical inconsistency was identified in the material reviewed. This does not certify the underlying calculations where raw data or original software output was not supplied."
        )

    _add_simple_heading(doc, "4. Chapter-by-chapter correction plan")
    plans = list(spec.get("chapter_plans") or [])
    if plans:
        for plan in plans:
            p = doc.add_paragraph()
            p.add_run(f"{_clean(plan.get('chapter'))}: ").bold = True
            corrections = [_clean(value) for value in plan.get("corrections") or [] if _clean(value)]
            p.add_run("; ".join(corrections) if corrections else "No material correction identified.")
    else:
        doc.add_paragraph("No chapter-specific correction plan was generated.")

    coverage = spec.get("coverage") or {}
    section_coverage = spec.get("section_coverage") or {}
    if coverage:
        _add_simple_heading(doc, "5. Review coverage assurance")
        counts = coverage.get("status_counts") or {}
        doc.add_paragraph(
            f"The review assessed {coverage.get('assessed_target_count', 0)} of {coverage.get('target_count', 0)} paragraphs and table rows across "
            f"{coverage.get('completed_units', 0)} of {coverage.get('unit_count', 0)} review units. "
            f"Coverage statuses: PASS {counts.get('PASS', 0)}, COMMENT {counts.get('COMMENT', 0)}, VERIFY SOURCE {counts.get('VERIFY SOURCE', 0)}, RE-ANALYSE {counts.get('RE-ANALYSE', 0)}. "
            "The number of comments was determined by the issues found, not by a predetermined quota."
        )
        next_number = 6
    else:
        next_number = 5

    if section_coverage:
        _add_simple_heading(doc, f"{next_number}. Required-section coverage")
        entries = list(section_coverage.get("entries") or [])
        material = [
            item for item in entries
            if _clean(item.get("status")) in {"MISSING", "PRESENT_BUT_INADEQUATE"}
        ]
        counts = section_coverage.get("counts") or {}
        doc.add_paragraph(
            f"The structural review assessed {section_coverage.get('applicable_section_count', 0)} applicable sections across the submitted chapters. "
            f"Present or equivalent: {int(counts.get('PRESENT', 0)) + int(counts.get('EQUIVALENT_HEADING', 0))}; "
            f"present but inadequate: {counts.get('PRESENT_BUT_INADEQUATE', 0)}; missing: {counts.get('MISSING', 0)}."
        )
        if material:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, ("Chapter", "Required section", "Status", "Correction")):
                _set_cell_shading(cell, BRAND)
                _set_cell_text(cell, text, True, "FFFFFF", 8.1)
            for item in material:
                cells = table.add_row().cells
                location_label = "Whole study" if int(item.get("chapter_number") or 0) == 0 else f"Chapter {item.get('chapter_number')}"
                _set_cell_text(cells[0], location_label, True, BRAND, 7.9)
                _set_cell_text(cells[1], item.get("label"), False, INK, 7.9)
                _set_cell_text(cells[2], item.get("status"), True, INK, 7.9)
                correction = item.get("required_action") or f"Add a clearly labelled {item.get('label')} section in the appropriate position."
                _set_cell_text(cells[3], correction, False, INK, 7.9)
        next_number += 1

    _add_simple_heading(doc, f"{next_number}. Numbered comments and detailed corrections")
    include_details = _env_bool("VPROF_REPORT_INCLUDE_DETAILED_FINDINGS", False)
    max_details = _env_int("VPROF_REPORT_MAX_DETAILED_FINDINGS", 30, 1, 200)
    bundle_ok = bool(summary.get("annotation_bundle_validation_passed"))
    if include_details:
        prioritised = sorted(ledger, key=lambda item: ({"critical": 0, "major": 1, "moderate": 2, "minor": 3}.get(str(item.get("severity") or "minor").lower(), 9), int(item.get("number") or 0)))
        _add_professional_finding_table(doc, prioritised[:max_details], "Selected detailed findings")
        if len(ledger) > max_details:
            if bundle_ok:
                doc.add_paragraph(
                    f"The reviewed thesis contains {len(ledger)} sequentially numbered findings. This report presents the {max_details} highest-priority findings; the complete guidance is represented in the validated native and inline annotated documents."
                )
            else:
                doc.add_paragraph(
                    f"The reviewed thesis contains {len(ledger)} sequentially numbered findings. This report presents the {max_details} highest-priority findings. The annotated delivery documents were not validated and must not be described as attached."
                )
    else:
        native_count = int(summary.get("native_docx_comment_count") or 0)
        inline_count = int(summary.get("inline_annotation_count") or 0)
        if bundle_ok:
            doc.add_paragraph(
                f"The review contains {len(ledger)} sequentially numbered findings represented across {native_count} native Word comment box{'es' if native_count != 1 else ''} and {inline_count} inline supervisor note{'s' if inline_count != 1 else ''}. "
                "The report summarises the decision, validity barriers, statistical audit and chapter correction plan without repeating every annotation word for word."
            )
        else:
            doc.add_paragraph(
                f"The review contains {len(ledger)} sequentially numbered comments in the canonical finding ledger. This report deliberately summarises the decision, validity barriers, statistical audit and chapter correction plan without repeating every comment word for word. "
                "The annotated delivery documents were not validated and must not be described as attached. Recover the export stage or submit a fresh review before treating the job as a complete supervisory package."
            )
    next_number += 1

    _add_simple_heading(doc, f"{next_number}. Evidence Required for Verification")
    evidence_required = list(spec.get("evidence_required") or [])
    if evidence_required:
        for item in evidence_required:
            p = doc.add_paragraph(style="List Bullet")
            if item.get("number"):
                p.add_run(f"Correction {item.get('number')}: ").bold = True
            p.add_run(_clean(item.get("evidence_needed") or item.get("required_correction") or item))
    else:
        doc.add_paragraph(
            "Retain the raw data, instrument or coding framework, sampling records, ethical approval and original analytical output so that any result can be verified during supervision or examination."
        )
    next_number += 1

    _add_simple_heading(doc, f"{next_number}. Order of revision")
    for text in (
        "Correct all measurement, scoring, model-specification and statistical inconsistencies first.",
        "Resolve major cross-chapter alignment, methodological and theoretical issues.",
        "Rebuild the results, discussion, conclusions and recommendations from verified analysis.",
        "Complete reference reconciliation, language editing, formatting and final document checks.",
    ):
        doc.add_paragraph(text, style="List Number")
    next_number += 1

    _add_simple_heading(doc, f"{next_number}. Professional recommendation")
    recommendation = package.get("recommendation") or {}
    doc.add_paragraph(f"{_clean(spec.get('overall_decision') or recommendation.get('decision'))}. {_clean(recommendation.get('meaning'))}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"VProfessor | {_clean(profile.get('role') or 'Professional academic reviewer')}")
    footer_run.italic = True
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_docx_report(review: Dict[str, Any]) -> bytes:
    """Create the canonical professional supervisor or examiner report.

    The report, native comments and inline annotated document are rendered from
    the same finding ledger. The report therefore preserves substantive academic
    judgement instead of reducing the review to a short comment summary.
    """
    if str(os.getenv("VPROF_SPEC_ALIGNED_SUPERVISORY_REPORT", "true")).strip().lower() in {"1", "true", "yes", "on"}:
        return _build_spec_aligned_docx_report(review)

    doc = Document()
    _set_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    summary = review.get("summary") or {}
    package = review.get("professional_review") or build_professional_review_package(review)
    profile = package.get("profile") or {}
    ledger = package.get("finding_ledger") or []

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(_clean(profile.get("report_title") or summary.get("professional_report_title") or "PROFESSIONAL ACADEMIC REVIEW"))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(_clean(summary.get("filename", "Reviewed work")))
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    details = doc.add_table(rows=0, cols=4)
    details.style = "Table Grid"
    details.autofit = False
    pairs = [
        (("Review role", profile.get("role", "Professional academic reviewer")), ("Review scope", str(summary.get("review_scope", "chapter")).replace("_", " ").title())),
        (("Document", summary.get("document_label", "")), ("Academic level", summary.get("academic_level", ""))),
        (("Review depth", str(summary.get("review_depth", "standard")).title()), ("Canonical findings", len(ledger))),
        (("Coverage assessed", f"{summary.get('coverage_percent', 0)}%" if summary.get("systematic_coverage_review") else "Legacy section review"), ("Review units", f"{summary.get('coverage_units_completed', 0)}/{summary.get('coverage_units_total', 0)}" if summary.get("systematic_coverage_review") else summary.get("academic_review_units_completed", ""))),
        (("Submission stage", "Revised submission" if summary.get("revised_mode") else "Initial submission"), ("Overall judgement", package.get("recommendation", {}).get("decision") or summary.get("readiness_label", ""))),
    ]
    for left, right in pairs:
        cells = details.add_row().cells
        for cell, width in zip(cells, (1.25, 2.35, 1.35, 2.35)):
            _set_cell_width(cell, width)
        _set_cell_shading(cells[0], SOFT)
        _set_cell_text(cells[0], left[0], True, BRAND, 8.5)
        _set_cell_text(cells[1], left[1], False, INK, 8.5)
        _set_cell_shading(cells[2], SOFT)
        _set_cell_text(cells[2], right[0], True, BRAND, 8.5)
        _set_cell_text(cells[3], right[1], False, INK, 8.5)

    doc.add_heading("1. Overall Professional Judgement", level=1)
    doc.add_paragraph(_compact_overall_assessment(review))
    p = doc.add_paragraph()
    p.add_run("Scope-specific judgement: ").bold = True
    p.add_run(_clean(profile.get("primary_task")))
    judgement = doc.add_table(rows=1, cols=1)
    judgement.style = "Table Grid"
    cell = judgement.cell(0, 0)
    counts = Counter(item.get("severity") for item in ledger)
    fill = PALE_AMBER if counts.get("critical", 0) or counts.get("major", 0) else PALE_GREEN
    _set_cell_shading(cell, fill)
    recommendation = package.get("recommendation") or {}
    _set_cell_text(cell, f"{recommendation.get('decision', 'Review completed')}. {recommendation.get('meaning', '')}", True, INK, 9.2)

    _add_supervisory_readiness_section(
        doc,
        review,
        "2. Actions Required Before Supervisor Approval or Submission",
    )

    section_number = 3
    if summary.get("systematic_coverage_review"):
        doc.add_heading(f"{section_number}. Review Coverage Assurance", level=1)
        section_number += 1
        coverage = review.get("coverage_ledger") or {}
        complete = bool(coverage.get("complete"))
        status = "Complete" if complete else "Incomplete"
        doc.add_paragraph(
            f"Coverage status: {status}. The reviewer assessed {coverage.get('assessed_target_count', 0)} of "
            f"{coverage.get('target_count', 0)} substantive paragraphs and table rows across "
            f"{coverage.get('completed_units', 0)} of {coverage.get('unit_count', 0)} sequential review units. "
            "The number of comments was determined by the material issues found, not by a predetermined quota."
        )
        chapters = coverage.get("chapters") or {}
        if chapters:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            header = table.rows[0]
            for cell, text in zip(header.cells, ("Chapter", "Units", "Targets", "Coverage")):
                _set_cell_shading(cell, BRAND)
                _set_cell_text(cell, text, True, "FFFFFF", 8.3)
            for chapter, stats in sorted(chapters.items(), key=lambda item: str(item[0])):
                row = table.add_row().cells
                chapter_label = CHAPTER_NAMES.get(int(chapter), f"Chapter {chapter}") if str(chapter).isdigit() else "Whole thesis"
                target_total = int(stats.get("targets") or 0)
                assessed = int(stats.get("assessed_targets") or 0)
                percent = round(100.0 * assessed / max(1, target_total), 1)
                _set_cell_text(row[0], chapter_label, True, BRAND, 8.1)
                _set_cell_text(row[1], f"{stats.get('completed_units', 0)}/{stats.get('units', 0)}", False, INK, 8.1)
                _set_cell_text(row[2], f"{assessed}/{target_total}", False, INK, 8.1)
                _set_cell_text(row[3], f"{percent}%", False, INK, 8.1)

    blockers = list((package.get("priority_correction_plan") or {}).get("priority_1_validity_and_submission_blockers") or [])
    doc.add_heading(f"{section_number}. Critical Submission Blockers", level=1)
    section_number += 1
    if blockers:
        for item in blockers:
            q = doc.add_paragraph(style="List Number")
            q.add_run(f"Correction {item.get('number')} — {item.get('section')}: ").bold = True
            q.add_run(_compact_sentence(item.get("required_correction"), 420))
    else:
        doc.add_paragraph(f"No critical blocker was identified. Major and moderate corrections may still prevent approval or progression.")

    _add_chapter_judgement_table(doc, package.get("chapter_judgements") or [], f"{section_number}. Chapter or Section Judgements")
    section_number += 1

    doc.add_heading(f"{section_number}. Numbered Comments and Detailed Corrections", level=1)
    if _env_bool("VPROF_REPORT_INCLUDE_DETAILED_FINDINGS", False):
        max_details = _env_int("VPROF_REPORT_MAX_DETAILED_FINDINGS", 30, 1, 200)
        prioritised = sorted(ledger, key=lambda item: ({"critical": 0, "major": 1, "moderate": 2, "minor": 3}.get(str(item.get("severity") or "minor").lower(), 9), int(item.get("number") or 0)))
        _add_professional_finding_table(doc, prioritised[:max_details], "Selected detailed findings")
    else:
        doc.add_paragraph(f"The reviewed work contains {len(ledger)} sequentially numbered, passage-specific comments. The detailed guidance remains beside the relevant text in the reviewed thesis, while this report concentrates on the overall judgement and priorities.")
    section_number += 1

    section_number = _add_methods_results_audit(doc, package.get("methods_results_discussion_audit") or {}, section_number)

    alignment_rows = package.get("cross_chapter_alignment") or []
    if alignment_rows:
        doc.add_heading(f"{section_number}. Cross-Chapter Alignment", level=1)
        section_number += 1
        for row in alignment_rows[:20]:
            q = doc.add_paragraph(style="List Bullet")
            if row.get("number"):
                q.add_run(f"Correction {row.get('number')} — ").bold = True
            q.add_run(_compact_sentence(row.get("finding") or row.get("required_correction"), 360))
            if row.get("required_correction") and _normalised(row.get("required_correction")) not in _normalised(row.get("finding")):
                q.add_run(" Required action: ").bold = True
                q.add_run(_compact_sentence(row.get("required_correction"), 280))

    doc.add_heading(f"{section_number}. Strengths to Retain", level=1)
    section_number += 1
    strengths = review.get("academic_strengths") or []
    if strengths:
        for value in _unique((f"{_clean(item.get('section', 'Chapter'))}: {_compact_sentence(item.get('observation', ''), 320)}" for item in strengths), 12):
            doc.add_paragraph(value, style="List Bullet")
    else:
        doc.add_paragraph("The revision should preserve any accurate evidence, coherent section structure and well-supported arguments while addressing the corrections above.")

    section_number = _add_priority_plan(doc, package.get("priority_correction_plan") or {}, section_number)

    evidence_required = (package.get("methods_results_discussion_audit") or {}).get("evidence_required") or []
    doc.add_heading(f"{section_number}. Evidence Required for Verification", level=1)
    section_number += 1
    if evidence_required:
        for item in evidence_required:
            q = doc.add_paragraph(style="List Bullet")
            q.add_run(f"Correction {item.get('number')}: ").bold = True
            q.add_run(_compact_sentence(item.get("evidence_needed"), 360))
            loc = q.add_run(f" ({_compact_sentence(item.get('location'), 150)})")
            loc.italic = True
            loc.font.color.rgb = RGBColor.from_string(MUTED)
    else:
        doc.add_paragraph("No separate original-output request was generated. This does not remove the candidate’s responsibility to retain data, instruments, coding records and analytical output for examination.")

    if _add_compact_follow_up(doc, f"{section_number}. Response to Earlier Supervisor Comments", review.get("revision_results") or [], 12):
        section_number += 1

    doc.add_heading(f"{section_number}. Professional Recommendation", level=1)
    doc.add_paragraph(f"{recommendation.get('decision', summary.get('readiness_label', 'Review completed'))}. {recommendation.get('meaning', summary.get('readiness_meaning', ''))}")
    closing = doc.add_paragraph()
    closing.add_run("Next action: ").bold = True
    closing.add_run("Resolve Priority 1 matters first, verify all analytical findings against original evidence, complete Priority 2 scholarly revisions, and then address Priority 3 presentation issues before resubmission.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"VProfessor | {profile.get('role', 'Professional academic reviewer')}")
    footer_run.italic = True
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor.from_string(MUTED)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
