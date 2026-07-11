from __future__ import annotations

import io
import os
import re
from collections import defaultdict
from copy import deepcopy
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table

from .document_parser import clean_text, normalised
from .comment_quality import (
    comment_max_chars,
    public_text,
    sanitise_finding_row,
    sanitise_finding_rows,
    sentence_safe_trim,
)
from .review_rules import STATUS_MANUAL, STATUS_MISSING, STATUS_PARTIAL
from .review_enrichment import context_specific_example
from .finding_order import order_and_number_rows
from .reviewer_language import academic_level_label, professionalise_reviewer_language

ANNOTATION_EXPORT_VERSION = "1.9.9.21-expert-sequential-detailed-review"
ACTIONABLE_STATUSES = {STATUS_PARTIAL, STATUS_MISSING, STATUS_MANUAL}
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
COMMENT_RED = RGBColor(0xC0, 0x00, 0x00)
INLINE_BLUE = RGBColor(0x00, 0x70, 0xC0)
_RICH_RED_RE = re.compile(r"\[\[VPROF_RED:(.*?)\]\]")
_REFNO_RE = re.compile(r"\[\[VPROF_REFNO:(\d+)\]\]")


# Unresolved drafting placeholders inside the student's own document must always
# receive a native comment, even if the AI reviewer or the independent audit
# misses them. This is deliberately handled at export time because the exporter
# has access to the final Word paragraph/run map and can anchor the comment
# exactly on the placeholder text.
_BODY_PLACEHOLDER_RE = re.compile(
    r"\[(?:\s*(?:insert|add|specify|provide|complete|fill\s+in|replace|enter|start|end|month|year|date|x\b)[^\]\r\n]{0,120})\]",
    flags=re.I,
)

_NATURAL_LABEL_RE = re.compile(
    r"\b(?:Issue|Why this matters|Revise by|Guidance|Academic implication|Academic consequence)\s*:\s*",
    flags=re.I,
)


def _strip_visible_labels(value: str) -> str:
    return re.sub(r"\s{2,}", " ", _NATURAL_LABEL_RE.sub("", clean_text(value))).strip()




def _env_bool(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "on"}


def _missing_section_inline_bottom_enabled() -> bool:
    return _env_bool("VPROF_MISSING_SECTION_INLINE_BOTTOM", True)


def _native_group_location_markers_enabled() -> bool:
    return _env_bool("VPROF_NATIVE_GROUP_LOCATION_MARKERS", True)


def _sequential_reference_numbers_enabled() -> bool:
    return _env_bool("VPROF_SEQUENTIAL_COMMENT_REFERENCES", True)


def _specific_corrections_required_enabled() -> bool:
    return _env_bool("VPROF_SPECIFIC_CORRECTIONS_REQUIRED_BOTTOM", True)


def _missing_section_haystack(row: Dict[str, Any]) -> str:
    return normalised(" ".join(
        clean_text(str(row.get(field, "")))
        for field in (
            "item", "issue_title", "comment", "assessment", "required_action",
            "section", "section_reference", "reference_label", "problematic_quote",
        )
    ))


def _is_missing_section_finding(row: Dict[str, Any]) -> bool:
    """Return True for findings about an absent chapter section.

    These findings should not be attached to an unrelated existing section. A
    missing section has no exact anchor in the document, so it is handled as a
    blue inline note at the bottom of the reviewed chapter.
    """
    text = _missing_section_haystack(row)
    if not text:
        return False
    missing_tokens = (
        "missing", "not evident", "not present", "absent", "add or clearly label",
        "expected ucc thesis section", "no references", "reference list is missing",
        "bibliography section", "definition of terms",
    )
    section_tokens = (
        "section", "definition of terms", "operational definition", "references",
        "reference list", "bibliography", "glossary",
    )
    if "too thin" in text or "underdeveloped" in text or "not explicit" in text:
        return False
    return any(token in text for token in missing_tokens) and any(token in text for token in section_tokens)


def _missing_section_name(row: Dict[str, Any]) -> str:
    text = _missing_section_haystack(row)
    for label, needles in (
        ("Definition of Terms", ("definition of terms", "operational definition", "glossary")),
        ("References", ("reference list", "references", "bibliography")),
        ("Research Gap", ("research gap section", "separate research gap")),
        ("Delimitation of the Study", ("delimitation", "scope/delimitation")),
    ):
        if any(needle in text for needle in needles):
            return label
    raw = clean_text(row.get("section_reference") or row.get("section") or row.get("reference_label") or "Required section")
    raw = re.sub(r"^expected\s+ucc\s+thesis\s+section\s+is\s+not\s+evident\s*[:\-]?\s*", "", raw, flags=re.I)
    return raw or "Required section"


def _missing_section_bottom_comment(row: Dict[str, Any]) -> str:
    section_name = _missing_section_name(row)
    comment = _sanitise_guidance(row.get("comment", ""))
    action = _sanitise_guidance(row.get("required_action", ""))
    example = _sanitise_guidance(row.get("illustrative_guidance", "")) or _sanitise_guidance(context_specific_example(row))
    if not comment:
        comment = f"The chapter does not make the {section_name} section evident, so the reader cannot locate required material in the expected chapter structure."
    if not action:
        if normalised(section_name) == "definition of terms":
            action = "Add a clearly labelled Definition of Terms section near the end of Chapter One and define the core study constructs in measurable terms."
        elif normalised(section_name) == "references":
            action = "Add a complete References section after the chapter text and match every in-text citation to a full reference-list entry."
        else:
            action = f"Add a clearly labelled {section_name} section at the appropriate point in the chapter and ensure the content matches the programme format."
    if not example and normalised(section_name) == "definition of terms":
        example = "For example, define each central construct using the dimensions and indicators applied in the instrument, coding framework or analysis plan."
    elif not example and normalised(section_name) == "references":
        example = "For example, remove duplicate citation clusters and provide complete author, year, title, source and DOI or URL details for every source retained."
    parts = [f"Missing section: {section_name}.", comment.rstrip(" .") + ".", _normalise_action_start(action).rstrip(" .") + "."]
    if example:
        example = re.sub(r"^for example[:,]?\s*", "", example, flags=re.I).strip(" .")
        parts.append("For example, " + example[0].lower() + example[1:] + ".")
    text = " ".join(part for part in parts if part)
    return public_text(_shorten_comment(text, 1100), reject_placeholders=True, reject_incomplete=True)


def _insert_blue_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _last_chapter_body_paragraph(document) -> Optional[Paragraph]:
    for paragraph in reversed(document.paragraphs):
        text = clean_text(paragraph.text)
        if not text:
            continue
        low = normalised(text)
        if "supervisor comment" in low or "missing section" in low:
            continue
        return paragraph
    return _first_academic_anchor(document) or _first_native_anchor(document)


def _add_missing_section_inline_bottom_notes(document, rows: Sequence[Dict[str, Any]]) -> None:
    if not _missing_section_inline_bottom_enabled():
        return
    comments: List[str] = []
    seen = set()
    for row in rows:
        if row.get("status") not in ACTIONABLE_STATUSES:
            continue
        if row.get("annotation_eligible") is False:
            continue
        if not _is_missing_section_finding(row):
            continue
        comment = _missing_section_bottom_comment(row)
        key = normalised(comment)
        if not comment or key in seen:
            continue
        seen.add(key)
        comments.append(comment)
    if not comments:
        return
    anchor = _last_chapter_body_paragraph(document)
    if anchor is None:
        return
    # Insert in reverse so the final visible order is heading, then numbered notes.
    for idx, comment in reversed(list(enumerate(comments, start=1))):
        note = _insert_blue_paragraph_after(anchor)
        run = note.add_run(f"{idx}. {comment}")
        run.font.color.rgb = INLINE_BLUE
        run.font.italic = True
        try:
            note.paragraph_format.left_indent = anchor.paragraph_format.left_indent
            note.paragraph_format.space_before = anchor.paragraph_format.space_after
            note.paragraph_format.space_after = anchor.paragraph_format.space_after
        except Exception:
            pass
    heading = _insert_blue_paragraph_after(anchor)
    lead = heading.add_run("Additional comment(s):")
    lead.bold = True
    lead.font.color.rgb = INLINE_BLUE
    lead.font.italic = True


def _native_comment_style() -> str:
    return (os.getenv("VPROF_NATIVE_COMMENT_STYLE") or "anchored_grouped").strip().lower()


def _merge_comments_by_section() -> bool:
    """Return True when native Word comments should be grouped professionally.

    Version 1.9.9.10 forced one Word comment per finding. That exposed every
    internal finding but made the document look over-commented and repetitive.
    The professional default is now to keep the full findings in the report but
    merge closely related findings into one numbered native comment box per
    section/anchor. The new style flag deliberately takes precedence over the
    older VPROF_EXPORT_ONE_COMMENT_PER_FINDING variable so deployments can move
    back to grouped comments with a single new env setting.
    """
    style = _native_comment_style()
    if style in {"anchored_grouped", "evidence_grouped", "numbered_grouped", "grouped", "section_grouped", "professional"}:
        return True
    if style in {"one_per_finding", "separate", "individual"}:
        return False
    return _env_bool("VPROF_COMMENT_MERGE_BY_SECTION", True)


def _export_one_comment_per_finding() -> bool:
    if _merge_comments_by_section():
        return False
    return _env_bool("VPROF_EXPORT_ONE_COMMENT_PER_FINDING", False)


def _split_related_concerns() -> bool:
    # Related concerns are kept together in numbered grouped comments unless a
    # deployment explicitly switches back to one-comment-per-finding mode.
    if _merge_comments_by_section():
        return False
    return _env_bool("VPROF_SPLIT_RELATED_CONCERNS_INTO_SEPARATE_COMMENTS", False)


def _include_section_review_comments() -> bool:
    return _env_bool("VPROF_INCLUDE_SECTION_REVIEW_COMMENTS", False)


def _max_items_per_native_comment() -> int:
    raw = os.getenv("VPROF_MAX_ITEMS_PER_NATIVE_COMMENT") or "3"
    try:
        return max(2, min(5, int(raw)))
    except ValueError:
        return 3


def _prepare_comment_list(comments: Iterable[str]) -> List[str]:
    unique: List[str] = []
    seen = set()
    for value in comments:
        text = _strip_visible_labels(
            public_text(value, limit=comment_max_chars(), reject_placeholders=True, reject_incomplete=True)
        ).strip("[] ").rstrip(" ;.")
        text = re.sub(r"^Supervisor comments?\s*:\s*", "", text, flags=re.I)
        if not text:
            continue
        if _split_related_concerns() and " A related concern is that " in text:
            first, second = text.split(" A related concern is that ", 1)
            candidates = [first, second[:1].upper() + second[1:]]
        else:
            candidates = [text]
        for candidate in candidates:
            candidate = clean_text(candidate).strip(" ;.")
            key = normalised(candidate)
            if not candidate or key in seen:
                continue
            seen.add(key)
            unique.append(_shorten_comment(candidate, comment_max_chars()).rstrip(" .") + ".")
    return unique




def _sanitise_rows_for_export(rows: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not _export_one_comment_per_finding():
        return sanitise_finding_rows(rows)
    output: List[Dict[str, Any]] = []
    seen_exact = set()
    for row in rows:
        cleaned = sanitise_finding_row(row)
        if cleaned is None:
            continue
        exact_key = (
            normalised(cleaned.get("section", "")),
            normalised(cleaned.get("item", "")),
            normalised(cleaned.get("comment", "")),
            normalised(cleaned.get("required_action", "")),
            tuple(
                str(item.get("paragraph") or item.get("paragraph_id") or "")
                for item in cleaned.get("evidence") or []
            ),
        )
        if exact_key in seen_exact:
            continue
        seen_exact.add(exact_key)
        output.append(cleaned)
    return output



def _expected_native_anchor_key(row: Dict[str, Any], evidence: Sequence[Dict[str, Any]]) -> Tuple[str, str, str]:
    evidence = _preferred_evidence(row, evidence) if evidence else []
    if evidence:
        best = evidence[0]
        if best.get("table_index") or best.get("table_number"):
            return ("table", str(best.get("table_index") or best.get("table_number") or ""), str(best.get("table_row") or ""))
        paragraph = str(best.get("paragraph") or best.get("paragraph_id") or "")
        if paragraph:
            return ("paragraph", paragraph, normalised(row.get("problematic_quote", ""))[:80])
    return ("unanchored", normalised(_canonical_group_label(row)), normalised(row.get("issue_title") or row.get("item") or "")[:80])

def expected_native_comment_count(review: Dict[str, Any]) -> int:
    rows = _sanitise_rows_for_export(
        list(review.get("academic_findings", []))
        + list(review.get("alignment_results", []))
        + list(review.get("revision_results", []))
    )
    actionable = [
        row for row in rows
        if row.get("status") in ACTIONABLE_STATUSES
        and row.get("annotation_eligible") is not False
        and not _is_missing_section_finding(row)
    ]
    if not _merge_comments_by_section():
        return len(actionable)
    groups = set()
    for row in actionable:
        evidence = [
            item for item in (row.get("evidence") or [])
            if item.get("document_role", "current") == "current"
        ]
        groups.add(_expected_native_anchor_key(row, evidence))
    return len(groups)


def _set_run_colour(run_element, colour: str) -> None:
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)
    colour_node = rpr.find(qn("w:color"))
    if colour_node is None:
        colour_node = OxmlElement("w:color")
        rpr.append(colour_node)
    colour_node.set(qn("w:val"), colour)


def _set_italic(run_element) -> None:
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)
    if rpr.find(qn("w:i")) is None:
        rpr.append(OxmlElement("w:i"))


def _run_element(text: str, source_run=None, colour: Optional[str] = None, italic: bool = False):
    element = OxmlElement("w:r")
    if source_run is not None and source_run._r.rPr is not None:
        element.append(deepcopy(source_run._r.rPr))
    if colour:
        _set_run_colour(element, colour)
    if italic:
        _set_italic(element)
    node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        node.set(XML_SPACE, "preserve")
    node.text = text
    element.append(node)
    return element


def _sentence_spans(text: str) -> List[Tuple[int, int, str]]:
    spans: List[Tuple[int, int, str]] = []
    for match in re.finditer(r"[^.!?\n]+(?:[.!?]+|$)", text or ""):
        start, end = match.span()
        if match.group(0).strip():
            spans.append((start, end, match.group(0)))
    if not spans and text:
        spans.append((0, len(text), text))
    return spans


def _expand_to_safe_text_span(text: str, start: int, end: int) -> Tuple[int, int]:
    """Avoid inserting reference markers inside a word.

    When extracted evidence quotes are very short or align inside a word, the
    previous exporter could create output such as "teache [6] r education".
    Expand to a sentence where possible, otherwise at least expand to whole-word
    boundaries before the native comment and red number are inserted.
    """
    if not text:
        return (0, 0)
    start = max(0, min(len(text), int(start)))
    end = max(start, min(len(text), int(end)))
    if start == end:
        return (0, len(text))
    # Short spans are usually weak anchors. Use the containing sentence instead.
    if end - start < 24:
        for s, e, _sentence in _sentence_spans(text):
            if s <= start and end <= e and e > s:
                return (s, e)
    while start > 0 and text[start - 1].isalnum() and text[start:start + 1].isalnum():
        start -= 1
    while end < len(text) and text[end - 1:end].isalnum() and text[end:end + 1].isalnum():
        end += 1
    # Prefer the full sentence when the adjusted span is still fragmentary.
    fragment = text[start:end].strip()
    if len(fragment.split()) < 6:
        for s, e, _sentence in _sentence_spans(text):
            if s <= start and end <= e and e > s:
                return (s, e)
    return start, end


def _best_span(text: str, matched_terms: Iterable[str], problematic_quote: str = "") -> Tuple[int, int]:
    quote = clean_text(problematic_quote)
    if quote:
        exact_start = text.find(quote)
        if exact_start >= 0:
            return _expand_to_safe_text_span(text, exact_start, exact_start + len(quote))
        normalised_quote = normalised(quote)
        for start, end, sentence in _sentence_spans(text):
            if normalised_quote and normalised_quote in normalised(sentence):
                return _expand_to_safe_text_span(text, start, end)
    terms = [normalised(term) for term in matched_terms if normalised(term)]
    spans = _sentence_spans(text)
    if not spans:
        return (0, len(text or ""))
    ranked = []
    for start, end, sentence in spans:
        low = normalised(sentence)
        hits = sum(1 for term in terms if term in low)
        ranked.append((hits, len(sentence.strip()), start, end))
    ranked.sort(key=lambda item: (item[0], item[1]), reverse=True)
    _, _, start, end = ranked[0]
    return _expand_to_safe_text_span(text, start, end)


def _sanitise_guidance(value: str) -> str:
    text = _strip_visible_labels(value)
    patterns = [
        r"^(?:retain|keep) this finding and\s+",
        r"^require (?:the student to )?",
        r"^ask the student to\s+",
        r"^the student should\s+",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text, flags=re.I)
    text = re.sub(r"^retain this point and\s+", "", text, flags=re.I)
    if text:
        text = text[0].upper() + text[1:]
    return text.rstrip()



def _normalise_action_start(text: str) -> str:
    """Return a natural supervisor instruction without template residue."""
    value = clean_text(text).rstrip(" .")
    if not value:
        return ""
    # Common model phrasing such as "by use" or "by undertake" must not
    # reach native Word comments. Convert it to direct supervision.
    value = re.sub(r"^revise(?: the)?(?: marked)? passage by\s+", "", value, flags=re.I)
    value = re.sub(r"^by\s+", "", value, flags=re.I)
    gerund_map = {
        "using": "use", "undertaking": "undertake", "applying": "apply",
        "providing": "provide", "linking": "link", "situating": "situate",
        "differentiating": "differentiate", "checking": "check", "clarifying": "clarify",
        "replacing": "replace", "rewriting": "rewrite", "aligning": "align",
        "expanding": "expand", "stating": "state", "defining": "define",
        "supporting": "support", "removing": "remove", "correcting": "correct",
        "ensuring": "ensure", "explaining": "explain", "adding": "add",
        "verifying": "verify", "inserting": "insert", "avoiding": "avoid",
        "developing": "develop", "formulating": "formulate", "showing": "show",
        "demonstrating": "demonstrate", "indicating": "indicate",
    }
    parts = value.split(maxsplit=1)
    if parts and parts[0].lower() in gerund_map:
        value = gerund_map[parts[0].lower()] + ((" " + parts[1]) if len(parts) > 1 else "")
    return value[0].upper() + value[1:] if value else value


def _shorten_comment(value: str, limit: Optional[int] = None) -> str:
    effective_limit = limit if limit is not None else comment_max_chars()
    return sentence_safe_trim(value, effective_limit)

def _comment_body(row: Dict[str, Any]) -> str:
    safe_row = sanitise_finding_row(row)
    if safe_row is None:
        return ""
    row = safe_row

    reference = clean_text(
        row.get("reference_label")
        or row.get("section_reference")
        or row.get("section")
    )
    if not row.get("table_reference"):
        table_evidence = next(
            (item for item in row.get("evidence") or [] if item.get("table_number")),
            None,
        )
        if table_evidence:
            number = clean_text(table_evidence.get("table_number", ""))
            title = clean_text(table_evidence.get("table_title", ""))
            table_reference = f"Table {number}" if number else "Table"
            if title:
                table_reference += f": {title}"
            reference = f"{reference}, {table_reference}" if reference else table_reference

    issue = _sanitise_guidance(row.get("item", ""))
    section_label = clean_text(row.get("section_reference") or row.get("section") or "")
    if (reference and normalised(issue) == normalised(reference)) or (section_label and normalised(issue) == normalised(section_label)):
        issue = ""
    assessment = _sanitise_guidance(row.get("comment", "") or row.get("assessment", ""))
    consequence = _sanitise_guidance(
        row.get("academic_consequence", "")
        or row.get("consequence", "")
        or row.get("why_it_matters", "")
    )
    action = _sanitise_guidance(row.get("required_action", ""))
    example = _sanitise_guidance(row.get("illustrative_guidance", ""))
    if not example:
        example = _sanitise_guidance(context_specific_example(row))
    example = re.sub(r"^for example[:,]?\s*", "", example, flags=re.I)

    heading = reference or "Supervisor review"
    parts: List[str] = []

    if issue:
        parts.append(issue.rstrip(" .") + ".")
    if assessment:
        parts.append(assessment.rstrip(" .") + ".")
    if consequence and normalised(consequence) not in normalised(assessment):
        parts.append(consequence.rstrip(" .") + ".")

    level = academic_level_label(row.get("_academic_level") or row.get("academic_level"))
    combined_so_far = " ".join(parts)
    explicit_level_sentence = re.search(
        r"(?:^|[.!?]\s+)At\s+(?:PhD|MPhil|professional doctorate|Master's|non-research Master's|Bachelor's)\s+level\b",
        combined_so_far,
    )
    if level != "the applicable academic level" and not explicit_level_sentence:
        category_text = normalised(" ".join(clean_text(row.get(field, "")) for field in ("category", "section", "item", "comment")))
        if any(term in category_text for term in ("result", "statistic", "analysis", "regression", "anova", "sem", "mediation", "moderation", "table")):
            expectation = "the analysis should be sufficiently complete and internally consistent for an examiner to trace each reported conclusion to the relevant table, model and diagnostic evidence"
        elif any(term in category_text for term in ("method", "design", "sampling", "instrument", "validity", "reliability", "ethics")):
            expectation = "the methodological choices should be justified, reproducible and explicitly aligned with the objectives, data and analysis"
        elif any(term in category_text for term in ("discussion", "interpretation", "contribution", "theory")):
            expectation = "the work should demonstrate independent scholarly interpretation, theoretical integration and defensible contribution"
        else:
            expectation = "the argument should demonstrate the precision, depth and independent scholarly judgement expected"
        parts.append(f"At {level}, {expectation}.")

    if action:
        action_text = _normalise_action_start(action)
        if re.match(r"^(?:revise|rewrite|replace|align|clarify|expand|state|define|support|remove|correct|ensure|explain|add|verify|use|undertake|apply|provide|insert|avoid|check|develop|formulate|show|demonstrate|indicate|link|situate|differentiate|populate|supply|fix|interpret|qualify|separate|clean)\b", action_text, flags=re.I):
            parts.append(action_text + ".")
        else:
            parts.append("Revise the marked passage so that it " + action_text[0].lower() + action_text[1:] + ".")
    elif assessment:
        parts.append("Revise the marked passage so the academic point is clear, properly supported and aligned with the section purpose.")
    if example:
        example_text = _normalise_action_start(example)
        if example_text:
            parts.append("For example, " + example_text[0].lower() + example_text[1:] + ".")

    deduped_parts: List[str] = []
    seen_parts = set()
    for part in parts:
        key = normalised(part)
        if not key or key in seen_parts:
            continue
        if any(_comment_similarity(key, existing) >= 0.88 for existing in seen_parts):
            continue
        seen_parts.add(key)
        deduped_parts.append(part)
    parts = deduped_parts

    body = f"{heading}: " + " ".join(parts) if parts else f"{heading}: Revise this passage to address the identified academic weakness."
    # Manual-confirmation and provider-failure status belongs in the internal
    # audit trail, never in a student's Word comment. Student-facing comments
    # remain developmental but must read as natural supervision, not as a
    # labelled template.
    body = _strip_visible_labels(body)
    body = professionalise_reviewer_language(body, row.get("_academic_level") or row.get("academic_level"))
    return public_text(_shorten_comment(body), reject_placeholders=True, reject_incomplete=True)


_LEVEL_PHRASE_RE = re.compile(
    r"\s*At MPhil level,? the section should show independent research judgement, conceptual clarity, methodological defensibility and traceable scholarly contribution\.?:?",
    flags=re.I,
)


def _remove_level_repetition(value: str) -> str:
    # Preserve a genuine level-specific expectation such as "At PhD level".
    # Only collapse whitespace; old versions removed the level statement and
    # made the comment sound generic.
    return re.sub(r"\s{2,}", " ", clean_text(value)).strip()


def _split_example(value: str) -> Tuple[str, str]:
    parts = re.split(r"\bFor example,\s*", value, maxsplit=1, flags=re.I)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip(" .")
    return value.strip(), ""


def _compact_group_item(value: str) -> Tuple[str, str]:
    value = _strip_comment_reference(value)
    text = _strip_visible_labels(
        public_text(value, limit=comment_max_chars(), reject_placeholders=True, reject_incomplete=True)
    ).strip("[] ").rstrip(" ;.")
    text = re.sub(r"^Supervisor comments?\s*:\s*", "", text, flags=re.I)
    text = _remove_level_repetition(text)
    core, example = _split_example(text)
    # When a grouped comment is anchored on a section heading, repeating
    # "Problem Statement:" or "Introduction:" inside every numbered item looks
    # mechanical. Remove only short heading prefixes; the section location is
    # already provided by the Word anchor and the report.
    core = re.sub(
        r"^.{2,180}:\s*(?=(?:Add|Align|Ask|Avoid|Check|Clarify|Clean|Correct|Define|Develop|Ensure|Explain|Expand|Formulate|Interpret|Insert|Link|Provide|Qualify|Remove|Replace|Revise|Rewrite|Separate|Show|State|Support|Use|Verify)\b)",
        "",
        core,
        flags=re.I,
    ).strip()
    core = re.sub(r"^(?:\d+(?:\.\d+){0,4}\s*)?[A-Za-z][A-Za-z0-9/&() \-]{2,90}:\s*", "", core).strip()
    # Keep the full local guidance where possible. The comment is already
    # grouped by the exact evidence passage, and using only the first two
    # sentences can accidentally cut decimal headings such as 4.2 or table
    # numbers such as Table 4.1.
    core = _shorten_comment(core, 560).rstrip(" .")
    example = _shorten_comment(example, 320).rstrip(" .") if example else ""
    return core, example


def _anchor_context_text(anchor_context: str) -> str:
    text = clean_text(anchor_context)
    if not text:
        return "this marked passage"
    text = re.sub(r"\s+", " ", text).strip()
    text = _shorten_comment(text, 150).strip(" .")
    if len(text.split()) <= 3:
        return f"the phrase '{text}'"
    return f"the marked passage beginning '{text}'"


def _red_marker(text: str) -> str:
    if not text:
        return ""
    return "[[VPROF_RED:" + text.replace("]]", "") + "]]"


def _with_comment_reference(number: int, comment: str) -> str:
    if not _sequential_reference_numbers_enabled():
        return comment
    text = clean_text(comment)
    if not text:
        return ""
    return f"[[VPROF_REFNO:{int(number)}]]" + text


def _comment_reference_number(comment: str) -> Optional[int]:
    match = _REFNO_RE.search(comment or "")
    if not match:
        return None
    try:
        return int(match.group(1))
    except (TypeError, ValueError):
        return None


def _strip_comment_reference(comment: str) -> str:
    return _REFNO_RE.sub("", comment or "", count=1).strip()


def _format_comment_group(comments: Iterable[str], anchor_context: str = "") -> str:
    """Format grouped text for the Word Review comment pane.

    Numbering is now global across the chapter. The same red number appears
    beside the exact passage in the document body and beside the corresponding
    item in the native Word comment box. This keeps grouped comments
    professional without making students guess which point applies where.
    """
    unique: List[Tuple[Optional[int], str]] = []
    examples: List[str] = []
    seen = set()
    for value in comments:
        ref_no = _comment_reference_number(value)
        item, example = _compact_group_item(value)
        key = normalised(item)
        if not item or not key:
            continue
        if key in seen:
            continue
        if any(_comment_similarity(key, existing_key) >= 0.74 for existing_key in seen):
            continue
        seen.add(key)
        unique.append((ref_no, item.rstrip(" .") + "."))
        if example and all(_comment_similarity(normalised(example), normalised(existing)) < 0.60 for existing in examples):
            examples.append(example.rstrip(" .") + ".")
        if len(unique) >= _max_items_per_native_comment():
            break
    if not unique:
        return ""
    parts = []
    for local_idx, (ref_no, item) in enumerate(unique, start=1):
        number = ref_no if ref_no is not None else local_idx
        parts.append(_red_marker(f"{number}. ") + item)
    body = " ".join(parts)
    if examples:
        example = examples[0]
        body = body.rstrip() + " For example, " + example[0].lower() + example[1:]
    return _shorten_comment(body, comment_max_chars())


def _is_synthetic_section_heading(value: str) -> bool:
    low = normalised(value)
    return any(term in low for term in (
        "whole chapter coherence",
        "whole chapter consistency",
        "cross chapter coherence",
        "cross chapter alignment",
        "supervisor comment compliance audit",
        "alignment audit",
        "revision audit",
    ))


_GENERIC_SECTION_ASSESSMENT_RE = re.compile(
    r"\b(?:reviewed against|selected academic level|no major issue|appears adequate|looks adequate|section was reviewed|this section has been reviewed|check that its purpose|should be checked|contribution to the chapter\'s argument|generally acceptable|meets the expected standard)\b",
    flags=re.I,
)


def _section_comment_template(heading: str, academic_level: Any = None) -> str:
    """Return a section-specific supervisory note rather than a generic coverage stamp."""
    low = normalised(heading)
    level = academic_level_label(academic_level)
    level_phrase = f"At {level}" if level != "the applicable academic level" else "At the applicable academic level"
    if "background" in low:
        return (
            "This part should move logically from the broad sustainability debate to the specific Ghanaian and sectoral context of the study. "
            "It should introduce the main constructs, show how they relate and prepare the reader for the problem statement rather than merely listing prior studies. "
            "Strengthen the progression, localise the evidence and ensure that every central construct in the objectives is introduced before the problem is stated."
        )
    if "statement" in low and "problem" in low:
        return (
            "This part should establish a defensible research problem, not only a general topic of interest. "
            "It should separate the practical problem, empirical gap, contextual gap and methodological gap, then show why the selected study setting requires investigation. "
            "Revise it so the gap directly leads to the purpose, objectives and questions."
        )
    if "purpose" in low:
        return (
            "The purpose statement should express the central intent of the whole study in one coherent frame. "
            "For MPhil-level work, it must cover every principal construct and outcome that later appears in the objectives and questions. "
            "Revise the statement so a reader can trace the design, analysis and conclusions back to this purpose."
        )
    if "objective" in low:
        return (
            "The objectives should be measurable, ordered and fully aligned with the title, purpose, research questions and proposed method. "
            "Avoid mixing descriptive, relational and causal intentions without explaining the analytical logic. "
            "Revise the objectives so each one can be answered by a clearly identifiable data source and analysis procedure."
        )
    if "question" in low:
        return (
            "The research questions should mirror the objectives one-to-one and use language that matches the intended design. "
            "Terms such as relationship, effect and impact imply different levels of analysis and should not be used interchangeably. "
            "Edit the questions for alignment, punctuation and methodological consistency."
        )
    if "significance" in low:
        return (
            "This part should explain the expected value of the study without reporting findings that have not yet been produced. "
            f"{level_phrase}, the contribution should be tied to evidence, policy, practice and scholarship in a balanced way. "
            "Rewrite stakeholder benefits prospectively and avoid claims that imply the study has already demonstrated an effect."
        )
    if "limitation" in low:
        return (
            "This part should identify unavoidable weaknesses in the design and explain how they may affect interpretation. "
            "The tense must match the stage of the work: proposed studies should not describe sampling, response or fieldwork constraints as completed events. "
            "Revise the section so the limitations are realistic, method-linked and consistently expressed."
        )
    if "delimitation" in low:
        return (
            "This part should define the deliberate boundaries of the study, including sector, location, respondents, constructs and time scope. "
            "A reader should be able to tell exactly what is included and excluded from the inquiry. "
            "Replace unresolved prompts and ensure the stated boundaries match the methodology chapter."
        )
    if "definition" in low or "terms" in low:
        return (
            "This part should define key constructs in a way that is conceptually clear and measurable. "
            "Avoid circular definitions, absolute claims and definitions that do not correspond to the proposed indicators. "
            "Revise each definition so it states the construct boundary, relevant dimensions and link to the study context."
        )
    if "organisation" in low or "organization" in low:
        return (
            "This part should provide a concise map of the remaining chapters without making claims about results not yet produced. "
            "The description should match the actual thesis structure and maintain the same proposal or completed-study tense used elsewhere. "
            "Revise it after the chapter structure is finalised to avoid inconsistency."
        )
    if "reference" in low:
        return (
            "This part should correspond exactly with the in-text citations and follow the required referencing style. "
            "Check author spelling, publication year, DOI completeness, source credibility and whether every listed item is cited in the chapter. "
            "Remove uncited or unverifiable entries and correct mismatches before submission."
        )
    return (
        "This section should be assessed for its role in the chapter, the quality of evidence used, conceptual clarity and alignment with the study purpose. "
        "Revise any wording, citation or structural element that weakens the reader's ability to trace the argument from the problem to the methodology. "
        f"{level_phrase}, keep the section focused on the depth, rigour and approved thesis format expected."
    )


def _is_weak_section_assessment(value: str) -> bool:
    text = clean_text(value)
    if not text:
        return True
    low = normalised(text)
    if _GENERIC_SECTION_ASSESSMENT_RE.search(text):
        return True
    # A useful section comment should contain at least one supervision cue, not merely a coverage stamp.
    cues = ("align", "evidence", "construct", "gap", "method", "purpose", "objective", "revise", "clarify", "support", "scope", "contribution", "citation")
    if any(cue in low for cue in cues):
        return False
    return len(text) < 140


def _polish_section_assessment(value: str) -> str:
    text = _strip_visible_labels(value)
    text = re.sub(r"\bDocument-level review note\.\s*", "", text, flags=re.I)
    text = re.sub(r"\bA separate model response for this section remained unavailable[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"\bThe section '[^']+' is present, but its separate expert review could not be completed after focused recovery[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"\bIt has therefore not been treated as absent and no unverified finding has been added[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"\bNo unsupported criticism has been inserted[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"\bManual confirmation of this section is recommended[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"\bRecovery detail:[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"[^.!?]*\b(?:manifest|document map|paragraph id|section packet|parser|fallback|recovery|focused recovery|model response|P\d{1,4})\b[^.!?]*(?:[.!?]|$)", "", text, flags=re.I)
    text = re.sub(r"\bThis section (?:has been|was) reviewed(?: against [^.]+)?\.?\s*", "", text, flags=re.I)
    text = re.sub(r"\bIt appears adequate\.?\s*", "", text, flags=re.I)
    text = re.sub(r"\bNo major issue(?:s)? (?:was|were) found\.?\s*", "", text, flags=re.I)
    text = re.sub(r"\s{2,}", " ", text).strip(" ;,.")
    if text:
        text = text[0].upper() + text[1:]
    return text


def _section_review_comment(row: Dict[str, Any]) -> str:
    """Build a natural, section-specific comment from the stored section review.

    Section comments must not be empty coverage stamps. They should tell the
    student what the section was assessed for and what kind of revision would
    improve it, while issue comments remain attached to exact defects.
    """
    heading = clean_text(row.get("heading") or row.get("section_name") or "Section")
    if not heading or _is_synthetic_section_heading(heading):
        return ""

    level = row.get("_academic_level") or row.get("academic_level")
    raw_assessment = professionalise_reviewer_language(
        _polish_section_assessment(clean_text(row.get("section_assessment", ""))), level
    )
    warning = professionalise_reviewer_language(
        _polish_section_assessment(clean_text(row.get("coverage_warning", ""))), level
    )

    # Do not release false section-coverage claims such as "no terms are
    # defined" when the heading exists. Definition sections are frequently
    # multi-paragraph and can be misread by a section packet. Exact issue
    # comments still flag weak or circular definitions where evidence exists.
    if ("definition" in normalised(heading) or "terms" in normalised(heading)) and re.search(
        r"\b(?:no|not|without)\b.{0,40}\b(?:definition|definitions|terms)\b|\b(?:definition|terms)\b.{0,40}\b(?:absent|missing|not defined|no definitions)\b",
        raw_assessment,
        flags=re.I,
    ):
        raw_assessment = ""

    if _is_weak_section_assessment(raw_assessment):
        assessment = _section_comment_template(heading, level)
    else:
        assessment = raw_assessment
        # Add a section-specific supervisory focus when the model assessment is
        # useful but too narrow to stand alone as a section coverage comment.
        if len(assessment) < 260:
            assessment = assessment.rstrip(" .") + ". " + _section_comment_template(heading, level)

    assessment = public_text(
        assessment,
        limit=760,
        reject_placeholders=True,
        reject_incomplete=True,
    )
    warning = public_text(
        warning,
        limit=320,
        reject_placeholders=True,
        reject_incomplete=True,
    )

    body = assessment or _section_comment_template(heading, level)
    if warning and normalised(warning) not in normalised(body):
        body = (body.rstrip(" .") + ". " + warning.rstrip(" .") + ".").strip()
    body = re.sub(r"^this section\s*[:\-]\s*", "", body, flags=re.I).strip()
    body = _strip_visible_labels(body)
    return _shorten_comment(f"{heading}: {body}", max(560, min(comment_max_chars(), 980)))


def _section_review_key(row: Dict[str, Any]) -> Tuple[Optional[int], Tuple[str, ...]]:
    heading = clean_text(row.get("heading") or row.get("section_name") or "")
    chapter_number = row.get("chapter_number")
    try:
        chapter_number = int(chapter_number) if chapter_number is not None else None
    except (TypeError, ValueError):
        chapter_number = None
    path = tuple(clean_text(value) for value in row.get("section_path") or [] if clean_text(value))
    if heading and (not path or normalised(path[-1]) != normalised(heading)):
        path = path + (heading,)
    elif not path and heading:
        path = (heading,)
    return chapter_number, path


def _add_section_review_comments(
    document,
    review: Dict[str, Any],
    *,
    author: str,
    initials: str,
    fallback_comments: List[str],
) -> None:
    """Add one native Word comment for every reviewed section/subsection.

    The AI engine already returns section assessments. Earlier exports showed
    only issue comments, so sections without issues looked unreviewed. This
    pass anchors each section assessment to the exact section heading where
    possible and falls back to the first paragraph only when the heading cannot
    be uniquely located.
    """
    seen: set[Tuple[Optional[int], Tuple[str, ...]]] = set()
    for row in review.get("academic_section_reviews") or []:
        heading = clean_text(row.get("heading") or row.get("section_name") or "")
        if not heading or _is_synthetic_section_heading(heading):
            continue
        key = _section_review_key(row)
        if key in seen:
            continue
        comment = _section_review_comment(row)
        if not comment:
            continue
        seen.add(key)
        chapter_number, headings = key
        target = _find_heading(document, headings or (heading,), chapter_number=chapter_number)
        if target is not None and _comment_on_paragraph(
            document, target, [comment], author=author, initials=initials
        ):
            continue
        # If the exact heading cannot be located, keep the section assessment
        # as a whole-chapter native comment rather than silently losing it.
        # The document-level anchor deliberately avoids title-page boilerplate.
        fallback_comments.append(comment)
        continue


def _comment_similarity(left_key: str, right_key: str) -> float:
    topical_clusters = [
        ("purpose", "objectiv", "align"),
        ("purpose", "narrow", "objectiv"),
        ("significance", "result", "proposal"),
        ("definition", "construct", "circular"),
        ("limitation", "tense", "proposal"),
        ("causal", "design", "cross"),
        ("problem", "determinant", "objective"),
        ("theor", "framework", "background"),
        ("gap", "context", "problem"),
        ("citation", "punctuat", "format"),
    ]
    for cluster in topical_clusters:
        if all(token in left_key for token in cluster) and all(token in right_key for token in cluster):
            return 0.92
    left_tokens = {token for token in re.findall(r"[a-z0-9]+", left_key) if len(token) >= 4}
    right_tokens = {token for token in re.findall(r"[a-z0-9]+", right_key) if len(token) >= 4}
    token_score = (len(left_tokens & right_tokens) / len(left_tokens | right_tokens)) if left_tokens and right_tokens else 0.0
    sequence_score = __import__("difflib").SequenceMatcher(None, left_key, right_key).ratio()
    return max(token_score, sequence_score)

def _replace_run_with_parts(run, before: str, marked: str, after: str):
    """Split a plain-text run without changing its visible formatting.

    Native Word comments must anchor to run boundaries. Splitting a run gives
    the comment an exact range while preserving the original text and run
    properties. No colour, inserted note, or tracked edit is added.
    """
    parent = run._r.getparent()
    index = parent.index(run._r)
    created = []
    for value in (before, marked, after):
        if not value:
            created.append(None)
            continue
        element = _run_element(value, source_run=run)
        parent.insert(index, element)
        index += 1
        created.append(element)
    parent.remove(run._r)
    return created


def _reviewer_initials(name: str) -> str:
    words = [part for part in re.split(r"\s+", clean_text(name)) if part]
    if not words:
        return "RV"
    initials = "".join(part[0] for part in words if part[0].isalnum()).upper()
    return (initials or "RV")[:8]


def _comment_identity(
    review: Dict[str, Any], comment_author: Optional[str] = None
) -> Tuple[str, str]:
    summary = review.get("summary") or {}
    metadata = review.get("assessment_metadata") or {}
    author = clean_text(
        comment_author
        or summary.get("reviewer_name")
        or summary.get("examiner_name")
        or metadata.get("examiner_name")
        or review.get("reviewer_name")
        or "Reviewer"
    )
    return author, _reviewer_initials(author)


def _write_rich_comment_text(comment_obj, text: str) -> None:
    """Write comment text, colouring VPROF_RED marker segments red.

    python-docx comments support rich runs, so grouped comments can show the
    numbered location cue in red while keeping the explanatory guidance in the
    normal comment text colour.
    """
    paragraph = comment_obj.paragraphs[0] if comment_obj.paragraphs else comment_obj.add_paragraph("")
    pos = 0
    for match in _RICH_RED_RE.finditer(text):
        before = text[pos:match.start()]
        if before:
            paragraph.add_run(before)
        red_text = match.group(1)
        if red_text:
            run = paragraph.add_run(red_text)
            run.font.color.rgb = COMMENT_RED
            run.bold = True
        pos = match.end()
    tail = text[pos:]
    if tail:
        paragraph.add_run(tail)


def _add_native_comment(
    document,
    runs: Sequence[Run],
    comment: str,
    *,
    author: str,
    initials: str,
) -> bool:
    usable = [run for run in runs if run is not None and clean_text(run.text)]
    if not usable:
        return False
    if not hasattr(document, "add_comment"):
        raise RuntimeError(
            "Native Word comments require python-docx 1.2.0 or newer."
        )
    text = clean_text(comment)
    if _RICH_RED_RE.search(text):
        comment_obj = document.add_comment(
            runs=usable,
            text="",
            author=author,
            initials=initials,
        )
        _write_rich_comment_text(comment_obj, text)
    else:
        document.add_comment(
            runs=usable,
            text=text,
            author=author,
            initials=initials,
        )
    return True


def _group_reference_numbers_from_comment(comment: str) -> List[int]:
    """Extract grouped comment item numbers that need visible body references."""
    numbers = []
    for match in re.finditer(r"\[\[VPROF_RED:(\d+)\.\s*", comment or ""):
        try:
            numbers.append(int(match.group(1)))
        except (TypeError, ValueError):
            continue
    # Only insert body markers for grouped comments. A single native comment is
    # already clear from the highlighted text and does not need an extra marker.
    unique = []
    for number in numbers:
        if number not in unique:
            unique.append(number)
    return unique


def _insert_red_reference_markers_after_span(
    paragraph: Paragraph,
    trailing_element,
    reference_numbers: Sequence[int],
    source_run: Optional[Run],
) -> None:
    """Insert red [1] [2] markers in the document body beside the anchor text.

    The numbered markers correspond to the numbered items in the native Word
    comment box. This gives the student an immediate visual map from the text to
    the grouped comment without using vague phrases such as "applies to the
    marked passage".
    """
    if not reference_numbers:
        return
    marker_text = " " + " ".join(f"[{number}]" for number in reference_numbers)
    marker = _run_element(marker_text, source_run=source_run, colour="C00000")
    parent = trailing_element.getparent() if trailing_element is not None else paragraph._p
    if trailing_element is not None and trailing_element.getparent() is not None:
        index = parent.index(trailing_element) + 1
        parent.insert(index, marker)
    else:
        paragraph._p.append(marker)


def _mark_span_and_insert_comment(
    document,
    paragraph: Paragraph,
    start: int,
    end: int,
    comment: str,
    *,
    author: str,
    initials: str,
) -> bool:
    """Anchor a native comment to an exact text range.

    The paragraph text and visible formatting are preserved. The selected text
    is not recoloured and no comment paragraph is inserted into the document.
    """
    if start >= end:
        return False
    runs = list(paragraph.runs)
    cursor = 0
    marked_elements = []
    trailing_element = None
    trailing_source_run = None

    for run in runs:
        text = run.text or ""
        run_start, run_end = cursor, cursor + len(text)
        cursor = run_end
        if not text or run_end <= start or run_start >= end:
            continue

        local_start = max(0, start - run_start)
        local_end = min(len(text), end - run_start)
        before = text[:local_start]
        marked = text[local_start:local_end]
        after = text[local_end:]
        created = _replace_run_with_parts(run, before, marked, after)
        if created[1] is not None:
            marked_elements.append(created[1])
            trailing_element = created[1]
            trailing_source_run = run

    if not marked_elements:
        return False
    reference_numbers = _group_reference_numbers_from_comment(comment)
    if _native_group_location_markers_enabled() and reference_numbers:
        _insert_red_reference_markers_after_span(
            paragraph, trailing_element, reference_numbers, trailing_source_run
        )
    anchor_runs = [Run(element, paragraph) for element in marked_elements]
    return _add_native_comment(
        document, anchor_runs, comment, author=author, initials=initials
    )

def _merge_nearby_span_groups(
    span_groups: Dict[Tuple[int, int], List[str]],
    max_gap: int = 24,
) -> List[Tuple[Tuple[int, int], List[str]]]:
    ordered = sorted(span_groups.items(), key=lambda item: item[0][0])
    merged: List[Tuple[Tuple[int, int], List[str]]] = []
    for (start, end), comments in ordered:
        if not merged:
            merged.append(((start, end), list(comments)))
            continue
        (previous_start, previous_end), previous_comments = merged[-1]
        if start <= previous_end + max_gap:
            merged[-1] = ((previous_start, max(previous_end, end)), previous_comments + list(comments))
        else:
            merged.append(((start, end), list(comments)))
    return merged


_TABLE_CAPTION_RE = re.compile(
    r"^\s*table\s+(?P<number>[A-Za-z]?\d+(?:\.\d+)*|[IVXLC]+)\b\s*[:.\-–—]?\s*(?P<title>.*)$",
    flags=re.I,
)


def _docx_blocks(document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _source_locator_map(document):
    output: Dict[int, Dict[str, Any]] = {}
    tables: Dict[int, Dict[str, Any]] = {}
    paragraph_no = 0
    table_index = 0
    pending_caption: Optional[Dict[str, Any]] = None
    active_chapter: Optional[int] = None
    chapter_words = {
        "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
        "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
    }

    for block in _docx_blocks(document):
        if isinstance(block, Table):
            table_index += 1
            table_info = {
                "table": block,
                "caption_paragraph": (pending_caption or {}).get("paragraph"),
                "table_number": (pending_caption or {}).get("table_number", str(table_index)),
                "table_title": (pending_caption or {}).get("table_title", ""),
            }
            tables[table_index] = table_info
            caption_number = int((pending_caption or {}).get("paragraph_number") or 0)
            if caption_number and caption_number in output:
                output[caption_number].update({
                    "kind": "table_caption",
                    "table_index": table_index,
                    "table": block,
                    **table_info,
                })
            for row_index, row in enumerate(block.rows, start=1):
                values = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
                if not values:
                    continue
                paragraph_no += 1
                cell_paragraphs = []
                for cell in row.cells:
                    cell_paragraphs.extend(
                        paragraph for paragraph in cell.paragraphs
                        if clean_text(paragraph.text)
                    )
                output[paragraph_no] = {
                    "kind": "table_row",
                    "chapter_number": active_chapter,
                    "table_index": table_index,
                    "table_row": row_index,
                    "table": block,
                    "cell_paragraphs": cell_paragraphs,
                    **table_info,
                }
            pending_caption = None
            continue

        text = clean_text(block.text)
        if not text:
            continue
        chapter_match = re.fullmatch(
            r"chapter\s+(one|two|three|four|five|six|seven|eight|nine|ten|[1-9]|10)",
            normalised(text),
        )
        if chapter_match:
            token = chapter_match.group(1)
            active_chapter = int(token) if token.isdigit() else chapter_words[token]
        paragraph_no += 1
        output[paragraph_no] = {
            "kind": "paragraph",
            "chapter_number": active_chapter,
            "paragraph": block,
        }
        match = _TABLE_CAPTION_RE.match(text)
        if match:
            pending_caption = {
                "paragraph": block,
                "paragraph_number": paragraph_no,
                "table_number": clean_text(match.group("number")),
                "table_title": clean_text(match.group("title")),
            }

    return output, tables


def _strip_heading_number(value: str) -> str:
    return re.sub(r"^\d+(?:\.\d+){0,4}\s+", "", normalised(value)).strip()


def _find_heading(
    document,
    headings: Iterable[str],
    chapter_number: Optional[int] = None,
) -> Optional[Paragraph]:
    """Find an exact heading only. Ambiguous or partial matches are rejected."""
    targets = {normalised(value) for value in headings if normalised(value)}
    stripped_targets = {_strip_heading_number(value) for value in headings if _strip_heading_number(value)}
    candidates: List[Paragraph] = []
    active_chapter: Optional[int] = None
    chapter_words = {
        "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
        "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
    }
    for paragraph in document.paragraphs:
        raw = clean_text(paragraph.text)
        low = normalised(raw)
        if not low or "supervisor comment" in low:
            continue
        chapter_match = re.fullmatch(
            r"chapter\s+(one|two|three|four|five|six|seven|eight|nine|ten|[1-9]|10)",
            low,
        )
        if chapter_match:
            token = chapter_match.group(1)
            active_chapter = int(token) if token.isdigit() else chapter_words[token]
        style_name = ""
        try:
            style_name = (paragraph.style.name or "").lower()
        except Exception:
            pass
        looks_like_heading = (
            "heading" in style_name
            or "title" in style_name
            or len(raw.split()) <= 15
            or bool(re.match(r"^\d+(?:\.\d+){0,3}\s+", raw))
        )
        if not looks_like_heading:
            continue
        exact = low in targets or _strip_heading_number(raw) in stripped_targets
        if not exact:
            continue
        if chapter_number is not None and active_chapter != chapter_number:
            continue
        candidates.append(paragraph)
    return candidates[0] if len(candidates) == 1 else None


def _comment_on_paragraph(
    document,
    paragraph: Paragraph,
    comments: List[str],
    *,
    author: str,
    initials: str,
) -> bool:
    runs = [run for run in paragraph.runs if clean_text(run.text)]
    if not runs:
        return False
    if _export_one_comment_per_finding():
        added = False
        for comment in _prepare_comment_list(comments):
            if _add_native_comment(document, runs, comment, author=author, initials=initials):
                added = True
        return added
    if len(comments) == 1 and re.match(r"^\s*1\.\s+", comments[0]) and re.search(r"\b2\.\s+", comments[0]):
        grouped = comments[0]
    else:
        grouped = _format_comment_group(comments, anchor_context=paragraph.text)
    if grouped and _native_group_location_markers_enabled():
        numbers = _group_reference_numbers_from_comment(grouped)
        if numbers:
            _insert_red_reference_markers_after_span(
                paragraph, runs[-1]._r, numbers, runs[-1]
            )
    return bool(
        grouped
        and _add_native_comment(
            document, runs, grouped, author=author, initials=initials
        )
    )


def _comment_on_table(
    document,
    table: Table,
    comments: List[str],
    caption: Optional[Paragraph] = None,
    *,
    author: str,
    initials: str,
) -> bool:
    prepared = _prepare_comment_list(comments) if _export_one_comment_per_finding() else [_format_comment_group(comments, anchor_context=(caption.text if caption is not None else "the table"))]
    prepared = [comment for comment in prepared if comment]
    if not prepared:
        return False
    if caption is not None:
        runs = [run for run in caption.runs if clean_text(run.text)]
        if runs:
            added = False
            for comment in prepared:
                if _add_native_comment(document, runs, comment, author=author, initials=initials):
                    added = True
            if added:
                return True
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                runs = [run for run in paragraph.runs if clean_text(run.text)]
                if not runs:
                    continue
                added = False
                for comment in prepared:
                    if _add_native_comment(
                        document,
                        runs,
                        comment,
                        author=author,
                        initials=initials,
                    ):
                        added = True
                if added:
                    return True
    return False



def _comment_on_table_row(
    document,
    table: Table,
    row_index: int,
    comments: List[str],
    *,
    author: str,
    initials: str,
) -> bool:
    """Anchor a finding to the reported table row rather than the table title."""
    if row_index <= 0 or row_index > len(table.rows):
        return False
    row = table.rows[row_index - 1]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            runs = [run for run in paragraph.runs if clean_text(run.text)]
            if not runs:
                continue
            grouped = _format_comment_group(comments, anchor_context=paragraph.text)
            if grouped:
                # Keep the student's table values unchanged. The native Word
                # comment range itself identifies the affected row, while the
                # sequential number remains visible in the comment pane.
                return _add_native_comment(
                    document, runs, grouped, author=author, initials=initials
                )
    return False

def _first_native_anchor(document) -> Optional[Paragraph]:
    """Return a stable existing paragraph for document-level comments."""
    for paragraph in document.paragraphs:
        if clean_text(paragraph.text):
            return paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if clean_text(paragraph.text):
                        return paragraph
    return None


def _first_academic_anchor(document):
    """Return a safe academic-body anchor, never title-page boilerplate."""
    preferred = (
        "chapter one", "chapter two", "chapter three", "chapter four", "chapter five",
        "introduction", "background to the study", "statement of the problem",
    )
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        low = normalised(text)
        if not text:
            continue
        if low in preferred or re.match(r"^chapter\s+(one|two|three|four|five|six|[1-9])\b", low):
            return paragraph
    title_page_terms = (
        "university of cape coast", "college of", "school of", "by", "june", "july",
        "august", "september", "emmanuel", "candidate", "supervisor"
    )
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        low = normalised(text)
        if len(text.split()) < 4:
            continue
        if any(term == low or low.startswith(term) for term in title_page_terms):
            continue
        return paragraph
    return None


def _attach_document_level_comments(
    document, comments: List[str], *, author: str, initials: str
) -> None:
    """Keep unplaced findings in the Review pane without changing body text.

    Public comments must not expose provider recovery, fallback, retry or
    manual-confirmation messages. Unplaced comments are therefore exported as
    polished whole-chapter guidance without the mechanical "Document-level
    review note" prefix.
    """
    cleaned: List[str] = []
    for value in comments:
        text = public_text(
            value,
            limit=comment_max_chars(),
            reject_placeholders=True,
            reject_incomplete=True,
        )
        text = re.sub(r"^Document-level review note\.\s*", "", text, flags=re.I).strip()
        low = normalised(text)
        if not text:
            continue
        if any(token in low for token in (
            "should be checked for its contribution",
            "focused recovery",
            "separate expert review",
            "separate model response",
            "manual confirmation",
            "provider fallback",
            "independent audit",
            "recovery detail",
        )):
            continue
        cleaned.append(text)
    unique = list(dict.fromkeys(cleaned))
    if not unique:
        return
    anchor = _first_academic_anchor(document) or _first_native_anchor(document)
    if anchor is None:
        raise RuntimeError(
            "The source document has no text that can anchor native Word comments."
        )
    batches = [[comment] for comment in unique] if _export_one_comment_per_finding() else [unique[start:start + 4] for start in range(0, len(unique), 4)]
    for batch in batches:
        if not _comment_on_paragraph(
            document, anchor, batch, author=author, initials=initials
        ):
            raise RuntimeError("A native Word comment could not be anchored.")


def _preferred_evidence(row: Dict[str, Any], evidence: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Order evidence so the annotation lands on the exact claimed location."""
    target_section = normalised(row.get("section_reference") or row.get("section") or "")
    table_reference = clean_text(row.get("table_reference", ""))
    table_match = re.search(r"\bTable\s+([A-Za-z]?\d+(?:\.\d+)*)\b", table_reference, flags=re.I)
    target_table = normalised(table_match.group(1)) if table_match else ""
    quote = clean_text(row.get("problematic_quote", ""))

    def rank(item: Dict[str, Any]):
        item_section = normalised(item.get("section_reference") or item.get("heading") or "")
        item_table = normalised(item.get("table_number", ""))
        text = clean_text(item.get("text", ""))
        return (
            0 if target_table and item_table == target_table else 1,
            0 if quote and quote in text else 1,
            0 if target_section and item_section == target_section else 1,
            0 if item.get("is_heading") else 1,
            int(item.get("paragraph") or 0),
        )
    return sorted(evidence, key=rank)


def _existing_placeholder_comment(review_rows: Sequence[Dict[str, Any]], paragraph_number: int, placeholder: str) -> bool:
    """Return True when the review already contains a usable placeholder comment.

    This prevents the exporter fallback from creating a second native comment on
    the same unresolved bracketed prompt.
    """
    placeholder_key = normalised(placeholder)
    for row in review_rows:
        if row.get("status") not in ACTIONABLE_STATUSES:
            continue
        text_blob = normalised(" ".join(
            clean_text(str(row.get(field, "")))
            for field in ("item", "comment", "required_action", "illustrative_guidance", "problematic_quote")
        ))
        if not ("placeholder" in text_blob or "bracketed" in text_blob or placeholder_key in text_blob):
            continue
        evidence_paragraphs = {
            int(item.get("paragraph"))
            for item in row.get("evidence") or []
            if str(item.get("paragraph") or "").isdigit()
        }
        if paragraph_number in evidence_paragraphs or placeholder_key in text_blob:
            return True
    return False


def _placeholder_finding_rows(source_map: Dict[int, Dict[str, Any]], existing_rows: Sequence[Dict[str, Any]] = ()) -> List[Dict[str, Any]]:
    """Create deterministic review rows for unresolved bracketed prompts.

    The main review engine also checks for placeholders, but export-time
    detection prevents a missed or filtered finding from leaving obvious
    author placeholders uncommented in the final DOCX. Existing placeholder
    comments are respected so the same issue is not exported twice.
    """
    rows: List[Dict[str, Any]] = []
    seen: set[Tuple[int, str]] = set()
    for paragraph_number, locator in sorted(source_map.items()):
        text = ""
        if locator.get("kind") == "paragraph" and locator.get("paragraph") is not None:
            text = clean_text(locator["paragraph"].text)
        elif locator.get("kind") == "table_row":
            text = clean_text(" ".join(clean_text(p.text) for p in locator.get("cell_paragraphs") or []))
        if not text:
            continue
        for match in _BODY_PLACEHOLDER_RE.finditer(text):
            placeholder = clean_text(match.group(0))
            key = (paragraph_number, normalised(placeholder))
            if not placeholder or key in seen:
                continue
            if _existing_placeholder_comment(existing_rows, paragraph_number, placeholder):
                continue
            seen.add(key)
            rows.append({
                "status": STATUS_MISSING,
                "annotation_eligible": True,
                "category": "presentation",
                "severity": "major",
                "confidence": 0.99,
                "reference_label": "Delimitation of the Study" if "date" in normalised(text) or "time scope" in normalised(text) else "Drafting placeholder",
                "item": "Unresolved drafting placeholder remains in the chapter",
                "comment": "The marked bracketed text is still a drafting prompt rather than final study information. This leaves the chapter incomplete and weakens the professional readiness of the submission.",
                "required_action": "Replace the placeholder with the correct verified study detail and check the full document for any remaining bracketed prompts before resubmission.",
                "illustrative_guidance": "state the actual data-collection period once it has been confirmed, instead of leaving a prompt in the text",
                "problematic_quote": placeholder,
                "evidence": [{
                    "document_role": "current",
                    "chapter_number": locator.get("chapter_number"),
                    "paragraph": paragraph_number,
                    "text": text,
                }],
            })
    return rows


def synchronise_export_fallback_findings(
    review: Dict[str, Any],
    source_map: Dict[int, Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """Add deterministic export-time findings to the canonical review once.

    The report is generated after the annotated DOCX during the normal job
    workflow. Persisting any export-time placeholder finding in the review
    therefore keeps the report, native comments, inline annotations and final
    correction numbering on the same evidence ledger.
    """
    academic_rows = list(review.get("academic_findings") or [])
    fallback_rows = _placeholder_finding_rows(source_map, academic_rows)
    if fallback_rows:
        for index, row in enumerate(fallback_rows, start=1):
            evidence = (row.get("evidence") or [{}])[0]
            row.setdefault(
                "finding_id",
                f"EXPORT-PLACEHOLDER-{evidence.get('chapter_number') or 0}-{evidence.get('paragraph') or index}-{index}",
            )
            academic_rows.append(row)
        review["academic_findings"] = academic_rows
        # A previously cached professional package was built before these
        # deterministic findings existed. Force a rebuild when the report is
        # exported so every output uses the same sequence.
        review.pop("professional_review", None)
    return (
        list(review.get("academic_findings") or [])
        + list(review.get("alignment_results") or [])
        + list(review.get("revision_results") or [])
    )


def native_comment_count(docx_bytes: bytes) -> int:
    """Return the number of native Word comments in an exported DOCX."""
    try:
        document = Document(io.BytesIO(docx_bytes))
        return len(list(document.comments))
    except Exception:
        return 0


def _canonical_group_label(row: Dict[str, Any]) -> str:
    raw = clean_text(
        row.get("reference_label")
        or row.get("section_reference")
        or row.get("section")
        or ((row.get("headings") or [""])[-1] if isinstance(row.get("headings"), list) else "")
        or "Chapter-level review"
    )
    low = normalised(raw)
    if "general objective" in low or "specific objective" in low or "research objectives" in low:
        return "Research Objectives"
    if "research question" in low or "hypoth" in low:
        return "Research Questions"
    if "background" in low:
        return "Background of the Study"
    if "problem" in low:
        return "Problem Statement"
    if "scope" in low or "delimitation" in low:
        return "Scope of the Study"
    if "significance" in low or "contribution" in low:
        return "Significance of the Study"
    if "limitation" in low:
        return "Limitations of the Study"
    if "organisation" in low or "organization" in low:
        return "Organization of the Study"
    if "definition" in low or "terms" in low or "construct" in low:
        return "Definition of Terms"
    if "citation" in low or "source" in low or "reference" in low:
        # Keep citation-only issues close to the location that supplied the evidence
        # where possible. If the row does not specify a section, it remains a
        # chapter-level source note.
        section = clean_text(row.get("section_reference") or row.get("section") or "")
        return section if section and not re.search(r"citation|source|reference", section, flags=re.I) else "Source and Reference Integrity"
    return raw or "Chapter-level review"


def _group_headings_for_row(row: Dict[str, Any], label: str) -> Tuple[str, ...]:
    headings = [clean_text(value) for value in (row.get("headings") or []) if clean_text(value)]
    # Prefer an actual document heading over an issue category.
    if headings:
        last = headings[-1]
        if normalised(last) not in {"citations and sources", "cross section coherence", "conceptual clarity", "critical analysis"}:
            if label and normalised(label) not in normalised(last):
                return tuple(headings + [label])
            return tuple(headings)
    if label == "Source and Reference Integrity":
        return ("CHAPTER ONE",)
    return (label,)


def _row_group_key(row: Dict[str, Any], evidence: Sequence[Dict[str, Any]]) -> Tuple[Optional[int], Tuple[str, ...], int]:
    label = _canonical_group_label(row)
    headings = _group_headings_for_row(row, label)
    chapter_number = row.get("chapter_number")
    if chapter_number is None:
        chapter_number = next((item.get("chapter_number") for item in evidence if item.get("chapter_number") is not None), None)
    try:
        chapter_number = int(chapter_number) if chapter_number is not None else None
    except (TypeError, ValueError):
        chapter_number = None
    paragraph_number = 0
    if not headings or headings == ("Chapter-level review",):
        paragraph_number = next((int(item.get("paragraph")) for item in evidence if str(item.get("paragraph") or "").isdigit()), 0)
    return chapter_number, headings, paragraph_number


def _paragraph_text_from_locator(locator: Dict[str, Any]) -> str:
    paragraph = locator.get("paragraph")
    if paragraph is not None:
        return clean_text(paragraph.text)
    if locator.get("kind") == "table_row":
        values: List[str] = []
        for paragraph in locator.get("cell_paragraphs") or []:
            values.append(clean_text(paragraph.text))
        return clean_text(" ".join(values))
    return ""


def _paragraph_looks_like_heading(locator: Dict[str, Any]) -> bool:
    paragraph = locator.get("paragraph")
    text = _paragraph_text_from_locator(locator)
    if not text:
        return False
    style_name = ""
    if paragraph is not None:
        try:
            style_name = (paragraph.style.name or "").lower()
        except Exception:
            style_name = ""
    if "heading" in style_name or "title" in style_name:
        return True
    if len(text.split()) <= 9 and re.match(r"^(?:chapter\s+\w+|\d+(?:\.\d+){0,4}\.?\s+)", text, flags=re.I):
        return True
    return False


def _next_substantive_paragraph_number(
    source_map: Dict[int, Dict[str, Any]],
    paragraph_number: int,
    *,
    max_ahead: int = 5,
) -> int:
    for number in range(paragraph_number + 1, paragraph_number + max_ahead + 1):
        locator = source_map.get(number) or {}
        text = _paragraph_text_from_locator(locator)
        if len(text.split()) >= 8 and not _paragraph_looks_like_heading(locator):
            return number
    return paragraph_number


def _first_paragraph_matching(
    source_map: Dict[int, Dict[str, Any]],
    patterns: Sequence[str],
) -> int:
    compiled = [re.compile(pattern, flags=re.I) for pattern in patterns]
    for number in sorted(source_map):
        text = _paragraph_text_from_locator(source_map[number])
        if text and any(pattern.search(text) for pattern in compiled):
            return number
    return 0


def _insertion_anchor_for_unanchored_row(
    row: Dict[str, Any],
    source_map: Dict[int, Dict[str, Any]],
) -> int:
    """Locate a neutral insertion point for genuinely absent material.

    The locator uses structural headings and wording from the current document
    only. It must never import names, variables or organisations from a prior
    submission. Missing sections remain listed at the chapter bottom.
    """
    haystack = normalised(" ".join([
        row.get("issue_title", ""), row.get("item", ""), row.get("section", ""),
        row.get("section_reference", ""), row.get("required_action", ""),
        row.get("comment", ""), row.get("assessment", ""),
    ]))
    structural_patterns = {
        "definition": [r"definition\s+of\s+(?:key\s+)?(?:terms|concepts)", r"operational\s+definition"],
        "references": [r"^\s*references\s*$", r"^\s*bibliography\s*$"],
        "gap": [r"research\s+gap", r"statement\s+of\s+the\s+problem", r"problem\s+statement"],
        "objective": [r"research\s+objectives?", r"specific\s+objectives?", r"general\s+objective"],
        "question": [r"research\s+questions?"],
        "scope": [r"scope\s+of\s+the\s+study", r"delimitations?\s+of\s+the\s+study"],
        "significance": [r"significance\s+of\s+the\s+study", r"contribution\s+to\s+knowledge"],
        "limitation": [r"limitations?\s+of\s+the\s+study"],
        "method": [r"research\s+methodology", r"research\s+methods", r"data\s+analysis"],
        "results": [r"results", r"findings", r"discussion"],
    }
    if "definition of terms" in haystack or "operational definition" in haystack:
        anchor = _first_paragraph_matching(source_map, structural_patterns["scope"] + structural_patterns["significance"])
        return anchor or _first_paragraph_matching(source_map, structural_patterns["definition"])
    if "reference list" in haystack or "references" in haystack or "bibliography" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["references"])
    if "research gap" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["gap"])
    if "objective" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["objective"])
    if "research question" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["question"])
    if "scope" in haystack or "delimitation" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["scope"])
    if "significance" in haystack or "contribution" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["significance"])
    if "limitation" in haystack or "generalisation" in haystack or "generalization" in haystack:
        return _first_paragraph_matching(source_map, structural_patterns["limitation"])
    if any(term in haystack for term in ("method", "sampling", "instrument", "validity", "reliability")):
        return _first_paragraph_matching(source_map, structural_patterns["method"])
    if any(term in haystack for term in ("result", "analysis", "discussion", "statistical")):
        return _first_paragraph_matching(source_map, structural_patterns["results"])
    if "citation" in haystack or "source attribution" in haystack:
        return _first_paragraph_matching(source_map, [r"\([^)]*(?:19|20)\d{2}[^)]*\)"])
    return 0


def _better_evidence_paragraph_number(
    row: Dict[str, Any],
    source_map: Dict[int, Dict[str, Any]],
    paragraph_number: int,
) -> int:
    if paragraph_number <= 0:
        return _insertion_anchor_for_unanchored_row(row, source_map)
    locator = source_map.get(paragraph_number) or {}
    if _paragraph_looks_like_heading(locator):
        replacement = _insertion_anchor_for_unanchored_row(row, source_map)
        if replacement and replacement != paragraph_number:
            return replacement
        return _next_substantive_paragraph_number(source_map, paragraph_number)
    return paragraph_number


def _row_span_for_paragraph(row: Dict[str, Any], paragraph_text: str) -> Tuple[int, int]:
    quote = clean_text(row.get("problematic_quote", ""))
    if quote:
        exact_start = paragraph_text.find(quote)
        if exact_start >= 0:
            return _expand_to_safe_text_span(paragraph_text, exact_start, exact_start + len(quote))
    terms = [
        row.get("issue_title", ""), row.get("item", ""), row.get("required_action", ""),
        row.get("comment", ""), row.get("assessment", ""), row.get("section", ""),
        row.get("section_reference", ""),
    ]
    return _best_span(paragraph_text, terms, quote)




def _specific_correction_text(row: Dict[str, Any], comment: str) -> str:
    """Return a detailed blue correction that mirrors the academic finding."""
    label = _canonical_group_label(row)
    issue = _sanitise_guidance(row.get("item", "") or row.get("issue_title", ""))
    assessment = _sanitise_guidance(row.get("comment", "") or row.get("assessment", ""))
    consequence = _sanitise_guidance(
        row.get("academic_consequence", "") or row.get("consequence", "") or row.get("why_it_matters", "")
    )
    action = _sanitise_guidance(row.get("required_action", ""))
    if not action:
        action = _strip_comment_reference(comment)
    example = _sanitise_guidance(row.get("illustrative_guidance", "")) or _sanitise_guidance(context_specific_example(row))
    example = re.sub(r"^(?:for\s+)?(?:context\s+)?example[:,]?\s*", "", example, flags=re.I).strip(" .")

    parts: List[str] = []
    if label and issue:
        parts.append(f"{label}: {issue.rstrip(' .')}.")
    elif issue:
        parts.append(issue.rstrip(" .") + ".")
    elif label:
        parts.append(f"{label} requires revision.")
    if assessment and normalised(assessment) not in normalised(" ".join(parts)):
        parts.append(assessment.rstrip(" .") + ".")
    if consequence and normalised(consequence) not in normalised(" ".join(parts)):
        parts.append(consequence.rstrip(" .") + ".")

    level = academic_level_label(row.get("_academic_level") or row.get("academic_level"))
    explicit_level_sentence = re.search(
        r"(?:^|[.!?]\s+)At\s+(?:PhD|MPhil|professional doctorate|Master's|non-research Master's|Bachelor's)\s+level\b",
        " ".join(parts),
    )
    if level != "the applicable academic level" and not explicit_level_sentence:
        parts.append(f"At {level}, the correction should demonstrate clear scholarly judgement, methodological or analytical defensibility, and traceable support from the study evidence.")

    if action:
        parts.append(_normalise_action_start(action).rstrip(" .") + ".")
    else:
        parts.append("Revise the marked passage so that the claim is clear, evidence-supported and aligned with the chapter purpose.")
    if example:
        example = _normalise_action_start(example).rstrip(" .")
        parts.append("For example, " + example[0].lower() + example[1:] + ".")

    text = professionalise_reviewer_language(" ".join(parts), row.get("_academic_level") or row.get("academic_level"))
    return public_text(_shorten_comment(text, 2000), reject_placeholders=True, reject_incomplete=True)


def _add_specific_corrections_required(
    document,
    numbered_rows: Sequence[Tuple[int, Dict[str, Any], str]],
) -> None:
    """Append a blue end-of-chapter checklist that matches red body numbers."""
    if not _specific_corrections_required_enabled() or not numbered_rows:
        return
    anchor = _last_chapter_body_paragraph(document)
    if anchor is None:
        return
    entries: List[Tuple[int, str]] = []
    seen = set()
    for number, row, comment in numbered_rows:
        text = _specific_correction_text(row, comment)
        key = (number, normalised(text))
        if not text or key in seen:
            continue
        seen.add(key)
        entries.append((number, text))
    if not entries:
        return
    entries.sort(key=lambda item: item[0])
    for number, text in reversed(entries):
        note = _insert_blue_paragraph_after(anchor)
        run = note.add_run(f"{number}. {text}")
        run.font.color.rgb = INLINE_BLUE
        run.font.italic = True
        try:
            note.paragraph_format.left_indent = anchor.paragraph_format.left_indent
            note.paragraph_format.space_before = anchor.paragraph_format.space_after
            note.paragraph_format.space_after = anchor.paragraph_format.space_after
        except Exception:
            pass
    heading = _insert_blue_paragraph_after(anchor)
    lead = heading.add_run("Specific corrections required")
    lead.bold = True
    lead.font.color.rgb = INLINE_BLUE
    lead.font.italic = True

def _build_grouped_annotated_docx(
    document,
    source_map: Dict[int, Dict[str, Any]],
    table_map: Dict[int, Dict[str, Any]],
    review_rows: Sequence[Dict[str, Any]],
    review: Dict[str, Any],
    *,
    author: str,
    initials: str,
) -> bytes:
    """Build professional grouped native comments anchored to the exact evidence.

    Earlier grouped exports placed one numbered comment on the section heading.
    That looked tidy, but it was difficult for a student to see the sentence or
    paragraph that needed revision. This builder groups related findings only
    when they share the same evidence passage, then anchors the numbered comment
    to the exact quote, best sentence, paragraph, or table row. Missing-section
    findings are placed on the nearest insertion point rather than on a chapter
    heading.
    """
    academic_level = (review.get("summary") or {}).get("academic_level")
    prepared_rows = []
    for source_row in review_rows:
        row = dict(source_row)
        row["_academic_level"] = academic_level
        prepared_rows.append(row)
    review_rows = order_and_number_rows(prepared_rows)

    by_paragraph: Dict[int, Dict[Tuple[int, int], List[str]]] = defaultdict(lambda: defaultdict(list))
    after_paragraph: Dict[int, List[str]] = defaultdict(list)
    by_table: Dict[Tuple[int, int], List[str]] = defaultdict(list)
    fallback_comments: List[str] = []
    missing_section_rows: List[Dict[str, Any]] = []
    numbered_rows: List[Tuple[int, Dict[str, Any], str]] = []
    def reference_number_for(row: Dict[str, Any]) -> int:
        try:
            return int(row.get("finding_number"))
        except (TypeError, ValueError):
            return 0

    for row in review_rows:
        if row.get("status") not in ACTIONABLE_STATUSES:
            continue
        if row.get("annotation_eligible") is False:
            continue
        if _is_missing_section_finding(row):
            raw_missing_comment = _missing_section_bottom_comment(row) or _comment_body(row)
            if raw_missing_comment:
                reference_number = reference_number_for(row)
                numbered_rows.append((reference_number, row, raw_missing_comment))
            missing_section_rows.append(row)
            continue
        raw_comment = _comment_body(row)
        if not raw_comment:
            continue
        reference_number = reference_number_for(row)
        numbered_rows.append((reference_number, row, raw_comment))
        comment = _with_comment_reference(reference_number, raw_comment)
        evidence = [
            item for item in (row.get("evidence") or [])
            if item.get("document_role", "current") == "current"
        ]
        evidence = _preferred_evidence(row, evidence)
        paragraph_number = 0
        best: Dict[str, Any] = evidence[0] if evidence else {}
        if best:
            try:
                paragraph_number = int(best.get("paragraph"))
            except (TypeError, ValueError):
                paragraph_number = 0
        paragraph_number = _better_evidence_paragraph_number(row, source_map, paragraph_number)
        locator = source_map.get(paragraph_number) if paragraph_number else None

        if locator is not None:
            if locator.get("kind") in {"table_row", "table_caption"} or best.get("table_index"):
                try:
                    table_index = int(locator.get("table_index") or best.get("table_index") or 0)
                except (TypeError, ValueError):
                    table_index = 0
                if table_index:
                    try:
                        table_row = int(locator.get("table_row") or best.get("table_row") or 0)
                    except (TypeError, ValueError):
                        table_row = 0
                    by_table[(table_index, table_row)].append(comment)
                    continue
            paragraph = locator.get("paragraph")
            paragraph_text = paragraph.text if paragraph is not None else ""
            if paragraph is not None and paragraph_text:
                start, end = _row_span_for_paragraph(row, paragraph_text)
                if start < end:
                    by_paragraph[paragraph_number][(start, end)].append(comment)
                else:
                    after_paragraph[paragraph_number].append(comment)
                continue

        insertion = _insertion_anchor_for_unanchored_row(row, source_map)
        if insertion and insertion in source_map:
            after_paragraph[insertion].append(comment)
        else:
            fallback_comments.append(comment)

    for paragraph_number, span_groups in by_paragraph.items():
        locator = source_map.get(paragraph_number) or {}
        paragraph = locator.get("paragraph")
        if paragraph is None:
            continue
        # Keep sentence-level references precise. Only overlapping spans are
        # grouped, so each red number remains attached to the exact sentence or
        # phrase it refers to rather than drifting to the end of a broad passage.
        merged_groups = _merge_nearby_span_groups(span_groups, max_gap=0)
        for (start, end), comments in reversed(merged_groups):
            combined = _format_comment_group(comments, anchor_context=(paragraph.text or "")[start:end])
            if not combined:
                continue
            if not _mark_span_and_insert_comment(
                document, paragraph, start, end, combined,
                author=author, initials=initials,
            ):
                after_paragraph[paragraph_number].extend(comments)

    for paragraph_number, comments in after_paragraph.items():
        locator = source_map.get(paragraph_number) or {}
        paragraph = locator.get("paragraph")
        if paragraph is not None:
            unique = list(dict.fromkeys(comments))
            if not _comment_on_paragraph(document, paragraph, unique, author=author, initials=initials):
                fallback_comments.extend(unique)
        else:
            fallback_comments.extend(comments)

    for table_key in sorted(by_table, reverse=True):
        table_index, table_row = table_key
        table_info = table_map.get(table_index) or {}
        table = table_info.get("table")
        comments = list(dict.fromkeys(by_table[table_key]))
        caption = table_info.get("caption_paragraph")
        added = False
        if table is not None and table_row:
            added = _comment_on_table_row(
                document, table, table_row, comments, author=author, initials=initials
            )
        if not added and table is not None:
            added = _comment_on_table(document, table, comments, caption=caption, author=author, initials=initials)
        if not added and caption is not None:
            added = _comment_on_paragraph(document, caption, comments, author=author, initials=initials)
        if not added:
            fallback_comments.extend(comments)

    if _include_section_review_comments():
        _add_section_review_comments(
            document,
            review,
            author=author,
            initials=initials,
            fallback_comments=fallback_comments,
        )

    _attach_document_level_comments(
        document,
        list(dict.fromkeys([comment for comment in fallback_comments if comment])),
        author=author,
        initials=initials,
    )
    if _specific_corrections_required_enabled():
        _add_specific_corrections_required(document, numbered_rows)
    else:
        _add_missing_section_inline_bottom_notes(document, missing_section_rows)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()

def build_annotated_docx(
    source_bytes: bytes,
    review: Dict[str, Any],
    comment_author: Optional[str] = None,
) -> bytes:
    document = Document(io.BytesIO(source_bytes))
    author, initials = _comment_identity(review, comment_author)
    source_map, table_map = _source_locator_map(document)

    # All annotations are native Word comments. Exact quotations are used as
    # comment anchors where available. The source text, formatting, pagination,
    # tables, and headings remain unchanged.
    by_paragraph: Dict[int, Dict[Tuple[int, int], List[str]]] = defaultdict(lambda: defaultdict(list))
    after_paragraph: Dict[int, List[str]] = defaultdict(list)
    by_table: Dict[int, List[str]] = defaultdict(list)
    missing_by_heading: Dict[Tuple[Optional[int], Tuple[str, ...]], List[str]] = defaultdict(list)
    fallback_comments: List[str] = []
    missing_section_rows: List[Dict[str, Any]] = []

    supplied_rows = synchronise_export_fallback_findings(review, source_map)
    review_rows = _sanitise_rows_for_export(supplied_rows)
    academic_level = (review.get("summary") or {}).get("academic_level")
    review_rows = order_and_number_rows([
        {**row, "_academic_level": academic_level} for row in review_rows
    ])
    if _merge_comments_by_section():
        return _build_grouped_annotated_docx(
            document,
            source_map,
            table_map,
            review_rows,
            review,
            author=author,
            initials=initials,
        )

    for row in review_rows:
        if row.get("status") not in ACTIONABLE_STATUSES:
            continue
        if row.get("annotation_eligible") is False:
            continue
        if _is_missing_section_finding(row):
            missing_section_rows.append(row)
            continue

        comment = _comment_body(row)
        if not comment:
            continue
        evidence = [
            item for item in (row.get("evidence") or [])
            if item.get("document_role", "current") == "current"
        ]
        evidence = _preferred_evidence(row, evidence)
        if evidence:
            best = evidence[0]
            try:
                paragraph_number = int(best.get("paragraph"))
            except (TypeError, ValueError):
                paragraph_number = 0
            locator = source_map.get(paragraph_number)
            if locator is not None:
                if locator.get("kind") in {"table_row", "table_caption"} or best.get("table_index"):
                    table_index = int(locator.get("table_index") or best.get("table_index") or 0)
                    if table_index:
                        by_table[table_index].append(comment)
                        continue
                paragraph = locator.get("paragraph")
                if paragraph is not None:
                    quote = clean_text(row.get("problematic_quote", ""))
                    exact_start = paragraph.text.find(quote) if quote else -1
                    if exact_start >= 0:
                        by_paragraph[paragraph_number][
                            (exact_start, exact_start + len(quote))
                        ].append(comment)
                    else:
                        after_paragraph[paragraph_number].append(comment)
                    continue

        headings = tuple(row.get("headings") or [row.get("section_reference") or row.get("section")])
        headings = tuple(value for value in headings if clean_text(value))
        if headings:
            chapter_number = row.get("chapter_number")
            if chapter_number is None:
                chapter_number = next(
                    (item.get("chapter_number") for item in evidence if item.get("chapter_number") is not None),
                    None,
                )
            try:
                chapter_number = int(chapter_number) if chapter_number is not None else None
            except (TypeError, ValueError):
                chapter_number = None
            missing_by_heading[(chapter_number, headings)].append(comment)
        else:
            fallback_comments.append(comment)

    for paragraph_number, span_groups in by_paragraph.items():
        locator = source_map.get(paragraph_number) or {}
        paragraph = locator.get("paragraph")
        if paragraph is None:
            continue
        merged_groups = _merge_nearby_span_groups(span_groups)
        for (start, end), comments in reversed(merged_groups):
            if _export_one_comment_per_finding():
                placed_any = False
                for comment in _prepare_comment_list(comments):
                    if _mark_span_and_insert_comment(
                        document, paragraph, start, end, comment,
                        author=author, initials=initials,
                    ):
                        placed_any = True
                    else:
                        after_paragraph[paragraph_number].append(comment)
                if not placed_any and comments:
                    after_paragraph[paragraph_number].extend(comments)
                continue
            combined = _format_comment_group(comments, anchor_context=(paragraph.text or "")[start:end])
            if not _mark_span_and_insert_comment(
                document, paragraph, start, end, combined,
                author=author, initials=initials,
            ):
                after_paragraph[paragraph_number].extend(comments)

    for paragraph_number, comments in after_paragraph.items():
        locator = source_map.get(paragraph_number) or {}
        paragraph = locator.get("paragraph")
        if paragraph is not None:
            unique = list(dict.fromkeys(comments))
            if not _comment_on_paragraph(
                document, paragraph, unique, author=author, initials=initials
            ):
                fallback_comments.extend(unique)
        else:
            fallback_comments.extend(comments)

    # Process tables in reverse order for stable comment anchoring.
    for table_index in sorted(by_table, reverse=True):
        table_info = table_map.get(table_index) or {}
        table = table_info.get("table")
        comments = list(dict.fromkeys(by_table[table_index]))
        caption = table_info.get("caption_paragraph")
        if table is not None:
            if not _comment_on_table(
                document, table, comments, caption=caption,
                author=author, initials=initials,
            ):
                fallback_comments.extend(comments)
        elif caption is not None:
            if not _comment_on_paragraph(
                document, caption, comments, author=author, initials=initials
            ):
                fallback_comments.extend(comments)
        else:
            fallback_comments.extend(comments)

    # Findings about absent or underdeveloped material are placed under the
    # exact supplied section or subsection heading.
    for (chapter_number, headings), comments in missing_by_heading.items():
        unique_comments = list(dict.fromkeys(comments))
        heading = _find_heading(document, headings, chapter_number=chapter_number)
        if heading is not None:
            if not _comment_on_paragraph(
                document, heading, unique_comments,
                author=author, initials=initials,
            ):
                fallback_comments.extend(unique_comments)
        else:
            fallback_comments.extend(unique_comments)

    # Add a section-level review note for every reviewed section or subsection.
    # These comments are separate from issue findings and make coverage visible
    # in the native Word Review pane. They are added after issue comments so
    # exact issue anchors remain the primary feedback.
    _add_section_review_comments(
        document,
        review,
        author=author,
        initials=initials,
        fallback_comments=fallback_comments,
    )

    _attach_document_level_comments(
        document,
        list(dict.fromkeys(fallback_comments)),
        author=author,
        initials=initials,
    )
    _add_missing_section_inline_bottom_notes(document, missing_section_rows)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
