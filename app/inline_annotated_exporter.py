from __future__ import annotations

import io
import re
from collections import defaultdict
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .annotated_exporter import (
    ACTIONABLE_STATUSES,
    _best_span,
    _expand_to_safe_text_span,
    _better_evidence_paragraph_number,
    _comment_body,
    _format_comment_group,
    _placeholder_finding_rows,
    synchronise_export_fallback_findings,
    _preferred_evidence,
    _run_element,
    _source_locator_map,
    _add_missing_section_inline_bottom_notes,
    _is_missing_section_finding,
    _RICH_RED_RE,
    COMMENT_RED,
    _with_comment_reference,
    _group_reference_numbers_from_comment,
    _specific_corrections_required_enabled,
    _add_specific_corrections_required,
    _specific_correction_text,
    _visible_runs,
    _visible_paragraph_text,
)
from .comment_quality import public_text
from .final_review_quality import build_canonical_finding_rows
from .reviewer_language import professionalise_reviewer_language
from .natural_supervisor_comment import natural_supervisor_comment
from .document_parser import clean_text, normalised

INLINE_ANNOTATION_EXPORT_VERSION = "2.7.1-atomic-inline-annotated-recovery"
PROFESSIONAL_INLINE_PRODUCT_VERSION = "2.7.1-atomic-inline-annotated-recovery"
REVISION_RED = "C00000"
COMMENT_BLUE = RGBColor(0x00, 0x70, 0xC0)


_PROHIBITED_PUBLIC_RE = re.compile(
    r"\b(?:manifest|document map|paragraph id|section packet|parser|fallback|recovery|focused recovery|model response|P\d{1,4})\b",
    flags=re.I,
)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _red_run_element(text: str, source_run: Optional[Run] = None):
    return _run_element(text, source_run=source_run, colour=REVISION_RED)


def _mark_span_red(paragraph: Paragraph, start: int, end: int, reference_numbers: Sequence[int] = ()) -> bool:
    if start < 0 or end <= start:
        return False
    start, end = _expand_to_safe_text_span(_visible_paragraph_text(paragraph), start, end)
    if start < 0 or end <= start:
        return False
    runs = _visible_runs(paragraph)
    cursor = 0
    changed = False
    trailing_element = None
    trailing_source_run = None
    for run in runs:
        text = run.text or ""
        run_start, run_end = cursor, cursor + len(text)
        cursor = run_end
        if not text or run_end <= start or run_start >= end:
            continue
        local_start = max(0, start - run_start)
        local_end = min(len(text), end - run_start)
        before, marked, after = text[:local_start], text[local_start:local_end], text[local_end:]
        parent = run._r.getparent()
        index = parent.index(run._r)
        for value, red in ((before, False), (marked, True), (after, False)):
            if not value:
                continue
            element = _red_run_element(value, run) if red else _run_element(value, source_run=run)
            parent.insert(index, element)
            if red:
                trailing_element = element
                trailing_source_run = run
            index += 1
        parent.remove(run._r)
        changed = True
    if changed and reference_numbers and trailing_element is not None:
        marker_text = " " + " ".join(f"[{number}]" for number in reference_numbers)
        marker = _red_run_element(marker_text, trailing_source_run)
        parent = trailing_element.getparent()
        if parent is not None:
            parent.insert(parent.index(trailing_element) + 1, marker)
    return changed


def _write_inline_comment_body(note: Paragraph, body: str) -> None:
    """Write grouped inline comments without exposing internal rich-text tokens."""
    pos = 0
    for match in _RICH_RED_RE.finditer(body):
        before = body[pos:match.start()]
        if before:
            run = note.add_run(before)
            run.font.color.rgb = COMMENT_BLUE
            run.font.italic = True
        marker = match.group(1)
        if marker:
            run = note.add_run(marker)
            run.font.color.rgb = COMMENT_RED
            run.bold = True
            run.font.italic = True
        pos = match.end()
    tail = body[pos:]
    if tail:
        run = note.add_run(tail)
        run.font.color.rgb = COMMENT_BLUE
        run.font.italic = True


def _add_inline_comment(paragraph: Paragraph, comments: Sequence[str]) -> None:
    if len(comments) == 1:
        value = comments[0]
        match = re.match(r"\[\[VPROF_REFNO:(\d+)\]\](.*)", value, flags=re.S)
        if match:
            body = f"[[VPROF_RED:{match.group(1)}. ]]" + match.group(2).strip()
        else:
            body = value
    else:
        body = _format_comment_group(comments, anchor_context=_visible_paragraph_text(paragraph))
    plain_body = _RICH_RED_RE.sub(lambda match: match.group(1), body)
    released = public_text(
        plain_body,
        limit=2200,
        reject_placeholders=True,
        reject_incomplete=True,
    )
    # ``body`` is assembled only from the validated canonical ledger. Do not
    # silently drop a numbered inline note when a quoted source fragment is
    # judged incomplete at this final presentation boundary.
    if not released:
        released = public_text(
            plain_body,
            limit=2200,
            reject_placeholders=True,
            reject_incomplete=False,
        )
    if not body or not released or _PROHIBITED_PUBLIC_RE.search(released):
        return
    note = _insert_paragraph_after(paragraph)
    lead = note.add_run("Detailed supervisor comment: ")
    lead.bold = True
    lead.font.color.rgb = COMMENT_BLUE
    lead.font.italic = True
    _write_inline_comment_body(note, body)
    try:
        note.paragraph_format.space_before = paragraph.paragraph_format.space_after
        note.paragraph_format.space_after = paragraph.paragraph_format.space_after
        note.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
    except Exception:
        pass


def _clean_comment(value: str) -> str:
    text = public_text(value, limit=2200, reject_placeholders=True, reject_incomplete=True)
    if not text or _PROHIBITED_PUBLIC_RE.search(text):
        return ""
    return text


def _inline_comment_limit() -> int:
    try:
        return max(260, min(720, int(__import__("os").getenv("VPROF_INLINE_COMMENT_MAX_CHARS", "480"))))
    except (TypeError, ValueError):
        return 480


def _row_comment(row: Dict[str, Any]) -> str:
    """Build a concise, natural inline note from the canonical finding record."""
    text = natural_supervisor_comment(
        row,
        compact=True,
        include_reason=False,
        include_verification=False,
        include_example=False,
    )
    if text:
        text = professionalise_reviewer_language(text, row.get("_academic_level") or row.get("academic_level"))
        released = _clean_comment(public_text(text, limit=_inline_comment_limit(), reject_placeholders=True, reject_incomplete=True))
        if released:
            return released
    fallback = clean_text(
        row.get("student_comment")
        or " ".join(value for value in (
            clean_text(row.get("item") or row.get("issue_title")),
            clean_text(row.get("required_action")),
        ) if value)
    )
    return public_text(fallback, limit=_inline_comment_limit(), reject_placeholders=True, reject_incomplete=False)


def _rows_for_inline(review: Dict[str, Any], source_map: Dict[int, Dict[str, Any]]) -> List[Dict[str, Any]]:
    # The exporter must use the same final, filtered and sequentially numbered
    # findings as the report and native-comment version. Export-time fallback
    # findings are synchronised before this function is called.
    return build_canonical_finding_rows(review)


def build_inline_annotated_docx(
    source_bytes: bytes,
    review: Dict[str, Any],
    comment_author: Optional[str] = None,
) -> bytes:
    """Create a non-native annotated DOCX.

    Areas needing revision are coloured red in the document body. The supervisor's
    explanatory comment is inserted immediately after the affected paragraph in
    blue italic text. No Word comments.xml or native comment boxes are used, so
    users can compare this output with the native-comment version.
    """
    document = Document(io.BytesIO(source_bytes))
    source_map, _ = _source_locator_map(document)
    academic_level = (review.get("summary") or {}).get("academic_level")
    synchronise_export_fallback_findings(review, source_map)
    # Rebuild after export fallbacks have been synchronised. Numbering is then
    # fixed once and reused everywhere.
    review_rows = [
        {**row, "_academic_level": academic_level}
        for row in build_canonical_finding_rows(review, force=bool(review.pop("_export_fallback_added", False)))
    ]

    # All findings tied to the same paragraph share one numbered inline note.
    # This mirrors the native Word comment grouping and keeps the annotated
    # chapter readable while preserving every canonical finding number.
    after_anchor: Dict[int, Dict[str, Any]] = defaultdict(lambda: {"spans": [], "comments": []})
    missing_section_rows: List[Dict[str, Any]] = []
    numbered_rows: List[Tuple[int, Dict[str, Any], str]] = []
    def reference_number_for(row: Dict[str, Any]) -> int:
        try:
            return int(row.get("finding_number"))
        except (TypeError, ValueError):
            return 0

    for row in review_rows:
        if row.get("status") not in ACTIONABLE_STATUSES:
            continue
        if _is_missing_section_finding(row):
            raw_missing_comment = _row_comment(row)
            if raw_missing_comment:
                reference_number = reference_number_for(row)
                numbered_rows.append((reference_number, row, raw_missing_comment))
            missing_section_rows.append(row)
            continue
        if row.get("annotation_eligible") is False:
            continue
        raw_comment = _row_comment(row)
        if not raw_comment:
            continue
        reference_number = reference_number_for(row)
        numbered_rows.append((reference_number, row, raw_comment))
        comment = _with_comment_reference(reference_number, raw_comment)
        evidence = [
            item for item in (row.get("evidence") or [])
            if item.get("document_role", "current") == "current"
        ]
        evidence = _preferred_evidence(row, evidence)
        best = evidence[0] if evidence else {}
        try:
            paragraph_number = int(best.get("paragraph")) if best else 0
        except (TypeError, ValueError):
            paragraph_number = 0
        paragraph_number = _better_evidence_paragraph_number(row, source_map, paragraph_number)
        locator = source_map.get(paragraph_number)
        if not locator or locator.get("kind") not in {"paragraph", "table_caption"}:
            continue
        paragraph = locator.get("paragraph")
        if paragraph is None:
            continue
        quote = clean_text(row.get("problematic_quote", ""))
        text = _visible_paragraph_text(paragraph)
        if quote and quote in text:
            quote_start = text.find(quote)
            start, end = _expand_to_safe_text_span(text, quote_start, quote_start + len(quote))
        else:
            terms = [row.get("issue_title", ""), row.get("item", ""), row.get("section", "")]
            start, end = _best_span(text, terms, quote)
        start, end = _expand_to_safe_text_span(text, start, end)
        if start >= end:
            start, end = 0, len(text)
        after_anchor[paragraph_number]["spans"].append((start, end))
        after_anchor[paragraph_number]["comments"].append(comment)

    # Missing-section findings must also be visible in the inline annotated
    # document. They cannot be highlighted inside absent text, so add a numbered
    # inline note at the end of the submitted chapter before the correction register.
    _add_missing_section_inline_bottom_notes(document, missing_section_rows)
    if _specific_corrections_required_enabled():
        _add_specific_corrections_required(document, numbered_rows)

    for paragraph_number, group in sorted(after_anchor.items(), reverse=True):
        locator = source_map.get(paragraph_number) or {}
        paragraph = locator.get("paragraph")
        if paragraph is None:
            continue
        comments = list(dict.fromkeys(group.get("comments") or []))
        spans = list(group.get("spans") or [])
        if not comments or not spans:
            continue
        text = _visible_paragraph_text(paragraph)
        start = min(item[0] for item in spans)
        end = max(item[1] for item in spans)
        start, end = _expand_to_safe_text_span(text, start, end)
        grouped = _format_comment_group(comments, anchor_context=text[start:end])
        _mark_span_red(
            paragraph,
            start,
            end,
            _group_reference_numbers_from_comment(grouped),
        )
        _add_inline_comment(paragraph, comments)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def inline_annotation_count(source_bytes: bytes, review: Dict[str, Any]) -> int:
    document = Document(io.BytesIO(source_bytes))
    source_map, _ = _source_locator_map(document)
    count = 0
    for row in _rows_for_inline(review, source_map):
        if row.get("status") in ACTIONABLE_STATUSES and row.get("annotation_eligible") is not False and _row_comment(row):
            count += 1
    return count


def inline_annotation_audit(docx_bytes: bytes, review: Dict[str, Any]) -> Dict[str, Any]:
    """Verify that every actionable finding number appears in an inline note."""
    from .annotated_exporter import expected_annotation_finding_numbers

    expected = expected_annotation_finding_numbers(review)
    expected_set = set(expected)
    represented: set[int] = set()
    note_count = 0
    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception:
        return {
            "note_count": 0,
            "expected_finding_numbers": expected,
            "represented_finding_numbers": [],
            "missing_finding_numbers": expected,
            "passed": False,
        }
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text.startswith("Detailed supervisor comment:"):
            continue
        note_count += 1
        for match in re.finditer(r"(?:^|\s)(\d{1,4})\.\s", text):
            try:
                number = int(match.group(1))
            except (TypeError, ValueError):
                continue
            if not expected_set or number in expected_set:
                represented.add(number)
    missing = sorted(expected_set - represented)
    return {
        "note_count": note_count,
        "expected_finding_numbers": expected,
        "represented_finding_numbers": sorted(represented),
        "missing_finding_numbers": missing,
        "passed": bool(expected) and note_count > 0 and not missing,
    }
