from __future__ import annotations

import io
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "17233A"
BLUE = "4557D5"
MUTED = "667085"
LIGHT = "EEF2FF"
LINE = "D9DFEA"
RED = "A61B1B"
GREEN = "176B3A"
AMBER = "9A6700"

DOMAIN_LABELS = {
    "chapter_one_assessment": "Assessment of Chapter One or the Foundational Chapter",
    "research_problem_and_purpose": "Research Problem, Purpose, Objectives and Questions",
    "literature_and_theoretical_foundation": "Literature and Theoretical Foundation",
    "methodology_and_procedures": "Research Methodology and Procedures",
    "results_or_findings": "Results or Findings",
    "discussion_and_interpretation": "Discussion and Interpretation",
    "conclusions_recommendations_and_contribution": "Conclusions, Recommendations and Contribution",
    "structural_coherence_and_alignment": "Structural Coherence and Cross-Chapter Alignment",
    "academic_writing_and_presentation": "Academic Writing and Presentation",
    "ethics_and_research_integrity": "Ethics and Research Integrity",
    "originality_and_contribution": "Originality and Degree-Level Contribution",
}

JUDGEMENT_LABELS = {
    "strong_and_fully_appropriate": "Strong and Fully Appropriate",
    "appropriate_with_minor_refinement": "Appropriate with Minor Refinement",
    "partly_appropriate_major_revision_required": "Partly Appropriate: Major Revision Required",
    "fundamentally_deficient": "Fundamentally Deficient",
    "not_applicable": "Not Applicable",
}

COVERAGE_LABELS = {
    "fully_assessed": "Fully assessed from supplied evidence",
    "partly_assessed": "Partly assessed from supplied evidence",
    "not_assessed_due_to_retrieval_limit": "Not assessed because source retrieval was insufficient",
    "not_applicable": "Not applicable",
}


STAGE_LABELS = {
    "initial_examination": "Initial Examination",
    "re_examination": "Re-examination",
    "corrected_thesis_verification": "Corrected Thesis Verification",
}


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), LINE)
        borders.append(tag)


def _base_document(title: str, subtitle: str = "") -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for level, size in ((1, 16), (2, 13), (3, 11)):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(MUTED)
    return doc


def _add_info_table(doc: Document, assessment: Dict[str, Any]) -> None:
    meta = assessment.get("assessment_metadata") or {}
    rows = [
        ("Candidate", meta.get("candidate_name")),
        ("Candidate number", meta.get("candidate_number")),
        ("Degree programme", meta.get("degree_programme")),
        ("Academic level", meta.get("academic_level")),
        ("Department", meta.get("department")),
        ("Institution", meta.get("institution")),
        ("Thesis or dissertation title", meta.get("thesis_title")),
        ("Research approach", meta.get("research_approach")),
        ("Examination stage", STAGE_LABELS.get(meta.get("assessment_stage"), _clean(meta.get("assessment_stage")).replace("_", " ").title())),
        ("External examiner", meta.get("examiner_name")),
        ("Examiner affiliation", meta.get("examiner_department")),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.0)
    _borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cells[0], LIGHT)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(_clean(value) or "Not supplied")


def _add_paragraph_section(doc: Document, heading: str, value: Any, level: int = 1) -> None:
    doc.add_heading(heading, level=level)
    doc.add_paragraph(_clean(value) or "No substantive assessment was recorded.")


def _add_bullets(doc: Document, heading: str, values: Iterable[Any]) -> None:
    cleaned = [_clean(value) for value in values if _clean(value)]
    if not cleaned:
        return
    p = doc.add_paragraph()
    p.add_run(heading).bold = True
    for value in cleaned:
        doc.add_paragraph(value, style="List Bullet")


def _add_source_assurance(doc: Document, assessment: Dict[str, Any]) -> None:
    quality = assessment.get("quality_assurance") or {}
    manifest = assessment.get("source_manifest") or {}
    if not quality and not manifest:
        return
    doc.add_heading("Source Coverage and Evidence Assurance", level=1)
    toc = manifest.get("toc_reconciliation") or {}
    rows = [
        ("Evidence audit", _clean(quality.get("audit_status")).replace("_", " ").title()),
        ("Source coverage", _clean(quality.get("coverage_status") or manifest.get("coverage_status")).replace("_", " ").title()),
        ("Coverage score", quality.get("coverage_score") or manifest.get("coverage_score")),
        ("Detected chapters", ", ".join(str(value) for value in (quality.get("detected_chapters") or manifest.get("detected_chapters") or []))),
        ("Table of contents reconciliation", _clean(toc.get("status")).replace("_", " ").title()),
        ("Extracted words", f"{int(quality.get('word_count') or manifest.get('word_count') or 0):,}"),
        ("Extracted tables", quality.get("table_count") if quality.get("table_count") is not None else manifest.get("table_count")),
        ("Evidence references used", quality.get("evidence_reference_count")),
        ("Derivative findings excluded", quality.get("derivative_findings_filtered")),
        ("Contradictions after audit", quality.get("presence_contradiction_count")),
    ]
    table = doc.add_table(rows=0, cols=2)
    _borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        _shade(cells[0], LIGHT)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(_clean(value) or "Not available")
    warnings = manifest.get("extraction_warnings") or []
    _add_bullets(doc, "Extraction warnings", warnings)


def _add_domain(doc: Document, label: str, domain: Dict[str, Any]) -> None:
    doc.add_heading(label, level=1)
    judgement = JUDGEMENT_LABELS.get(
        domain.get("judgement"),
        _clean(domain.get("judgement")).replace("_", " ").title(),
    )
    p = doc.add_paragraph()
    p.add_run("Examiner’s judgement: ").bold = True
    run = p.add_run(judgement)
    run.bold = True
    if domain.get("judgement") == "fundamentally_deficient":
        run.font.color.rgb = RGBColor.from_string(RED)
    elif domain.get("judgement") == "strong_and_fully_appropriate":
        run.font.color.rgb = RGBColor.from_string(GREEN)
    else:
        run.font.color.rgb = RGBColor.from_string(AMBER)
    coverage = COVERAGE_LABELS.get(
        domain.get("coverage_status"),
        _clean(domain.get("coverage_status")).replace("_", " ").title(),
    )
    if coverage:
        p2 = doc.add_paragraph()
        p2.add_run("Evidence coverage: ").bold = True
        p2.add_run(coverage)
    evidence_ids = [
        _clean(value) for value in (domain.get("evidence_ids") or []) if _clean(value)
    ]
    if evidence_ids:
        p3 = doc.add_paragraph()
        p3.add_run("Evidence references: ").bold = True
        p3.add_run(", ".join(evidence_ids))
    doc.add_paragraph(_clean(domain.get("assessment")))
    _add_bullets(doc, "Strengths", domain.get("strengths") or [])
    _add_bullets(doc, "Concerns", domain.get("concerns") or [])
    _add_bullets(doc, "Required corrections", domain.get("required_corrections") or [])


def _add_evidence_register(doc: Document, assessment: Dict[str, Any]) -> None:
    entries = assessment.get("source_evidence_register") or []
    if not entries:
        return
    doc.add_heading("Source Evidence Register", level=1)
    doc.add_paragraph(
        "The following source excerpts support the evidence identifiers used in the "
        "assessment. They are provided for traceability and should be read against "
        "the complete thesis before the examiner signs the report."
    )
    table = doc.add_table(rows=1, cols=4)
    _borders(table)
    headers = ("Evidence ID", "Location", "Heading", "Source excerpt")
    for cell, label in zip(table.rows[0].cells, headers):
        _shade(cell, BLUE)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for item in entries:
        cells = table.add_row().cells
        location_parts = []
        if item.get("chapter_number") is not None:
            location_parts.append(f"Chapter {item.get('chapter_number')}")
        if item.get("page") is not None:
            location_parts.append(f"page {item.get('page')}")
        if item.get("paragraph") is not None:
            location_parts.append(f"paragraph {item.get('paragraph')}")
        if item.get("table_index") is not None:
            table_location = f"table {item.get('table_index')}"
            if item.get("table_row") is not None:
                table_location += f", row {item.get('table_row')}"
            location_parts.append(table_location)
        values = (
            item.get("id"),
            ", ".join(location_parts) or "Document source",
            item.get("heading"),
            item.get("text"),
        )
        for cell, value in zip(cells, values):
            run = cell.paragraphs[0].add_run(_clean(value) or "Not available")
            run.font.size = Pt(8.5)


def _add_recommendation_box(doc: Document, assessment: Dict[str, Any]) -> None:
    table = doc.add_table(rows=2, cols=1)
    _borders(table)
    _shade(table.cell(0, 0), BLUE)
    run = table.cell(0, 0).paragraphs[0].add_run("FINAL RECOMMENDATION")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    value = table.cell(1, 0)
    value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = value.paragraphs[0].add_run(_clean(assessment.get("recommendation_label")))
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(NAVY)


def _bytes(doc: Document) -> bytes:
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_external_examination_report(review: Dict[str, Any]) -> bytes:
    assessment = review.get("external_assessment") or {}
    doc = _base_document(
        "External Examination Report",
        "Thesis or Dissertation Assessment",
    )
    _add_info_table(doc, assessment)
    _add_source_assurance(doc, assessment)
    _add_paragraph_section(doc, "1. Summary of the Study", assessment.get("study_summary"))
    _add_paragraph_section(doc, "2. Overall Academic Judgement", assessment.get("overall_academic_judgement"))
    _add_paragraph_section(doc, "3. Degree-Level Standard", assessment.get("degree_standard_judgement"))

    for field in assessment.get("domain_order") or DOMAIN_LABELS:
        domain = assessment.get(field) or {}
        if domain:
            _add_domain(doc, DOMAIN_LABELS.get(field, field.replace("_", " ").title()), domain)

    _add_evidence_register(doc, assessment)

    doc.add_heading("Major Strengths", level=1)
    strengths = assessment.get("major_strengths") or []
    if strengths:
        for item in strengths:
            doc.add_paragraph(_clean(item), style="List Bullet")
    else:
        doc.add_paragraph("No major strengths were separately recorded.")

    _add_paragraph_section(doc, "Publication Potential", assessment.get("publication_potential"))
    stage = (assessment.get("assessment_metadata") or {}).get("assessment_stage")
    if stage in {"re_examination", "corrected_thesis_verification"}:
        _add_paragraph_section(
            doc,
            "Assessment of Earlier Examiner Corrections",
            assessment.get("corrections_verification_assessment"),
        )
    _add_recommendation_box(doc, assessment)
    _add_paragraph_section(doc, "Rationale for Recommendation", assessment.get("recommendation_rationale"))
    _add_bullets(doc, "Priority corrections before award", assessment.get("priority_corrections_before_award") or [])

    doc.add_heading("Examiner Declaration", level=1)
    doc.add_paragraph(_clean(assessment.get("examiner_declaration")))
    doc.add_paragraph("Examiner’s signature: ______________________________________")
    doc.add_paragraph("Date: ______________________________")
    return _bytes(doc)


def build_corrections_schedule(review: Dict[str, Any]) -> bytes:
    assessment = review.get("external_assessment") or {}
    doc = _base_document("Schedule of Required Corrections")
    _add_info_table(doc, assessment)
    corrections = assessment.get("corrections") or []
    table = doc.add_table(rows=1, cols=8)
    _borders(table)
    headers = ["No.", "Class", "Chapter/section", "Location", "Evidence", "Issue", "Required correction", "Verification"]
    for cell, label in zip(table.rows[0].cells, headers):
        _shade(cell, BLUE)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for item in corrections:
        cells = table.add_row().cells
        values = [
            item.get("number"), item.get("classification", "").title(),
            item.get("chapter_or_section"), item.get("location"),
            ", ".join(_clean(value) for value in (item.get("evidence_ids") or []) if _clean(value)),
            item.get("issue"), item.get("required_correction"), item.get("verification_by"),
        ]
        for cell, value in zip(cells, values):
            cell.paragraphs[0].add_run(_clean(value))
        if item.get("classification") in {"critical", "major"}:
            _shade(cells[1], "FDECEC" if item.get("classification") == "critical" else "FFF4E5")
        if _clean(item.get("rationale")):
            p = cells[6].add_paragraph()
            r = p.add_run("Rationale: " + _clean(item.get("rationale")))
            r.italic = True
            r.font.size = Pt(8.5)
    if not corrections:
        doc.add_paragraph("No formal corrections were recorded.")
    _add_evidence_register(doc, assessment)
    doc.add_paragraph()
    doc.add_paragraph("Corrections verified by: ______________________________________")
    doc.add_paragraph("Signature: ______________________________   Date: ______________________________")
    return _bytes(doc)


def build_confidential_recommendation(review: Dict[str, Any]) -> bytes:
    assessment = review.get("external_assessment") or {}
    doc = _base_document(
        "Confidential Recommendation to the University",
        "Not for release to the candidate",
    )
    _add_info_table(doc, assessment)
    _add_source_assurance(doc, assessment)
    _add_evidence_register(doc, assessment)
    _add_recommendation_box(doc, assessment)
    _add_paragraph_section(doc, "Rationale", assessment.get("recommendation_rationale"))
    _add_paragraph_section(doc, "Confidential Comments", assessment.get("confidential_comments_to_university"))
    _add_paragraph_section(doc, "Confidence in Recommendation", _clean(assessment.get("recommendation_confidence")).title())
    _add_paragraph_section(doc, "Viva Voce Recommendation", _clean(assessment.get("viva_recommendation")).replace("_", " ").title())
    _add_paragraph_section(doc, "Verification of Corrections", assessment.get("corrections_verification_by"))
    doc.add_heading("Examiner Confirmation", level=1)
    doc.add_paragraph("I confirm that this confidential recommendation represents my independent academic judgement.")
    doc.add_paragraph("Examiner’s signature: ______________________________________")
    doc.add_paragraph("Date: ______________________________")
    return _bytes(doc)


def build_oral_examination_questions(review: Dict[str, Any]) -> bytes:
    assessment = review.get("external_assessment") or {}
    doc = _base_document("Oral Examination Question Bank")
    _add_info_table(doc, assessment)
    questions = assessment.get("oral_examination_questions") or []
    if not questions:
        doc.add_paragraph("No oral examination questions were generated.")
    for index, item in enumerate(questions, start=1):
        doc.add_heading(f"Question {index}: {_clean(item.get('category'))}", level=2)
        p = doc.add_paragraph()
        p.add_run(_clean(item.get("question"))).bold = True
        if _clean(item.get("rationale")):
            p2 = doc.add_paragraph()
            p2.add_run("Examiner’s purpose: ").bold = True
            p2.add_run(_clean(item.get("rationale")))
    return _bytes(doc)
